# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx")

print("=== All Headings ===")
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith("Heading"):
        print(f"P{i} [{p.style.name}]: {p.text}")

print("\n=== 3.2 area (verify 3.4 merged) ===")
for i in range(96, 135):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    text = p.text[:80]
    has_img = any(run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')) for run in p.runs)
    style = p.style.name if p.style else "?"
    prefix = "[IMG] " if has_img else ""
    if style.startswith("Heading") or has_img or "表3." in text or "图3." in text:
        print(f"P{i} [{style}]: {prefix}{text}")

print("\n=== 6.4 text ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if "6.4" in p.text and p.style and p.style.name.startswith("Heading"):
        for j in range(i, min(i+5, len(doc.paragraphs))):
            print(f"P{j}: {doc.paragraphs[j].text[:100]}")
        break

print("\n=== 7.1 last para ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if "7.1" in p.text and p.style and p.style.name.startswith("Heading"):
        for j in range(i, min(i+5, len(doc.paragraphs))):
            print(f"P{j}: {doc.paragraphs[j].text[:120]}")
        break

print("\n=== References (first 5) ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if p.text.strip() == "参考文献" and p.style and p.style.name.startswith("Heading"):
        for j in range(i+1, min(i+16, len(doc.paragraphs))):
            text = doc.paragraphs[j].text
            if text.strip():
                print(f"  {text[:80]}")
        break

print("\n=== Citation order in text ===")
seen = set()
for i in range(50, 400):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    refs = re.findall(r'\[(\d+)\]', p.text)
    for r in refs:
        if r not in seen:
            seen.add(r)
            print(f"  [{r}] first at P{i}")
