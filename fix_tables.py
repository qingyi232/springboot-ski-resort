# -*- coding: utf-8 -*-
"""
修改三线表格式:
- 上下边线 1.5pt
- 表头与表体分隔线 1pt
- 表格文字 11pt 宋体居中
- 段前3pt 段后3pt
- 单倍行距
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"
doc = Document(DOC)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Each param is a dict like {'sz': '12', 'val': 'single', 'color': '000000'}"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge, data in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        if data is not None:
            elem = tcBorders.find(qn(edge))
            if elem is None:
                elem = OxmlElement(edge)
                tcBorders.append(elem)
            for k, v in data.items():
                elem.set(qn(f'w:{k}'), v)


def format_three_line_table(table):
    """Format a table as a three-line table."""
    nrows = len(table.rows)
    ncols = len(table.columns)
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Clear all borders first
            no_border = {'val': 'none', 'sz': '0', 'color': '000000'}
            set_cell_border(cell, top=no_border, bottom=no_border, left=no_border, right=no_border)
            
            # Top border for first row (1.5pt = 12 eighth-points)
            if row_idx == 0:
                set_cell_border(cell, top={'val': 'single', 'sz': '12', 'color': '000000'})
                # Bottom of header row (1pt = 8 eighth-points)
                set_cell_border(cell, bottom={'val': 'single', 'sz': '8', 'color': '000000'})
            
            # Bottom border for last row (1.5pt)
            if row_idx == nrows - 1:
                set_cell_border(cell, bottom={'val': 'single', 'sz': '12', 'color': '000000'})
            
            # Format cell text: 11pt SimSun, centered, spacing 3pt before/after
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Paragraph spacing
                pPr = para._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._element.insert(0, pPr)
                
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                spacing.set(qn('w:before'), str(int(3 * 20)))  # 3pt
                spacing.set(qn('w:after'), str(int(3 * 20)))   # 3pt
                spacing.set(qn('w:line'), '240')                # single line spacing
                spacing.set(qn('w:lineRule'), 'auto')
                
                # Font: 11pt SimSun
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:eastAsia'), '宋体')
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


# Apply to all tables in the document
for i, table in enumerate(doc.tables):
    format_three_line_table(table)
    print(f"  Table {i} formatted")

doc.save(DOC)
print(f"Saved: {DOC}")
