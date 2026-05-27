# -*- coding: utf-8 -*-
"""
检查图中 9 条修改要求是否全部完成
"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('5.26.2.docx')

print("=" * 70)
print("要求1: 3.4要和3.2的内容结合起来，把3.4的图放到3.2对应的文字下边")
print("=" * 70)
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style.startswith('Heading') and ('3.2' in text or '3.3' in text or '3.4' in text):
        print(f"  [{i}] [{style}] {text}")
    if style == 'Heading 1' and '第4章' in text:
        break

# Check if 3.4 exists as a separate section
has_34 = False
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '3.4' in p.text:
        has_34 = True
        print(f"\n  发现 3.4 节: [{i}] {p.text.strip()}")

if not has_34:
    print("\n  3.4 节不存在（可能已合并）")

# Check 3.2 sub-sections and their images
print("\n  3.2 的子节和图片:")
in_32 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '3.2' in text and style.startswith('Heading'):
        in_32 = True
    if in_32 and style == 'Heading 2' and '3.3' in text:
        break
    if in_32:
        if style.startswith('Heading 3'):
            print(f"    [{i}] [{style}] {text}")
        if '图3' in text or '图 3' in text:
            print(f"    [{i}] [图题] {text}")
        drawings = p._element.findall('.//' + qn('w:drawing'))
        if drawings:
            print(f"    [{i}] [IMG] (图片)")

print("\n" + "=" * 70)
print("要求2: 第5章要按角色分类写（1.管理员 2.用户 3.工作人员 4.教练）")
print("       5.1不用单独开一小节，可以保留5.7")
print("=" * 70)
in_ch5 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and ('第5章' in text or '第五章' in text):
        in_ch5 = True
        print(f"  [{i}] [{style}] {text}")
    elif in_ch5 and style == 'Heading 1':
        break
    elif in_ch5 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n" + "=" * 70)
print("要求3: 文献要按顺序（找不到文献6）")
print("=" * 70)
# Check for reference citations
refs_found = set()
for i, p in enumerate(doc.paragraphs):
    text = p.text
    import re
    citations = re.findall(r'\[(\d+)\]', text)
    for c in citations:
        refs_found.add(int(c))
refs_sorted = sorted(refs_found)
print(f"  文中引用的文献编号: {refs_sorted}")
missing = []
if refs_sorted:
    for n in range(1, max(refs_sorted) + 1):
        if n not in refs_found:
            missing.append(n)
if missing:
    print(f"  缺失的文献编号: {missing}")
else:
    print(f"  文献编号连续，无缺失")

print("\n" + "=" * 70)
print("要求5: 4.5要放到4.4前边，把4.4应该是最后三段跟4.5的业务流程图对应起来")
print("=" * 70)
in_ch4 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第4章' in text:
        in_ch4 = True
        print(f"  [{i}] [{style}] {text}")
    elif in_ch4 and style == 'Heading 1':
        break
    elif in_ch4 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n" + "=" * 70)
print("要求6: 第5章 5.1-5.6 这几节的字有点太多了")
print("=" * 70)
in_ch5 = False
current_section = ""
section_chars = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and ('第5章' in text or '第五章' in text):
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style.startswith('Heading 2'):
        current_section = text[:20]
        section_chars[current_section] = 0
    elif in_ch5 and current_section and not style.startswith('Heading'):
        section_chars[current_section] = section_chars.get(current_section, 0) + len(text)

for sec, chars in section_chars.items():
    print(f"  {sec}: {chars} 字符")

print("\n" + "=" * 70)
print("要求8: 6.4口语化检查")
print("=" * 70)
in_64 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
        print(f"  [{i}] [{style}] {text}")
    elif in_64 and style.startswith('Heading') and '6.4' not in text:
        break
    elif in_64 and text:
        # Check for oral expressions
        oral_words = ['还有一个问题', '理论上', '都能', '其实', '比较', '挺', '可能']
        found = [w for w in oral_words if w in text]
        if found:
            print(f"  [{i}] 口语化词: {found}")
            print(f"         {text[:100]}")

print("\n" + "=" * 70)
print("要求9: 7.1最后一段检查 - 笔者 / '这些在小型练习中难以遇到'")
print("=" * 70)
in_71 = False
last_paras_71 = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '7.1' in text and style.startswith('Heading'):
        in_71 = True
        print(f"  [{i}] [{style}] {text}")
    elif in_71 and style.startswith('Heading'):
        break
    elif in_71 and text:
        last_paras_71.append((i, text))

if last_paras_71:
    print(f"\n  7.1 最后一段:")
    idx, txt = last_paras_71[-1]
    print(f"  [{idx}] {txt[:200]}")
    if '笔者' in txt:
        print(f"\n  ⚠ 仍包含'笔者'")
    else:
        print(f"\n  ✓ 不包含'笔者'")
    if '小型练习' in txt:
        print(f"  ⚠ 仍包含'小型练习'")
    elif '课程中' in txt or '在课程' in txt:
        print(f"  ✓ 已改为'课程'相关表述")
