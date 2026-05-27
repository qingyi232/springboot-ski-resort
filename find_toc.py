# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document('5.26.2.docx')

# 搜索目录相关的段落
print("=== 搜索包含'目录'的段落 ===")
for i, p in enumerate(doc.paragraphs):
    if '目录' in p.text or '目 录' in p.text:
        print(f"[{i}] [{p.style.name}] {p.text.strip()[:60]}")

# 搜索TOC样式的段落
print("\n=== 搜索TOC样式的段落 ===")
for i, p in enumerate(doc.paragraphs):
    if 'TOC' in p.style.name or 'toc' in p.style.name.lower():
        print(f"[{i}] [{p.style.name}] {p.text.strip()[:80]}")

# 搜索 fldChar / instrText (Word field codes)
print("\n=== 搜索TOC字段代码 ===")
body = doc.element.body
for i, elem in enumerate(body):
    if elem.tag == qn('w:p'):
        instrTexts = elem.findall('.//' + qn('w:instrText'))
        for instr in instrTexts:
            if instr.text and 'TOC' in instr.text:
                text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
                print(f"  body元素[{i}]: instrText={instr.text.strip()}")

# 检查 sdt (structured document tag，Word TOC 常用)
sdts = body.findall(qn('w:sdt'))
print(f"\n=== 找到 {len(sdts)} 个 sdt 元素 ===")
for idx, sdt in enumerate(sdts):
    alias = sdt.find('.//' + qn('w:alias'))
    tag = sdt.find('.//' + qn('w:tag'))
    print(f"  sdt[{idx}]: alias={alias.get(qn('w:val')) if alias is not None else 'N/A'}, tag={tag.get(qn('w:val')) if tag is not None else 'N/A'}")
    # Check if it has TOC content
    content = sdt.find(qn('w:sdtContent'))
    if content is not None:
        paras = content.findall(qn('w:p'))
        print(f"    内含 {len(paras)} 个段落")
        for j, p in enumerate(paras[:5]):
            text = ''.join(r.text or '' for r in p.findall('.//' + qn('w:r')))[:60]
            if text.strip():
                print(f"    [{j}] {text}")
