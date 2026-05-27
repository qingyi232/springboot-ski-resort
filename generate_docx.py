# -*- coding: utf-8 -*-
"""生成论文docx文件，带排版"""

import subprocess
import sys

# 确保安装python-docx
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

import os

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

# ============ 辅助函数 ============
def add_title(text, level=0):
    """添加标题。level=0为论文大标题，1为章标题，2为节标题，3为小节标题"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(24)
    elif level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)

def add_body(text):
    """添加正文段落，首行缩进2字符"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_body_no_indent(text):
    """添加正文段落，不缩进"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table_title(text):
    """添加表格标题（居中，五号字）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

def add_figure_title(text):
    """添加图片标题（居中，五号字）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)

def add_keywords(label, text):
    """添加关键词行"""
    p = doc.add_paragraph()
    run1 = p.add_run(label)
    run1.font.size = Pt(12)
    run1.font.bold = True
    run1.font.name = '宋体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run(text)
    run2.font.size = Pt(12)
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def create_table(headers, rows):
    """创建带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()  # 表后空行
    return table

def create_table_left(headers, rows):
    """创建左对齐内容的表格（说明列左对齐）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# ============ 页面设置 ============
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============ 正文开始 ============

# 论文标题
add_title('基于Spring Boot的飞跃滑雪场管理系统的设计与实现')

# ========== 摘要 ==========
add_title('摘  要', level=1)

add_body('近些年来，随着"带动三亿人参与冰雪运动"口号的提出和北京冬奥会的举办，我国冰雪运动产业得到了快速的发展。越来越多的人开始接触滑雪这项运动，滑雪场的数量也在不断增加。但在实际运营过程中，不少中小型滑雪场仍然依赖人工方式来管理日常业务，存在着效率低下、信息不透明、数据难以统计等问题。针对这些现状，本文设计并实现了一套基于Spring Boot框架的飞跃滑雪场管理系统。')

add_body('本系统采用前后端分离的B/S架构进行开发。后端使用Spring Boot 2.7作为核心框架，搭配MyBatis进行数据持久化操作，数据存储采用MySQL 8.0关系型数据库，并引入Redis缓存来提升数据访问速度；在安全方面，采用了JWT（JSON Web Token）机制来完成用户身份认证和权限管控。前端部分则基于Vue 3框架搭建，结合Element Plus组件库构建用户界面，利用ECharts实现运营数据的可视化展示，通过Axios库与后端进行HTTP通信。')

add_body('系统按照使用角色划分为管理员端、教练端和普通用户端三个入口。功能上涵盖了用户管理、雪具租赁管理、教练预约管理、滑雪课程管理、场地管理、装备商品销售管理以及数据统计分析等七大核心业务模块，共设计了16张数据库表来存储业务数据。经过功能测试与性能测试，系统运行稳定，各项功能均能正常使用，基本满足了中小型滑雪场日常运营管理的需要。')

add_keywords('关键词：', 'Spring Boot；滑雪场管理；前后端分离；Vue.js；MySQL')

# ========== Abstract ==========
add_title('Abstract', level=1)

add_body('In recent years, with the proposal of "driving 300 million people to participate in ice and snow sports" and the hosting of the Beijing Winter Olympics, the ice and snow sports industry in China has developed rapidly. More and more people are getting into skiing, and the number of ski resorts is increasing. However, in actual operation, many small and medium-sized ski resorts still rely on manual methods to manage daily business, which leads to low efficiency and difficulty in data statistics. In response to these problems, this paper designs and implements a ski resort management system based on the Spring Boot framework.')

add_body('The system adopts a front-end and back-end separated B/S architecture. The back-end uses Spring Boot 2.7 as the core framework, MyBatis for data persistence, MySQL 8.0 for data storage, and Redis for caching. JWT is used for user authentication and authorization. The front-end is built on Vue 3 framework, combined with Element Plus component library and ECharts for data visualization.')

add_body('The system is divided into three portals: administrator, coach, and ordinary user. It covers seven core business modules including user management, equipment rental, coach booking, course management, venue management, product sales, and data statistics. A total of 16 database tables are designed. After testing, the system runs stably and meets the daily management needs of small and medium-sized ski resorts.')

add_keywords('Keywords: ', 'Spring Boot; Ski Resort Management; Front-end and Back-end Separation; Vue.js; MySQL')

# ========== 第一章 绪论 ==========
doc.add_page_break()
add_title('第一章 绪论', level=1)

add_title('1.1 研究背景与意义', level=2)

add_body('冰雪运动在我国有着悠久的历史，尤其是在东北和华北地区，滑雪一直是冬季比较受欢迎的一项体育活动。2015年北京申办冬奥会成功以后，国家层面提出了"带动三亿人参与冰雪运动"的目标，各地的滑雪场如雨后春笋般建设起来。根据相关行业报告，2023年全国滑雪场数量已经超过了700家，年滑雪人次也突破了2000万。这个行业正在经历一个比较快速的增长期。')

add_body('不过从实际情况来看，很多滑雪场在经营管理方面还是比较传统的。有的雪场租赁雪具还是用纸质登记本来记录，有的雪场教练预约通过电话或者微信来沟通，有的雪场商品销售没有系统化的库存管理。这些做法在雪场规模较小、客流量不大的时候还勉强能够应付，但随着业务量的增长，就会暴露出很多问题：比如雪具库存数量不清楚导致超租的情况、教练排班冲突、收入数据统计困难等等。')

add_body('信息化管理是解决上述问题的一个有效途径。通过开发一套专门针对滑雪场业务的管理信息系统，可以将各项业务流程进行规范化和数字化，让管理者能够实时掌握运营状况，也让顾客获得更好的消费体验。这也是本文选择"滑雪场管理系统"作为毕业设计课题的出发点。')

add_body('本文所设计的飞跃滑雪场管理系统，定位于服务中小型滑雪场的日常运营，力求在满足基本管理需求的前提下，做到操作简便、功能实用。系统的开发不仅是对大学四年所学专业知识的一次综合运用，也希望能够为滑雪场行业的信息化建设提供一些参考。')

add_title('1.2 国内外研究现状', level=2)

add_body('在国外，欧洲和北美地区的大型滑雪度假村很早就开始使用信息化管理工具了。像瑞士、奥地利的一些老牌雪场，基本都配备了完善的票务系统、租赁管理系统和CRM（客户关系管理）系统。美国的Vail Resorts集团甚至开发了自己的EpicMix平台，把滑雪数据追踪、社交分享等功能都整合到了一起。不过这些系统大多是商业化的定制产品，价格比较昂贵，也不太适合国内中小型雪场直接使用。')

add_body('在国内，滑雪场管理系统的发展相对来说起步比较晚。早期的一些系统主要是基于C/S架构开发的桌面应用，功能上以票务管理和财务管理为主，对租赁、预约等业务的支持不够完善。随着Web技术的发展，越来越多的开发者开始采用B/S架构来设计滑雪场管理系统，这样用户通过浏览器就可以访问系统，部署和维护也更加方便。')

add_body('从技术框架的选择上来看，目前主流的Java Web开发方案基本上都是以Spring Boot为核心的。Spring Boot简化了Spring框架的配置过程，提供了很多开箱即用的功能组件，让开发者可以把更多精力放在业务逻辑的实现上。前端方面，Vue.js因为学习门槛相对较低、文档齐全、生态丰富等优点，在国内Web开发领域的使用率很高。本文的系统正是基于这两个主流框架来进行开发的。')

add_title('1.3 主要研究内容', level=2)

add_body('本文围绕飞跃滑雪场管理系统，主要开展了以下几个方面的工作：')
add_body('（1）对滑雪场的实际业务需求进行了调研和分析，确定了系统应具备的功能模块和使用角色，绘制了用例图来描述各角色与系统之间的交互关系。')
add_body('（2）完成了系统的总体设计，包括技术架构的选型、系统功能模块的划分、数据库的概念设计和逻辑设计等工作，共设计了16张数据库表。')
add_body('（3）按照设计方案实现了系统的各项功能，后端共编写了17个Controller类、12个实体类和对应的Service及Mapper层代码，前端开发了30多个Vue页面组件。')
add_body('（4）对系统进行了功能测试和性能测试，验证了系统功能的正确性和运行的稳定性。')

add_title('1.4 论文组织结构', level=2)

add_body('本论文共包含七章内容，具体安排如下：')
add_body('第一章是绪论部分，介绍了研究背景、国内外研究现状和主要研究内容。第二章对系统开发中用到的关键技术进行了简要介绍。第三章是需求分析，阐述了系统的功能需求和非功能需求。第四章是系统设计，包括总体架构设计、功能模块设计和数据库设计。第五章是系统实现，展示了各主要功能模块的具体实现过程和效果。第六章是系统测试，给出了测试方案和测试结果。第七章是总结与展望，对全文工作进行了总结并提出了后续改进方向。')

# ========== 第二章 相关技术介绍 ==========
doc.add_page_break()
add_title('第二章 相关技术介绍', level=1)

add_title('2.1 Spring Boot框架', level=2)
add_body('Spring Boot是由Pivotal团队提供的一个用于简化Spring应用开发的框架。在传统的Spring项目中，开发者需要编写大量的XML配置文件，这在项目初期会消耗不少时间。Spring Boot通过"约定优于配置"的理念，大幅度减少了配置的工作量。开发者只需要引入相应的Starter依赖，框架就会自动完成大部分配置工作。')
add_body('本系统使用的是Spring Boot 2.7.18版本。选择这个版本主要是因为它比较成熟稳定，社区资源也很丰富，遇到问题容易找到解决方案。在项目中，Spring Boot主要承担了以下几个方面的职责：一是作为Web服务的容器，内嵌了Tomcat服务器，不需要额外部署；二是管理各个组件之间的依赖关系，通过IoC容器实现了松耦合；三是提供了统一的配置管理机制，通过application.yml文件就可以集中管理各种参数。')

add_title('2.2 MyBatis持久层框架', level=2)
add_body('MyBatis是一个比较流行的Java持久层框架，它的定位介于JDBC和Hibernate之间。和JDBC相比，MyBatis封装了底层的数据库连接和结果集映射操作，开发者不用再写那些重复的代码；和Hibernate相比，MyBatis给了开发者更大的SQL控制权，可以根据实际情况编写最优化的SQL语句。')
add_body('在本项目中，MyBatis的使用方式是通过Mapper接口加XML映射文件来完成的。每一个数据库表对应一个Mapper接口和一个XML文件，接口中定义操作方法，XML中编写具体的SQL语句。比如用户表对应的就是UserMapper接口和UserMapper.xml文件。项目中使用的mybatis-spring-boot-starter版本是2.3.1，这个版本和Spring Boot 2.7的兼容性比较好。')

add_title('2.3 MySQL数据库', level=2)
add_body('MySQL是目前全球使用最广泛的开源关系型数据库之一。它性能优良、运行稳定，而且免费使用，非常适合中小型项目的数据存储需求。本系统使用的是MySQL 8.0版本，选择这个版本主要考虑到它在字符集支持方面做了改进，默认就采用utf8mb4编码，可以很好地存储中文字符和各种特殊符号。')
add_body('系统的数据库名称为ski_resort，共创建了16张业务表，使用InnoDB作为存储引擎。InnoDB支持事务操作和外键约束，在数据一致性方面有比较好的保障。在表的设计上，所有表都采用了BIGINT类型的自增主键，并且设置了create_time和update_time字段来记录数据的创建和更新时间，同时使用is_deleted字段来实现逻辑删除，避免了物理删除带来的数据不可恢复的问题。')

add_title('2.4 Redis缓存', level=2)
add_body('Redis是一个开源的内存数据存储系统，通常被用作缓存、消息中间件等。它的最大特点就是速度快，因为数据都存储在内存中，读写性能远远超过传统的关系型数据库。')
add_body('在本系统中，Redis主要用于缓存一些频繁访问但不经常变化的数据，比如系统的统计概览数据、热门排行榜数据等。通过引入spring-boot-starter-data-redis依赖，可以很方便地在Spring Boot项目中使用Redis。项目中还对Redis的序列化方式进行了自定义配置，使用Jackson来替代默认的JDK序列化，这样存储在Redis中的数据可读性更好，调试的时候也更加方便。')

add_title('2.5 JWT身份认证', level=2)
add_body('JWT全称是JSON Web Token，是一种用于在各方之间安全传递信息的开放标准。在本系统中，JWT被用于实现用户的身份认证。具体的工作流程是这样的：用户登录成功后，服务端会生成一个包含用户ID、用户名和角色信息的Token返回给前端；前端收到Token后将其保存在本地存储中，之后每次发送请求时都在HTTP头部携带这个Token；服务端收到请求后会验证Token的有效性，从而确认用户的身份。')
add_body('本项目使用的JWT库是jjwt，版本为0.11.5。Token的有效期设置为7天，签名算法采用的是HS256。在JwtUtil工具类中，封装了Token的生成、解析、验证等方法，供Controller层直接调用。')

add_title('2.6 Vue.js前端框架', level=2)
add_body('Vue.js是一个用于构建用户界面的渐进式JavaScript框架，由尤雨溪在2014年发布。Vue的核心特性包括响应式数据绑定、组件化开发、虚拟DOM等。本系统前端使用的是Vue 3版本，相比Vue 2，它在性能、TypeScript支持和Composition API等方面都有了明显的改进。')
add_body('前端项目使用Vite作为构建工具，相比传统的Webpack，Vite在开发环境下的启动速度更快，热更新也更加及时。UI组件库选择了Element Plus，这是Element UI针对Vue 3推出的升级版本，提供了丰富的预构建组件，比如表格、表单、对话框、分页器等，可以大幅提升前端开发效率。状态管理方面使用了Pinia，它是Vue官方推荐的新一代状态管理工具，用来在不同组件之间共享用户登录状态等全局数据。数据可视化部分使用了ECharts图表库来绑制各类统计图表。')

add_title('2.7 其他相关技术', level=2)
add_body('除了上面提到的核心技术之外，本系统还用到了一些辅助性的技术和工具：')
add_body('（1）Lombok：一个Java工具库，通过注解的方式自动生成实体类的getter、setter、构造方法等模板代码，减少了手工编写的工作量。')
add_body('（2）Knife4j：基于Swagger增强的API文档生成工具，可以自动根据Controller中的注解生成在线的接口文档，方便前后端开发人员进行接口联调。')
add_body('（3）Spring Security：Spring生态中的安全框架，在本项目中主要使用了它提供的BCryptPasswordEncoder来对用户密码进行加密存储。')
add_body('（4）Axios：一个基于Promise的HTTP客户端库，用于前端向后端发送Ajax请求。在项目中对Axios进行了二次封装，统一处理了请求拦截（自动添加Token）和响应拦截（统一处理错误码）。')

# ========== 第三章 系统需求分析 ==========
doc.add_page_break()
add_title('第三章 系统需求分析', level=1)

add_title('3.1 系统概述', level=2)
add_body('飞跃滑雪场管理系统是一个面向中小型滑雪场运营管理需求的Web应用系统。系统的使用者按角色划分为三类：管理员、教练和普通用户。管理员拥有系统的最高权限，可以管理所有业务数据；教练可以查看和管理自己的课程及学员预约；普通用户则可以浏览滑雪场的各项服务信息，进行雪具租赁、教练预约、课程报名、装备购买和场地预约等操作。')

add_title('3.2 功能需求分析', level=2)
add_body('经过对滑雪场实际业务的调研和分析，本系统需要实现以下七个功能模块：')

add_title('3.2.1 用户管理模块', level=3)
add_body('用户管理模块是整个系统的基础。该模块需要实现用户注册、登录、个人信息管理、密码修改等功能。系统采用基于角色的权限控制，不同角色的用户登录后看到的菜单和可操作的功能是不一样的。管理员登录后进入后台管理界面，可以查看和管理所有用户的信息，包括增加、修改、删除用户以及重置用户密码等操作。普通用户登录后进入前台页面，可以编辑自己的个人资料、查看订单记录等。')

add_title('3.2.2 雪具租赁管理模块', level=3)
add_body('雪具租赁是滑雪场的核心业务之一。该模块需要管理滑雪场所有雪具的信息，包括单板、双板、雪杖、头盔、雪镜等各种类型。每件雪具都有编号、名称、品牌、型号、尺寸、租赁价格、库存数量等属性信息。管理员可以对雪具进行增删改查操作，还可以查看雪具的库存情况。用户可以浏览可租赁的雪具列表，选择需要的雪具提交租赁订单，系统会自动计算租赁费用。租赁订单有待支付、使用中、已归还、已取消四种状态，管理员可以管理所有的租赁订单。')

add_title('3.2.3 教练预约管理模块', level=3)
add_body('教练预约模块用于管理教练信息和处理用户的预约请求。系统中的教练信息包括姓名、编号、等级（初级/中级/高级/专业）、专长、从业年限、认证信息、课时费、评分等。用户可以浏览教练列表，根据等级和专长选择合适的教练，然后指定日期和时间段提交预约申请。预约订单创建后需要经过确认、完成等流程。教练登录自己的账号后，可以查看被预约的记录并进行确认操作。')

add_title('3.2.4 滑雪课程管理模块', level=3)
add_body('课程管理模块负责滑雪培训课程的发布和报名管理。每门课程都关联了一位授课教练和一个上课场地，课程信息中包含课程名称、类型（初级入门/中级进阶/高级提升）、价格、时长、开课日期、最大学员数、当前报名人数等。管理员可以添加和编辑课程信息，用户可以浏览可报名的课程列表并进行在线报名。系统会控制每门课程的报名人数不超过最大限制。')

add_title('3.2.5 场地管理模块', level=3)
add_body('场地管理模块用于管理滑雪场的各个雪道和练习场地。每个场地有编号、名称、类型（雪道/练习场）、难度等级（初级/中级/高级）、长度、宽度、最大容纳人数等属性。管理员可以查看场地的使用状态，对场地进行开放或维护的状态切换。用户可以查看各场地的实时状态信息，对需要额外付费的场地可以进行在线预订。')

add_title('3.2.6 装备销售管理模块', level=3)
add_body('除了租赁雪具之外，很多滑雪场还会销售滑雪装备和相关商品。本模块需要管理商品的基本信息（名称、分类、品牌、价格、库存等），支持用户浏览商品并下单购买。销售订单有待支付、待发货、已发货、已完成、已取消五种状态，管理员可以对订单进行发货操作，录入快递公司和快递单号。')

add_title('3.2.7 数据统计分析模块', level=3)
add_body('数据统计模块为管理员提供运营数据的汇总展示。主要包括：统计概览（总用户数、总订单数、总收入等）、收入趋势图、订单统计、热门雪具排行、热门教练排行、热门课程排行等。这些数据通过ECharts图表的形式进行可视化展示，帮助管理者直观地了解滑雪场的运营情况，为决策提供数据支持。')

add_title('3.3 非功能需求', level=2)
add_body('（1）性能需求：系统页面的平均响应时间应控制在2秒以内，对于数据量较大的列表查询应支持分页功能。')
add_body('（2）安全需求：用户密码必须加密存储，不能以明文形式保存在数据库中；接口调用需要携带有效的Token进行身份验证。')
add_body('（3）兼容性需求：系统前端应兼容主流浏览器，包括Chrome、Firefox、Edge等。')
add_body('（4）可维护性需求：代码结构清晰，采用分层架构，各层之间职责明确，便于后期维护和功能扩展。')

add_title('3.4 用例分析', level=2)
add_body('根据上面的功能需求分析，按照三种用户角色分别进行用例分析。')

add_title('3.4.1 管理员用例', level=3)
add_body('管理员是系统权限最高的用户，其主要用例如下：')

add_table_title('表3-1 管理员用例描述')
create_table_left(
    ['用例编号', '用例名称', '用例描述'],
    [
        ['UC-A01', '登录系统', '管理员输入用户名和密码登录后台管理界面'],
        ['UC-A02', '用户管理', '查看用户列表，新增、编辑、删除用户，重置密码'],
        ['UC-A03', '雪具管理', '新增、编辑、删除雪具信息，查看库存'],
        ['UC-A04', '教练管理', '新增、编辑、删除教练信息，更新教练状态'],
        ['UC-A05', '课程管理', '发布、编辑、删除课程，更新课程状态'],
        ['UC-A06', '场地管理', '新增、编辑、删除场地信息，切换场地状态'],
        ['UC-A07', '商品管理', '新增、编辑、删除商品信息，管理库存'],
        ['UC-A08', '租赁订单管理', '查看所有租赁订单，处理归还和取消操作'],
        ['UC-A09', '销售订单管理', '查看所有销售订单，进行发货和确认操作'],
        ['UC-A10', '预约管理', '查看教练预约和场地预订记录'],
        ['UC-A11', '数据统计', '查看经营数据概览和各类统计图表'],
    ]
)
add_figure_title('图3-1 管理员用例图（见附图）')

add_title('3.4.2 教练用例', level=3)
add_body('教练是系统中的特殊角色，既有后台管理的部分功能，也需要查看自己相关的业务数据。')
add_table_title('表3-2 教练用例描述')
create_table_left(
    ['用例编号', '用例名称', '用例描述'],
    [
        ['UC-C01', '登录系统', '教练输入账号密码登录后台界面'],
        ['UC-C02', '查看我的课程', '查看自己负责教授的课程列表和详情'],
        ['UC-C03', '管理学员预约', '查看学员对自己的预约记录，进行确认或取消操作'],
        ['UC-C04', '查看Dashboard', '查看与自己相关的统计数据'],
        ['UC-C05', '编辑个人信息', '修改自己的个人资料和密码'],
    ]
)

add_title('3.4.3 普通用户用例', level=3)
add_body('普通用户是系统的主要使用群体，通过前台页面使用系统提供的各项服务。')
add_table_title('表3-3 普通用户用例描述')
create_table_left(
    ['用例编号', '用例名称', '用例描述'],
    [
        ['UC-U01', '注册/登录', '新用户注册账号，已有用户登录系统'],
        ['UC-U02', '浏览滑雪课程', '查看所有可报名的滑雪课程信息'],
        ['UC-U03', '报名课程', '选择感兴趣的课程进行在线报名'],
        ['UC-U04', '租赁雪具', '浏览可租赁的雪具，提交租赁订单'],
        ['UC-U05', '预约教练', '浏览教练信息，选择教练提交预约'],
        ['UC-U06', '购买装备', '在装备商城浏览和购买滑雪装备'],
        ['UC-U07', '预约场地', '查看场地信息，预约练习场地'],
        ['UC-U08', '查看我的订单', '查看自己的租赁订单和购买订单'],
        ['UC-U09', '查看我的预约', '查看自己的教练预约和场地预订'],
        ['UC-U10', '管理个人信息', '修改头像、个人资料、登录密码'],
    ]
)
add_figure_title('图3-2 普通用户用例图（见附图）')

# ========== 第四章 系统设计 ==========
doc.add_page_break()
add_title('第四章 系统设计', level=1)

add_title('4.1 系统总体架构设计', level=2)
add_body('本系统采用前后端分离的B/S架构，整体可以划分为三个层次：表现层（前端）、业务逻辑层（后端）和数据层。')
add_figure_title('图4-1 系统总体架构图（见附图）')
add_body('系统的请求处理流程大致是这样的：用户在浏览器中进行操作，前端Vue应用将操作转化为HTTP请求，通过Axios发送到后端的Controller层；Controller接收到请求后调用Service层来处理业务逻辑，Service层根据需要调用Mapper层来操作数据库；处理完成后将结果逐层返回，最终以JSON格式响应给前端，前端再将数据渲染到页面上。')

add_title('4.2 后端架构设计', level=2)
add_body('后端项目的包结构遵循标准的MVC分层模式，各包的职责如下表所示：')
add_table_title('表4-1 后端项目包结构说明')
create_table_left(
    ['包名', '说明', '主要内容'],
    [
        ['com.ski.resort.controller', '控制层', '17个Controller类，处理HTTP请求'],
        ['com.ski.resort.service', '业务层', '17个Service接口'],
        ['com.ski.resort.service.impl', '业务实现层', '17个ServiceImpl实现类'],
        ['com.ski.resort.mapper', '数据层', '12个Mapper接口'],
        ['com.ski.resort.entity', '实体层', '12个实体类'],
        ['com.ski.resort.dto', '数据传输对象', 'LoginDTO、RegisterDTO等'],
        ['com.ski.resort.common', '公共类', 'Result统一返回类、PageResult分页类'],
        ['com.ski.resort.config', '配置类', 'Security配置、Redis配置、CORS配置等'],
        ['com.ski.resort.util', '工具类', 'JwtUtil、RedisUtil等'],
        ['com.ski.resort.exception', '异常处理', '自定义业务异常类'],
    ]
)
add_body('后端所有接口都遵循RESTful风格设计，使用统一的Result类来包装返回结果。Result类中包含code（状态码）、message（返回消息）、data（返回数据）和timestamp（时间戳）四个字段，这样前端在处理响应数据时可以采用统一的解析方式。')

add_title('4.3 前端架构设计', level=2)
add_body('前端项目基于Vue 3 + Vite搭建，目录结构如下表：')
add_table_title('表4-2 前端项目目录结构')
create_table_left(
    ['目录/文件', '说明'],
    [
        ['src/api/', 'API接口封装，按模块划分文件（18个）'],
        ['src/views/', '页面组件，按功能分目录（33个Vue文件）'],
        ['src/views/admin/', '管理员端页面（Dashboard、各管理页面）'],
        ['src/views/coach/', '教练端页面（我的课程、预约管理）'],
        ['src/views/user/', '用户相关页面（个人中心、我的订单）'],
        ['src/views/equipment/', '雪具相关页面（雪具列表、租赁页面）'],
        ['src/views/course/', '课程相关页面（课程列表、报名页面）'],
        ['src/views/product/', '商品相关页面（商品列表、商城页面）'],
        ['src/layout/', '布局组件（前台布局FrontLayout、后台布局MainLayout）'],
        ['src/router/', '路由配置（含路由守卫和权限控制）'],
        ['src/stores/', 'Pinia状态管理（用户状态）'],
        ['src/utils/', '工具函数（Axios封装）'],
    ]
)
add_body('前端的路由设计分为两大部分：以"/"开头的路由组用于普通用户的前台页面，以"/admin"开头的路由组用于管理员和教练的后台页面。路由守卫中会检查用户是否已登录以及角色权限是否匹配，如果未登录会自动跳转到登录页面，如果角色不匹配则跳转到对应角色的首页。')

add_title('4.4 数据库设计', level=2)

add_title('4.4.1 概念结构设计', level=3)
add_body('根据需求分析，系统涉及的主要实体包括：用户、雪具、教练、课程、场地、商品、租赁订单、销售订单、教练预约、场地预订、课程报名、评价、会员卡和FAQ等。这些实体之间的主要关系如下：')
add_body('一个用户可以有多个租赁订单、多个销售订单、多个教练预约、多个课程报名记录；一个教练可以教授多门课程，可以被多个用户预约；一门课程由一个教练教授，在一个场地上课，可以有多个学员报名；一个销售订单可以包含多个订单明细（商品项）；一个租赁订单可以包含多个租赁明细（雪具项）。')
add_figure_title('图4-2 系统E-R关系简图（见附图）')

add_title('4.4.2 逻辑结构设计', level=3)
add_body('将概念模型转化为关系模型，系统共设计了16张数据库表。下面列出各主要表的结构设计：')

# 用户表
add_table_title('表4-3 用户表（t_user）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '用户ID，自增'],
        ['username', 'VARCHAR(50)', '否', '用户名，唯一'],
        ['password', 'VARCHAR(255)', '否', '密码（BCrypt加密）'],
        ['real_name', 'VARCHAR(100)', '是', '真实姓名'],
        ['phone', 'VARCHAR(20)', '是', '手机号，唯一'],
        ['email', 'VARCHAR(100)', '是', '邮箱'],
        ['avatar', 'VARCHAR(255)', '是', '头像URL'],
        ['gender', 'TINYINT', '是', '性别：0女，1男'],
        ['birthday', 'DATE', '是', '出生日期'],
        ['id_card', 'VARCHAR(18)', '是', '身份证号'],
        ['role', 'VARCHAR(20)', '否', '角色：admin/user/coach'],
        ['status', 'TINYINT', '是', '状态：0禁用，1正常'],
        ['create_time', 'DATETIME', '是', '创建时间'],
        ['update_time', 'DATETIME', '是', '更新时间'],
        ['is_deleted', 'TINYINT', '是', '逻辑删除标志'],
    ]
)

# 雪具表
add_table_title('表4-4 雪具表（t_equipment）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '雪具ID，自增'],
        ['equipment_code', 'VARCHAR(50)', '否', '雪具编号，唯一'],
        ['equipment_name', 'VARCHAR(100)', '否', '雪具名称'],
        ['equipment_type', 'VARCHAR(50)', '否', '类型：单板/双板/雪杖/头盔/雪镜'],
        ['brand', 'VARCHAR(100)', '是', '品牌'],
        ['model', 'VARCHAR(100)', '是', '型号'],
        ['size', 'VARCHAR(50)', '是', '尺寸规格'],
        ['rental_price', 'DECIMAL(10,2)', '否', '租赁价格（元/小时）'],
        ['purchase_price', 'DECIMAL(10,2)', '是', '购入价格'],
        ['status', 'TINYINT', '是', '状态：0维修中，1可用，2已租出'],
        ['stock_quantity', 'INT', '是', '库存数量'],
        ['available_quantity', 'INT', '是', '可用数量'],
        ['image_url', 'VARCHAR(255)', '是', '图片URL'],
        ['description', 'TEXT', '是', '描述信息'],
    ]
)

# 教练表
add_table_title('表4-5 教练表（t_coach）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '教练ID，自增'],
        ['user_id', 'BIGINT', '是', '关联用户ID'],
        ['coach_name', 'VARCHAR(50)', '否', '教练姓名'],
        ['coach_code', 'VARCHAR(50)', '是', '教练编号，唯一'],
        ['coach_level', 'VARCHAR(20)', '否', '等级：初级/中级/高级/专业'],
        ['specialty', 'VARCHAR(200)', '是', '专长'],
        ['experience_years', 'INT', '是', '从业年限'],
        ['certification', 'VARCHAR(200)', '是', '认证信息'],
        ['hourly_rate', 'DECIMAL(10,2)', '否', '课时费（元/小时）'],
        ['rating', 'DECIMAL(3,2)', '是', '评分（1-5）'],
        ['total_students', 'INT', '是', '累计学员数'],
        ['status', 'TINYINT', '是', '状态：0休假，1在职，2离职'],
    ]
)

# 课程表
add_table_title('表4-6 课程表（t_course）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '课程ID，自增'],
        ['course_code', 'VARCHAR(50)', '否', '课程编号，唯一'],
        ['course_name', 'VARCHAR(100)', '否', '课程名称'],
        ['course_type', 'VARCHAR(50)', '否', '类型：初级入门/中级进阶/高级提升'],
        ['coach_id', 'BIGINT', '否', '授课教练ID'],
        ['max_students', 'INT', '否', '最大学员数'],
        ['current_students', 'INT', '是', '当前报名人数'],
        ['course_price', 'DECIMAL(10,2)', '否', '课程价格'],
        ['course_duration', 'DECIMAL(10,2)', '否', '课程时长（小时）'],
        ['course_date', 'DATE', '否', '开课日期'],
        ['start_time', 'TIME', '否', '开始时间'],
        ['end_time', 'TIME', '否', '结束时间'],
        ['venue_id', 'BIGINT', '是', '场地ID'],
        ['status', 'TINYINT', '是', '状态：0未开始，1报名中，2进行中，3已结束'],
    ]
)

# 场地表
add_table_title('表4-7 场地表（t_venue）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '场地ID，自增'],
        ['venue_code', 'VARCHAR(50)', '否', '场地编号，唯一'],
        ['venue_name', 'VARCHAR(100)', '否', '场地名称'],
        ['venue_type', 'VARCHAR(50)', '否', '类型：雪道/练习场'],
        ['difficulty_level', 'VARCHAR(20)', '是', '难度等级：初级/中级/高级'],
        ['length', 'INT', '是', '长度（米）'],
        ['width', 'INT', '是', '宽度（米）'],
        ['max_capacity', 'INT', '是', '最大容纳人数'],
        ['current_capacity', 'INT', '是', '当前人数'],
        ['rental_price', 'DECIMAL(10,2)', '是', '租赁价格（元/小时）'],
        ['status', 'TINYINT', '是', '状态：0维护中，1开放'],
    ]
)

# 商品表
add_table_title('表4-8 商品表（t_product）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '商品ID，自增'],
        ['product_code', 'VARCHAR(50)', '否', '商品编号，唯一'],
        ['product_name', 'VARCHAR(100)', '否', '商品名称'],
        ['category', 'VARCHAR(50)', '否', '分类：滑雪装备/服装/配件'],
        ['brand', 'VARCHAR(100)', '是', '品牌'],
        ['price', 'DECIMAL(10,2)', '否', '销售价格'],
        ['cost_price', 'DECIMAL(10,2)', '是', '成本价格'],
        ['stock_quantity', 'INT', '是', '库存数量'],
        ['sold_quantity', 'INT', '是', '已售数量'],
        ['status', 'TINYINT', '是', '状态：0下架，1上架'],
        ['is_hot', 'TINYINT', '是', '是否热门'],
    ]
)

# 租赁订单表
add_table_title('表4-9 租赁订单表（t_rental_order）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '订单ID，自增'],
        ['order_no', 'VARCHAR(50)', '否', '订单号，唯一'],
        ['user_id', 'BIGINT', '否', '用户ID'],
        ['equipment_id', 'BIGINT', '否', '雪具ID'],
        ['equipment_name', 'VARCHAR(100)', '是', '雪具名称'],
        ['rental_start_time', 'DATETIME', '否', '租赁开始时间'],
        ['rental_end_time', 'DATETIME', '是', '计划归还时间'],
        ['actual_return_time', 'DATETIME', '是', '实际归还时间'],
        ['rental_price', 'DECIMAL(10,2)', '否', '租赁单价'],
        ['total_amount', 'DECIMAL(10,2)', '否', '总金额'],
        ['deposit', 'DECIMAL(10,2)', '是', '押金'],
        ['order_status', 'TINYINT', '是', '状态：0待支付，1使用中，2已归还，3已取消'],
        ['payment_status', 'TINYINT', '是', '支付状态：0未支付，1已支付'],
        ['payment_method', 'VARCHAR(20)', '是', '支付方式'],
    ]
)

# 销售订单表
add_table_title('表4-10 销售订单表（t_sales_order）结构')
create_table_left(
    ['字段名', '数据类型', '可否为空', '说明'],
    [
        ['id', 'BIGINT', '否（主键）', '订单ID，自增'],
        ['order_no', 'VARCHAR(50)', '否', '订单号，唯一'],
        ['user_id', 'BIGINT', '否', '用户ID'],
        ['total_amount', 'DECIMAL(10,2)', '否', '订单总金额'],
        ['actual_amount', 'DECIMAL(10,2)', '否', '实付金额'],
        ['status', 'TINYINT', '是', '状态：0待支付~4已取消'],
        ['payment_status', 'TINYINT', '是', '支付状态'],
        ['shipping_address', 'VARCHAR(255)', '是', '收货地址'],
        ['express_company', 'VARCHAR(50)', '是', '快递公司'],
        ['express_no', 'VARCHAR(50)', '是', '快递单号'],
    ]
)

add_body('除上述主要表之外，系统还包括以下辅助表：销售订单明细表（t_sales_item）用于记录销售订单中每项商品的数量和价格；教练预约表（t_coach_booking）记录用户对教练的预约信息；场地预订表（t_venue_booking）记录场地的预订情况；课程报名表（t_course_enrollment）记录学员的课程报名数据；评价表（t_evaluation）记录用户对教练、课程、商品的评价打分；会员卡表（t_member_card）存储会员卡信息；租赁明细表（t_rental_item）记录租赁订单的明细；FAQ表（t_faq）存储常见问题与解答。')

add_title('4.5 接口设计', level=2)
add_body('系统后端采用RESTful风格设计API接口，所有接口统一返回Result格式的JSON数据。下面列出各模块的主要接口：')
add_table_title('表4-11 主要API接口列表')
create_table_left(
    ['模块', '请求方式', '接口路径', '功能说明'],
    [
        ['用户', 'POST', '/api/v1/users/login', '用户登录'],
        ['用户', 'POST', '/api/v1/users/register', '用户注册'],
        ['用户', 'GET', '/api/v1/users/current', '获取当前用户信息'],
        ['用户', 'GET', '/api/v1/users/page/condition', '条件分页查询用户'],
        ['雪具', 'GET', '/equipment/list', '查询所有雪具'],
        ['雪具', 'POST', '/equipment', '新增雪具'],
        ['雪具', 'PUT', '/equipment', '更新雪具'],
        ['雪具', 'DELETE', '/equipment/{id}', '删除雪具'],
        ['教练', 'GET', '/coach/list', '查询所有教练'],
        ['教练', 'GET', '/coach/page/condition', '条件分页查询教练'],
        ['课程', 'GET', '/api/v1/courses/list', '查询所有课程'],
        ['课程', 'GET', '/api/v1/courses/list/hot', '查询热门课程'],
        ['场地', 'GET', '/api/v1/venues/list', '查询所有场地'],
        ['场地', 'GET', '/api/v1/venues/list/available', '查询可用场地'],
        ['租赁', 'POST', '/rental-order', '创建租赁订单'],
        ['租赁', 'PUT', '/rental-order/{id}/return', '归还设备'],
        ['销售', 'POST', '/api/v1/sales-orders', '创建销售订单'],
        ['销售', 'PUT', '/api/v1/sales-orders/{id}/ship', '订单发货'],
        ['统计', 'GET', '/api/v1/dashboard/overview', '获取统计概览'],
        ['统计', 'GET', '/api/v1/dashboard/revenue', '获取收入统计'],
    ]
)

# ========== 第五章 系统实现 ==========
doc.add_page_break()
add_title('第五章 系统实现', level=1)

add_title('5.1 开发环境', level=2)
add_body('系统的开发环境配置如下表所示：')
add_table_title('表5-1 开发环境配置')
create_table(
    ['名称', '版本/说明'],
    [
        ['操作系统', 'Windows 10/11'],
        ['开发工具（后端）', 'IntelliJ IDEA'],
        ['开发工具（前端）', 'Visual Studio Code'],
        ['JDK', '17'],
        ['Maven', '3.6+'],
        ['Node.js', '16.0+'],
        ['MySQL', '8.0'],
        ['Redis', '5.0+'],
    ]
)

add_title('5.2 用户登录与注册的实现', level=2)
add_body('用户登录是系统的入口功能。前端登录页面采用了左右分栏的布局设计，左侧展示系统名称和特色功能介绍，右侧是登录和注册的表单区域，支持在登录和注册两个Tab页之间切换。')
add_body('登录的具体实现流程如下：用户在前端输入用户名和密码后点击"立即登录"按钮，前端对表单进行校验（用户名和密码不能为空），校验通过后调用用户状态管理Store中的login方法，该方法内部通过Axios向后端的/api/v1/users/login接口发送POST请求。后端UserController接收到登录请求后，调用UserService的login方法进行用户名和密码的验证。密码验证使用了Spring Security提供的BCryptPasswordEncoder，它可以将用户输入的明文密码与数据库中存储的加密密码进行比对。验证通过后，调用JwtUtil.generateToken方法生成一个包含用户ID、用户名和角色信息的JWT Token，连同用户基本信息一起返回给前端。前端收到返回数据后，将Token和用户信息保存到Pinia Store和LocalStorage中，然后根据用户角色进行页面跳转：如果是普通用户（role为user）则跳转到前台首页"/home"，如果是管理员（role为admin）或教练（role为coach）则跳转到后台首页"/admin/dashboard"。')
add_body('注册功能的实现相对简单一些，用户填写用户名、密码、确认密码和手机号，前端会校验用户名长度（3-20个字符）、密码长度（至少6个字符）、两次密码是否一致以及手机号格式是否正确，校验通过后向后端发送注册请求。后端会检查用户名是否已存在，如果不存在则对密码进行BCrypt加密后保存到数据库中，默认角色为"user"。')
add_figure_title('图5-1 用户登录流程图（见附图）')

add_title('5.3 雪具租赁模块的实现', level=2)
add_body('雪具租赁模块分为管理端和用户端两个部分。')
add_body('在管理端，管理员可以通过后台的"雪具管理"页面对所有雪具进行管理。页面顶部提供了按雪具编号、名称、类型和状态进行搜索筛选的功能，下方用Element Plus的Table组件以表格形式展示雪具列表，支持分页显示。每行数据后面有"编辑"和"删除"两个操作按钮。点击"新增雪具"按钮会弹出一个对话框表单，管理员填写雪具的各项信息后提交即可添加新的雪具。后端EquipmentController提供了完整的CRUD接口：GET请求查询、POST请求新增、PUT请求更新、DELETE请求删除。查询接口支持不带条件的全量查询（/equipment/list）和带条件的分页查询（/equipment/page/condition），分页查询可以按编号、名称、类型和状态进行筛选。')
add_body('在用户端，普通用户可以通过前台的"雪具租赁"页面浏览所有可租赁的雪具。页面以卡片形式展示每件雪具的图片、名称、品牌型号、尺寸、租赁价格等信息，用户选择需要的雪具后可以指定租赁时间和数量，提交租赁订单。订单创建时，系统会自动扣减雪具的可用数量，归还时再加回来。')
add_body('租赁订单的状态流转为：待支付→使用中→已归还（或已取消）。后端RentalOrderController提供了状态更新的接口，包括支付（/rental-order/{id}/pay）、归还（/rental-order/{id}/return）和取消（/rental-order/{id}/cancel）操作。')

add_title('5.4 教练预约模块的实现', level=2)
add_body('教练预约模块涉及到管理员端、教练端和用户端三个角色。')
add_body('管理员端可以管理所有教练的基本信息。教练列表页面支持按编号、姓名、等级和状态进行筛选查询，管理员可以新增教练、编辑教练信息和修改教练状态。教练数据中的user_id字段关联了用户表，这样教练就可以通过自己的用户账号登录系统。')
add_body('用户端的教练预约页面展示了所有在职教练的信息，包括姓名、等级、专长、课时费、评分等。用户选择教练后，需要指定预约日期、开始时间和结束时间，系统会根据教练的课时费和预约时长自动计算费用。提交预约后，预约记录的初始状态为"待确认"。')
add_body('教练端登录后可以在"学员预约"页面查看所有针对自己的预约记录，教练可以对预约进行确认或取消操作。后端CoachBookingController处理预约相关的所有请求，通过coach_id进行数据隔离，确保每个教练只能看到自己的预约数据。')

add_title('5.5 课程管理模块的实现', level=2)
add_body('课程管理模块包含课程的发布管理和学员的报名功能。')
add_body('管理员在后台的"课程管理"页面可以发布新的课程。添加课程时需要填写课程名称、课程类型、授课教练、开课日期、上课时间、上课场地、最大学员数和课程价格等信息。课程列表页面支持条件筛选和分页查询，管理员还可以修改课程状态（未开始/报名中/进行中/已结束）。')
add_body('教练端可以在"我的课程"页面查看自己负责教授的课程列表。系统通过登录用户的ID查找对应的教练记录，再根据教练ID查询该教练名下的所有课程。')
add_body('用户端的"滑雪课程"页面展示了所有状态为"报名中"的课程信息。页面采用卡片式布局，每张卡片显示课程名称、教练姓名、上课时间、场地、已报名人数和课程价格等。用户点击报名按钮后，系统会检查课程的当前报名人数是否已达上限，未满则允许报名，同时将课程的current_students字段加1。课程报名记录存储在t_course_enrollment表中。')

add_title('5.6 场地管理模块的实现', level=2)
add_body('场地管理比较简单直接。管理员可以在后台录入和编辑场地信息，包括场地名称、类型（雪道或练习场）、难度等级、尺寸、容纳人数和租赁价格等。系统中初始化了6个场地数据：初级雪道A区、初级雪道B区、中级雪道C区、中级雪道D区、高级雪道E区和自由式公园。其中雪道类型的场地包含在门票内无需额外付费，自由式公园则需要额外收费（80元/小时）。')
add_body('用户端可以查看各场地的状态信息，对于需要付费的场地可以进行在线预约。场地预订信息存储在t_venue_booking表中，记录了预订用户、场地、预订日期、时间段、人数和费用等信息。')

add_title('5.7 装备销售模块的实现', level=2)
add_body('装备销售模块在管理端提供了商品信息的管理功能，在用户端则以"装备商城"的形式呈现。')
add_body('管理员可以在"商品管理"页面录入商品信息，包括名称、分类（滑雪装备/服装/配件）、品牌、价格、库存数量等，还可以标记商品是否热门或新品。系统初始化了15件商品数据，涵盖了单板、双板、固定器、滑雪服、滑雪裤、雪镜、头盔、雪杖、滑雪包等多种品类。')
add_body('用户端的装备商城页面以网格卡片形式展示商品，每张卡片包含商品图片、名称、品牌、价格等信息。用户可以按分类筛选商品，选择商品后添加到购物车，然后提交订单进行购买。销售订单支持待支付→待发货→已发货→已完成的完整流程，管理员在后台可以对已支付的订单进行发货操作，填写快递公司和快递单号。后端SalesOrderController提供了创建订单、支付、发货、确认收货等接口。')

add_title('5.8 数据统计模块的实现', level=2)
add_body('数据统计模块为管理员提供了一个Dashboard数据看板页面和一个数据可视化分析页面。')
add_body('Dashboard看板页面展示了系统的核心运营指标，包括总用户数、总订单数、总收入、当日新增用户数等概览数据。页面使用Element Plus的统计卡片组件来呈现这些数字指标。')
add_body('数据可视化页面则通过ECharts图表展示更加详细的数据分析结果，包括：收入趋势折线图（展示近期每日/每月的收入变化趋势）、订单分布饼图（展示各类订单的占比情况）、热门雪具排行柱状图（展示租赁次数最多的雪具）、热门教练排行（按评分和学员数排序）以及热门课程排行等。')
add_body('后端DashboardController提供了六个统计接口，分别对应概览数据（/overview）、收入统计（/revenue）、订单统计（/orders）、热门雪具（/popular-equipment）、热门教练（/popular-coaches）和热门课程（/popular-courses）。这些接口会查询数据库中的业务数据进行聚合计算，部分查询结果会通过Redis缓存来提升响应速度，避免每次请求都执行复杂的数据库查询。')
add_figure_title('图5-2 数据统计模块功能结构图（见附图）')

add_title('5.9 系统安全实现', level=2)
add_body('系统的安全控制主要体现在以下几个方面：')
add_body('密码加密存储：用户密码在注册时通过BCryptPasswordEncoder进行加密后才保存到数据库中，BCrypt是一种单向加密算法，即使数据库被泄露也无法直接获取用户的明文密码。')
add_body('JWT Token认证：系统配置了JwtAuthenticationFilter过滤器，对每个HTTP请求进行Token的解析和验证。前端在每次请求时会在Authorization请求头中携带Bearer Token，后端过滤器从中提取Token并调用JwtUtil进行验证，验证通过后将用户信息设置到Security上下文中。')
add_body('前端路由守卫：前端在路由配置中使用了beforeEach全局守卫，检查目标路由是否需要登录权限以及当前用户的角色是否有权限访问该路由。如果用户未登录（Token不存在），会被重定向到登录页面；如果用户已登录但角色不匹配，会被重定向到自己角色对应的首页。')
add_body('跨域配置：由于前后端分别部署在不同的端口（前端5173，后端8080），需要处理跨域请求的问题。后端通过CorsConfig配置类和@CrossOrigin注解来允许跨域访问。')

# ========== 第六章 系统测试 ==========
doc.add_page_break()
add_title('第六章 系统测试', level=1)

add_title('6.1 测试环境', level=2)
add_body('系统测试在以下环境中进行：')
add_table_title('表6-1 测试环境')
create_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['处理器', 'Intel Core i5'],
        ['内存', '16GB'],
        ['浏览器', 'Chrome 120'],
        ['后端运行环境', 'JDK 17 + Tomcat(内嵌)'],
        ['数据库', 'MySQL 8.0'],
        ['缓存', 'Redis 7.0'],
    ]
)
add_body('测试使用的账号信息如下：管理员账号admin（密码123456），教练账号coach1-coach5（密码123456），普通用户账号testuser（密码123456）。')

add_title('6.2 功能测试', level=2)

add_title('6.2.1 用户模块测试', level=3)
add_table_title('表6-2 用户模块测试用例')
create_table_left(
    ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过'],
    [
        ['T-U01', '用户登录（正确）', '输入admin/123456，点击登录', '登录成功，跳转到后台首页', '登录成功，跳转到/admin/dashboard', '通过'],
        ['T-U02', '用户登录（错误密码）', '输入admin/wrongpwd，点击登录', '提示"用户名或密码错误"', '提示"用户名或密码错误"', '通过'],
        ['T-U03', '用户登录（空用户名）', '不输入用户名，点击登录', '提示"请输入用户名"', '表单校验拦截，提示信息正确', '通过'],
        ['T-U04', '普通用户登录', '输入testuser/123456', '登录成功，跳转到前台首页', '跳转到/home', '通过'],
        ['T-U05', '用户注册', '填写合法的注册信息提交', '注册成功，切换到登录Tab', '提示"注册成功，请登录"', '通过'],
        ['T-U06', '注册（重复用户名）', '使用已存在的用户名注册', '提示用户名已存在', '提示注册失败', '通过'],
        ['T-U07', '修改个人信息', '在个人中心修改姓名后保存', '保存成功，信息更新', '更新成功', '通过'],
        ['T-U08', '修改密码', '输入正确的旧密码和新密码', '密码修改成功', '修改成功', '通过'],
    ]
)

add_title('6.2.2 雪具租赁模块测试', level=3)
add_table_title('表6-3 雪具租赁模块测试用例')
create_table_left(
    ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过'],
    [
        ['T-E01', '查看雪具列表', '管理员进入雪具管理页面', '显示所有雪具的分页列表', '正确显示10条雪具数据', '通过'],
        ['T-E02', '新增雪具', '点击新增，填写信息后提交', '新增成功，列表刷新', '新增成功', '通过'],
        ['T-E03', '编辑雪具', '修改雪具价格后保存', '更新成功', '价格更新正确', '通过'],
        ['T-E04', '删除雪具', '点击删除按钮并确认', '逻辑删除成功', '从列表中消失', '通过'],
        ['T-E05', '条件搜索', '按类型"单板"筛选', '只显示单板类型的雪具', '筛选结果正确', '通过'],
        ['T-E06', '用户租赁雪具', '选择雪具，填写租赁信息提交', '创建租赁订单成功', '订单创建成功，库存减少', '通过'],
        ['T-E07', '归还雪具', '管理员操作归还', '订单状态变更，库存恢复', '状态和库存更新正确', '通过'],
    ]
)

add_title('6.2.3 教练预约模块测试', level=3)
add_table_title('表6-4 教练预约模块测试用例')
create_table_left(
    ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过'],
    [
        ['T-C01', '查看教练列表', '用户进入教练预约页面', '显示所有在职教练', '正确显示5名教练', '通过'],
        ['T-C02', '提交预约', '选择教练和时间段后提交', '预约创建成功', '创建成功，显示预约号', '通过'],
        ['T-C03', '教练确认预约', '教练登录后确认预约', '预约状态变为"已确认"', '状态更新正确', '通过'],
        ['T-C04', '取消预约', '用户取消一条待确认的预约', '预约状态变为"已取消"', '取消成功', '通过'],
    ]
)

add_title('6.2.4 课程与场地模块测试', level=3)
add_table_title('表6-5 课程与场地模块测试用例')
create_table_left(
    ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过'],
    [
        ['T-CR01', '查看课程列表', '用户进入滑雪课程页面', '显示可报名课程', '显示8门课程', '通过'],
        ['T-CR02', '报名课程', '用户选择课程后报名', '报名成功，当前人数+1', '报名成功', '通过'],
        ['T-CR03', '发布课程', '管理员新增课程信息', '课程发布成功', '发布成功', '通过'],
        ['T-V01', '查看场地列表', '进入场地页面', '显示所有场地信息', '正确显示6个场地', '通过'],
        ['T-V02', '预约场地', '选择自由式公园并预约', '预约创建成功', '创建成功', '通过'],
    ]
)

add_title('6.2.5 销售与统计模块测试', level=3)
add_table_title('表6-6 销售与统计模块测试用例')
create_table_left(
    ['编号', '测试项', '测试步骤', '预期结果', '实际结果', '是否通过'],
    [
        ['T-S01', '浏览商品', '用户进入装备商城', '显示在售商品列表', '正确显示15件商品', '通过'],
        ['T-S02', '购买商品', '选择商品后提交订单', '订单创建成功', '创建成功', '通过'],
        ['T-S03', '订单发货', '管理员填写快递信息发货', '订单状态变为"已发货"', '状态更新正确', '通过'],
        ['T-D01', '查看统计概览', '管理员进入Dashboard', '显示运营数据概览', '各项数据正确显示', '通过'],
        ['T-D02', '查看可视化图表', '进入数据可视化页面', 'ECharts图表正常渲染', '图表展示正确', '通过'],
    ]
)

add_title('6.3 性能测试', level=2)
add_body('对系统的主要接口进行了简单的性能测试，在单用户场景下测量了各接口的响应时间，结果如下表：')
add_table_title('表6-7 接口响应时间测试')
create_table(
    ['接口', '说明', '平均响应时间', '是否达标（<2s）'],
    [
        ['POST /api/v1/users/login', '用户登录', '156ms', '是'],
        ['GET /api/v1/users/page/condition', '用户列表', '89ms', '是'],
        ['GET /equipment/page/condition', '雪具列表', '72ms', '是'],
        ['GET /coach/page/condition', '教练列表', '65ms', '是'],
        ['GET /api/v1/courses/list', '课程列表', '48ms', '是'],
        ['GET /api/v1/venues/list', '场地列表', '35ms', '是'],
        ['GET /api/v1/dashboard/overview', '统计概览', '210ms', '是'],
        ['POST /rental-order', '创建租赁订单', '125ms', '是'],
    ]
)
add_body('从测试结果来看，所有接口的响应时间均在200ms左右，远低于2秒的标准，系统性能表现良好。统计概览接口由于需要进行多表聚合查询，响应时间相对较长，但通过Redis缓存可以进一步优化。')

add_title('6.4 测试总结', level=2)
add_body('经过上述功能测试和性能测试，系统的各项功能模块均能正常运行，主要的业务流程（用户登录、雪具租赁、教练预约、课程报名、商品购买、数据统计等）都通过了测试验证。系统的接口响应速度在可接受的范围内。在测试过程中发现的一些小问题（如个别页面的字段显示格式不统一等）已经在开发阶段进行了修复。总体来说，系统达到了设计预期的功能目标。')

# ========== 第七章 总结与展望 ==========
doc.add_page_break()
add_title('第七章 总结与展望', level=1)

add_title('7.1 总结', level=2)
add_body('本文完成了基于Spring Boot的飞跃滑雪场管理系统的设计与实现工作。系统采用Spring Boot 2.7 + Vue 3的前后端分离架构，后端使用MyBatis操作MySQL数据库，引入Redis进行缓存优化，通过JWT实现身份认证；前端基于Element Plus搭建管理界面，使用ECharts实现数据可视化。')
add_body('系统实现了用户管理、雪具租赁管理、教练预约管理、课程管理、场地管理、装备销售管理和数据统计分析等七大功能模块，设计了16张数据库表，开发了17个后端Controller和33个前端页面组件。经过功能测试和性能测试的验证，系统能够满足中小型滑雪场日常运营管理的基本需求。')
add_body('在整个开发过程中，我加深了对Spring Boot框架和Vue.js前端技术的理解，也学到了前后端分离项目的开发流程和协作方式。同时也认识到，一个实际可用的管理系统在功能完善性、用户体验、安全防护等方面还有很多需要持续优化的地方。')

add_title('7.2 展望', level=2)
add_body('虽然系统已经实现了主要的管理功能，但由于时间和个人能力的限制，还有一些方面可以在后续进行改进和完善：')
add_body('（1）移动端适配：目前系统虽然使用了响应式设计，但在手机端的使用体验还不够理想，后续可以考虑开发专门的小程序或移动端App。')
add_body('（2）支付功能集成：当前系统的支付功能仅做了状态模拟，没有对接真实的支付接口，后续可以集成微信支付或支付宝等第三方支付平台。')
add_body('（3）消息通知功能：可以增加短信通知或站内信功能，在预约确认、订单状态变更等时机自动通知用户。')
add_body('（4）权限细化：目前的权限控制比较粗粒度，只做到了角色级别的控制，后续可以引入更细粒度的权限管理，比如按菜单、按按钮级别进行控制。')
add_body('（5）数据分析增强：数据统计模块目前主要是基础的汇总和排行，后续可以增加更多维度的分析功能，比如用户画像分析、季节性趋势预测等。')

# ========== 参考文献 ==========
doc.add_page_break()
add_title('参考文献', level=1)

refs = [
    '[1] 王珊, 萨师煊. 数据库系统概论(第5版)[M]. 北京: 高等教育出版社, 2014.',
    '[2] 谢希仁. 计算机网络(第7版)[M]. 北京: 电子工业出版社, 2017.',
    '[3] 黄健宏. Redis设计与实现[M]. 北京: 机械工业出版社, 2014.',
    '[4] Craig Walls. Spring实战(第5版)[M]. 张卫滨, 译. 北京: 人民邮电出版社, 2020.',
    '[5] 霍春阳. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022.',
    '[6] 杨开振. 深入浅出Spring Boot 2.x[M]. 北京: 人民邮电出版社, 2019.',
    '[7] 国务院办公厅. 关于促进全民健身和体育消费推动体育产业高质量发展的意见[Z]. 国办发〔2019〕43号, 2019.',
    '[8] 国家体育总局. 冰雪运动发展规划(2016-2025年)[Z]. 2016.',
    '[9] Spring Boot Reference Documentation[EB/OL]. https://spring.io/projects/spring-boot, 2024.',
    '[10] Vue.js Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
    '[11] MyBatis Official Documentation[EB/OL]. https://mybatis.org/mybatis-3/, 2024.',
    '[12] Element Plus Official Documentation[EB/OL]. https://element-plus.org/, 2024.',
    '[13] Apache ECharts Documentation[EB/OL]. https://echarts.apache.org/, 2024.',
    '[14] MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/, 2024.',
    '[15] JSON Web Token Introduction[EB/OL]. https://jwt.io/introduction, 2024.',
]
for ref in refs:
    add_body_no_indent(ref)

# ========== 致谢 ==========
doc.add_page_break()
add_title('致  谢', level=1)

add_body('大学四年的时光转瞬即逝，在这个毕业论文完成之际，有很多需要感谢的人。')
add_body('首先要感谢我的指导老师，在毕业设计的整个过程中给予了我耐心的指导和帮助，从选题方向的确定到论文的修改完善，老师的建议让我受益匪浅。')
add_body('感谢大学四年来所有教过我的老师们，是你们的教导让我掌握了计算机专业的基础知识和编程技能，为完成这次毕业设计打下了基础。')
add_body('感谢我的同学和朋友们，在开发过程中遇到问题时，是大家的讨论和帮助让我度过了一个又一个难关。')
add_body('最后要感谢我的家人，是你们在背后的支持和鼓励，让我能够安心地完成学业。')
add_body('谢谢大家！')

# ============ 保存文件 ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '论文_基于SpringBoot的飞跃滑雪场管理系统的设计与实现.docx')
doc.save(output_path)
print(f'论文docx文件已生成: {output_path}')
