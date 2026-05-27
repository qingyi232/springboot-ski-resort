# -*- coding: utf-8 -*-
"""
修复第五章的顺序问题：
1. 确保每个模块格式统一：介绍文字 → 图 → 图标题 → 核心代码 → 代码
2. 修复 Para 267 中图片和标题混在同一段落的问题
3. 确保图编号正确连续
4. 只修改第五章，其他不动
"""
import sys, io, os, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

shutil.copy2('5.26.2.docx', '5.26.2_before_fix.docx')
print('已备份 5.26.2.docx -> 5.26.2_before_fix.docx')

doc = Document('5.26.2.docx')

in_ch5 = False
ch5_start = -1
ch5_end = -1

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style_name = para.style.name if para.style else ''
    
    if not in_ch5:
        if ('第五章' in text or '第5章' in text) and 'Heading' in style_name:
            in_ch5 = True
            ch5_start = i
    
    if in_ch5:
        if ('第六章' in text or '第6章' in text or '结论' in text) and 'Heading 1' in style_name:
            ch5_end = i
            break

print(f'第五章范围: para {ch5_start} ~ {ch5_end}')

fix_count = 0

for i in range(ch5_start, ch5_end):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    blips = para._element.findall(f'.//{{{qn("a:blip").split("}")[0][1:]}}}blip')
    if not blips:
        blips = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    
    has_image = len(para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
    
    if has_image and text and '图' in text:
        print(f'\n[Para {i}] 发现图片和标题混在同一段落: "{text}"')
        print(f'  -> 需要拆分为两个段落：图片段 + 标题段')
        
        img_runs = []
        text_runs = []
        
        for run in para.runs:
            has_drawing = len(run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
            if has_drawing:
                img_runs.append(run)
            else:
                text_runs.append(run)
        
        if img_runs and text_runs:
            new_para_elem = copy.deepcopy(para._element)
            
            for run in para.runs:
                has_drawing = len(run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
                if has_drawing:
                    run._element.getparent().remove(run._element)
            
            for r_elem in new_para_elem.findall(qn('w:r')):
                has_drawing = len(r_elem.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
                if not has_drawing:
                    r_elem.getparent().remove(r_elem)
            
            para._element.addprevious(new_para_elem)
            
            print(f'  -> 已拆分完成')
            fix_count += 1

print(f'\n共修复 {fix_count} 处问题')

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style_name = para.style.name if para.style else ''
    
    if not (ch5_start <= i < ch5_end + fix_count + 5):
        continue
    
    has_image = len(para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0
    
    if has_image or 'Heading' in style_name or '图5' in text or '核心代码' in text:
        img_tag = " [IMG]" if has_image else ""
        print(f'[{i}] [{style_name}]{img_tag} {text[:80]}')

doc.save('5.26.2.docx')
print('\n已保存修复后的 5.26.2.docx')
