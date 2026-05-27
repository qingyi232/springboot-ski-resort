# -*- coding: utf-8 -*-
"""
整段交换文档章节 + 修复6.4口语化
"""
import sys, io, copy, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

INPUT = '5.26.2.docx'
OUTPUT = '5.26.2.docx'

shutil.copy2(INPUT, '5.26.2_before_swap.docx')
print("已备份 -> 5.26.2_before_swap.docx")

doc = Document(INPUT)
body = doc.element.body

def get_elem_text(elem):
    return ''.join(r.text or '' for r in elem.findall(qn('w:r')))

def get_style_val(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'))
    return None

def find_heading_elem(body, text_contains, style_val):
    """按XML样式值和文字查找标题"""
    for elem in body:
        if elem.tag == qn('w:p'):
            sv = get_style_val(elem)
            if sv == style_val:
                text = get_elem_text(elem)
                if text_contains in text:
                    return elem
    return None

def set_all_runs_text(elem, new_text):
    """设置段落文字，合并到第一个run"""
    runs = elem.findall(qn('w:r'))
    if not runs:
        return
    first_set = False
    for r in runs:
        t = r.find(qn('w:t'))
        if t is not None:
            if not first_set:
                t.text = new_text
                t.set(qn('xml:space'), 'preserve')
                first_set = True
            else:
                t.text = ''

def collect_between(body, start_elem, end_elem):
    """收集两个元素之间的所有body子元素（不含start和end）"""
    elements = []
    collecting = False
    for elem in list(body):
        if elem is start_elem:
            collecting = True
            continue
        if elem is end_elem:
            break
        if collecting:
            elements.append(elem)
    return elements

def swap_two_sections(body, h1, h2, h_end, new_h1_text, new_h2_text):
    """交换两个连续章节的完整内容"""
    content_a = collect_between(body, h1, h2)
    content_b = collect_between(body, h2, h_end)
    
    print(f"  章节A内容: {len(content_a)} 个元素")
    print(f"  章节B内容: {len(content_b)} 个元素")
    
    clones_a = [copy.deepcopy(e) for e in content_a]
    clones_b = [copy.deepcopy(e) for e in content_b]
    
    for e in content_a:
        body.remove(e)
    for e in content_b:
        body.remove(e)
    
    cursor = h1
    for e in clones_b:
        cursor.addnext(e)
        cursor = e
    
    cursor = h2
    for e in clones_a:
        cursor.addnext(e)
        cursor = e
    
    set_all_runs_text(h1, new_h1_text)
    set_all_runs_text(h2, new_h2_text)

# ============================================================
# 1. 交换第5章 5.3(教练) ↔ 5.4(工作人员)
# ============================================================
print("\n" + "=" * 60)
print("1. 交换第5章 5.3(教练) ↔ 5.4(工作人员)")
print("=" * 60)

h53 = find_heading_elem(body, '5.3', '3')
h54 = find_heading_elem(body, '5.4', '3')
h55 = find_heading_elem(body, '5.5', '3')

if h53 and h54 and h55:
    t53 = get_elem_text(h53)
    t54 = get_elem_text(h54)
    print(f"  5.3: {t53}")
    print(f"  5.4: {t54}")
    print(f"  5.5边界: {get_elem_text(h55)}")
    
    new_53 = t53.replace('教练', '工作人员')
    new_54 = t54.replace('工作人员', '教练')
    
    swap_two_sections(body, h53, h54, h55, new_53, new_54)
    print(f"  新5.3: {new_53}")
    print(f"  新5.4: {new_54}")
    print("  交换完成!")
else:
    print(f"  [ERROR] h53={h53 is not None}, h54={h54 is not None}, h55={h55 is not None}")

# ============================================================
# 2. 交换第4章 4.4(业务流程) ↔ 4.5(数据库)
# ============================================================
print("\n" + "=" * 60)
print("2. 交换第4章 4.4(业务流程) ↔ 4.5(数据库)")
print("=" * 60)

h44 = find_heading_elem(body, '4.4', '3')
h45 = find_heading_elem(body, '4.5', '3')
h46 = find_heading_elem(body, '4.6', '3')

if h44 and h45 and h46:
    t44 = get_elem_text(h44)
    t45 = get_elem_text(h45)
    print(f"  4.4: {t44}")
    print(f"  4.5: {t45}")
    print(f"  4.6边界: {get_elem_text(h46)}")
    
    swap_two_sections(body, h44, h45, h46,
                      t44.replace('业务流程设计', '数据库设计'),
                      t45.replace('数据库设计', '业务流程设计'))
    
    # 更新子标题编号 4.5.1→4.4.1, 4.5.2→4.4.2
    for elem in body:
        if elem.tag == qn('w:p') and get_style_val(elem) == '4':
            text = get_elem_text(elem)
            if '4.5.1' in text:
                set_all_runs_text(elem, text.replace('4.5.1', '4.4.1'))
                print(f"  子标题: {text.strip()} -> {text.replace('4.5.1', '4.4.1').strip()}")
            elif '4.5.2' in text:
                set_all_runs_text(elem, text.replace('4.5.2', '4.4.2'))
                print(f"  子标题: {text.strip()} -> {text.replace('4.5.2', '4.4.2').strip()}")
    
    print("  交换完成!")
else:
    print(f"  [ERROR] h44={h44 is not None}, h45={h45 is not None}, h46={h46 is not None}")

# ============================================================
# 3. 更新第5章图编号
# 交换后: 新5.3=工作人员(有图5.5工作人员后台, 图5.6工作人员模块)
#         新5.4=教练(有图5.4教练预约)
# 应更新为: 图5.4工作人员后台, 图5.5工作人员模块, 图5.6教练预约
# ============================================================
print("\n" + "=" * 60)
print("3. 更新第5章图编号")
print("=" * 60)

for elem in body:
    if elem.tag == qn('w:p'):
        text = get_elem_text(elem)
        if '图5.5' in text and '工作人员后台' in text:
            new = text.replace('图5.5', '图5.4')
            set_all_runs_text(elem, new)
            print(f"  {text.strip()} -> {new.strip()}")
        elif '图5.6' in text and '工作人员' in text:
            new = text.replace('图5.6', '图5.5')
            set_all_runs_text(elem, new)
            print(f"  {text.strip()} -> {new.strip()}")
        elif '图5.4' in text and '教练预约' in text:
            new = text.replace('图5.4', '图5.6')
            set_all_runs_text(elem, new)
            print(f"  {text.strip()} -> {new.strip()}")

# ============================================================
# 4. 修复6.4口语化
# ============================================================
print("\n" + "=" * 60)
print("4. 修复6.4口语化")
print("=" * 60)

oral_map = [
    ('一个是商品管理页面的状态字段',
     '其一，商品管理页面的状态字段中，后端返回的是数字编码（0表示下架、1表示上架），前端未做转换而直接显示原始数字，导致用户无法理解其含义，后通过在前端增加状态映射予以修复。其二，租赁订单金额存在浮点精度问题，偶尔出现多位小数的情况，后在前端使用toFixed(2)进行格式化处理。其三，教练预约的时间段校验不够严格，未限制用户选择已过去的时间，后在前端为日期选择器添加disabledDate属性加以约束。以上均属非核心缺陷，发现后已及时修复。'),
    ('此外，在并发控制方面也存在待优化之处',
     '此外，在并发控制方面仍存在待优化之处。当多个用户同时租赁同一件商品时，存在库存出现负数的风险，其根本原因在于扣减库存时未引入并发控制机制。应在数据库层面采用乐观锁等手段加以处理。鉴于本系统定位为中小型雪场使用，并发量相对有限，该问题在当前阶段影响较小，但随着用户规模扩大，该部分仍有优化的必要。'),
    ('总的来说，系统基本达到了',
     '综上所述，系统基本达到了预期的功能目标，各项核心业务流程均可正常运行。'),
]

fix_count = 0
for elem in body:
    if elem.tag == qn('w:p'):
        text = get_elem_text(elem)
        for match_prefix, replacement in oral_map:
            if match_prefix in text:
                set_all_runs_text(elem, replacement)
                fix_count += 1
                print(f"  已修复: {match_prefix[:30]}...")

print(f"  共修复 {fix_count} 处")

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT)
print(f"\n已保存到 {OUTPUT}")

# ============================================================
# 验证
# ============================================================
print("\n" + "=" * 60)
print("验证结果")
print("=" * 60)

doc2 = Document(OUTPUT)

print("\n第4章结构:")
in_ch4 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第4章' in text:
        in_ch4 = True
    elif in_ch4 and style == 'Heading 1':
        break
    if in_ch4 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n第5章结构:")
in_ch5 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n第5章各图:")
in_ch5 = False
rels = doc2.part.rels
for i, p in enumerate(doc2.paragraphs):
    style = p.style.name
    text = p.text.strip()
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and '图5' in text and not style.startswith('Heading'):
        print(f"  [{i}] {text}")
    if in_ch5:
        drawings = p._element.findall('.//' + qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed and embed in rels:
                    print(f"  [{i}] IMG -> {rels[embed].target_ref}")

print("\n第5章角色顺序:")
in_ch5 = False
order = []
for p in doc2.paragraphs:
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style == 'Heading 2':
        order.append(text)
        
print(f"  {' -> '.join(order)}")

print("\n6.4内容:")
in_64 = False
for p in doc2.paragraphs:
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
        print(f"  [{style}] {text}")
    elif in_64 and style.startswith('Heading'):
        break
    elif in_64 and text:
        print(f"  {text[:80]}...")
