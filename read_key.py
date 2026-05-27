# -*- coding: utf-8 -*-
from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx")

key_paras = {
    "6.4": list(range(406, 411)),
    "7.1": list(range(412, 416)),
    "4.4_end": list(range(175, 180)),
    "ch5_5.5": list(range(271, 303)),
}

with open(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\key_text.txt", "w", encoding="utf-8") as f:
    for section, indices in key_paras.items():
        f.write(f"\n=== {section} ===\n")
        for i in indices:
            if i < len(doc.paragraphs):
                p = doc.paragraphs[i]
                f.write(f"P{i} [{p.style.name}]: {p.text}\n---\n")

    f.write("\n=== References ===\n")
    for i in range(422, 437):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            f.write(f"P{i}: {p.text}\n")

    f.write("\n=== Check citations in text ===\n")
    for i in range(56, 422):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text
            import re
            refs = re.findall(r'\[(\d+)\]', text)
            if refs:
                f.write(f"P{i}: refs={refs} | {text[:80]}\n")

print("Done")
