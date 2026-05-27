# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document('5.26.2.docx')
rels = doc.part.rels

print("=== rId 映射 ===")
for rid, rel in sorted(rels.items()):
    if 'image' in rel.reltype:
        print(f"  {rid} -> {rel.target_ref}")

print("\n=== 第5章图片 rId 值 ===")
in_ch5 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if not in_ch5:
        continue
    
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            target = rels[embed].target_ref if embed in rels else '?'
            next_text = doc.paragraphs[i+1].text.strip()[:50] if i+1 < len(doc.paragraphs) else ''
            print(f"  [{i}] r:embed={embed} -> {target}  下文: {next_text}")

print("\n=== 原始文件(5.26.2_pre_fix.docx)的 rId ===")
try:
    doc_orig = Document('5.26.2_pre_fix.docx')
    rels_orig = doc_orig.part.rels
    in_ch5 = False
    for i, p in enumerate(doc_orig.paragraphs):
        text = p.text.strip()
        style = p.style.name
        if style == 'Heading 1' and '第5章' in text:
            in_ch5 = True
        elif in_ch5 and style == 'Heading 1':
            break
        if not in_ch5:
            continue
        drawings = p._element.findall('.//' + qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                target = rels_orig[embed].target_ref if embed in rels_orig else '?'
                next_text = doc_orig.paragraphs[i+1].text.strip()[:50] if i+1 < len(doc_orig.paragraphs) else ''
                print(f"  [{i}] r:embed={embed} -> {target}  下文: {next_text}")
except Exception as e:
    print(f"  Error: {e}")
