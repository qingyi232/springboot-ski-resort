# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
doc = Document('5.26.2.docx')

print("=" * 60)
print("详查 6.4 全部文字")
print("=" * 60)
in_64 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
        print(f"[{i}] [{style}] {text}")
    elif in_64 and style.startswith('Heading'):
        break
    elif in_64 and text:
        print(f"[{i}] {text}")
        print()

print("\n" + "=" * 60)
print("详查 7.1 全部文字")
print("=" * 60)
in_71 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '7.1' in text and style.startswith('Heading'):
        in_71 = True
        print(f"[{i}] [{style}] {text}")
    elif in_71 and style.startswith('Heading'):
        break
    elif in_71 and text:
        print(f"[{i}] {text}")
        print()

print("\n" + "=" * 60)
print("详查第4章 4.4和4.5 标题及 4.4最后三段")
print("=" * 60)
in_44 = False
paras_44 = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '4.4' in text:
        in_44 = True
        print(f"[{i}] [{style}] {text}")
    elif in_44 and style.startswith('Heading'):
        print(f"[{i}] [{style}] {text}")
        if style == 'Heading 2' and ('4.5' in text or '4.6' in text):
            break
    elif in_44 and text:
        paras_44.append((i, text))

if paras_44:
    print(f"\n  4.4 共有 {len(paras_44)} 个段落")
    print(f"  最后3段:")
    for idx, txt in paras_44[-3:]:
        print(f"  [{idx}] {txt[:120]}")
        print()

print("\n" + "=" * 60)
print("详查 3.2 活动图/用例图位置")
print("=" * 60)
from docx.oxml.ns import qn
in_32 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '3.2' in text and style == 'Heading 2':
        in_32 = True
    if in_32 and style == 'Heading 2' and '3.3' in text:
        break
    if in_32:
        drawings = p._element.findall('.//' + qn('w:drawing'))
        has_img = len(drawings) > 0
        if style.startswith('Heading'):
            print(f"  [{i}] [{style}] {text}")
        elif has_img:
            print(f"  [{i}] [IMG]")
        elif '图' in text and '3.' in text:
            print(f"  [{i}] [图题] {text}")
        elif '表' in text and '3.' in text:
            print(f"  [{i}] [表题] {text}")
