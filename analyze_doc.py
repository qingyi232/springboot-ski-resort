import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx')

print(f'段落总数: {len(doc.paragraphs)}')
print(f'表格总数: {len(doc.tables)}')
print()

# 查看所有表格信息
for i, t in enumerate(doc.tables):
    first_cell = t.cell(0, 0).text[:60]
    print(f'表格{i+1}: {len(t.rows)}行x{len(t.columns)}列, 首格: {first_cell}')

print('\n--- 文档内容结构（段落+表格位置）---')

# 遍历文档body，按顺序输出段落和表格
from docx.oxml.ns import qn
table_idx = 0
para_idx = 0
for element in doc.element.body:
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    if tag == 'p':
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
            style = p.style.name if p.style else 'None'
            text = p.text[:80]
            if text.strip():
                print(f'[P{para_idx}][{style}] {text}')
            para_idx += 1
    elif tag == 'tbl':
        if table_idx < len(doc.tables):
            t = doc.tables[table_idx]
            print(f'\n  >>> [TABLE {table_idx+1}] {len(t.rows)}行x{len(t.columns)}列, 首格: {t.cell(0,0).text[:60]}')
            # 输出表格完整内容
            for ri, row in enumerate(t.rows):
                cells = [cell.text[:30] for cell in row.cells]
                print(f'      行{ri}: {" | ".join(cells)}')
            print()
            table_idx += 1
