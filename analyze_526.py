# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
import json
import sys

doc = Document(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx")

output = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    text = para.text.strip()
    
    is_heading = style_name.startswith("Heading") or "标题" in style_name
    has_image = any(run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')) for run in para.runs)
    
    if is_heading or has_image or (text and len(text) < 100):
        prefix = ""
        if is_heading:
            prefix = f"[H-{style_name}] "
        if has_image:
            prefix += "[IMG] "
        output.append(f"P{i}: {prefix}{text[:120]}")
    elif text:
        output.append(f"P{i}: [{len(text)}chars] {text[:80]}...")

with open(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\doc_structure.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(output))

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Output lines: {len(output)}")

tables_info = []
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_row = [cell.text.strip()[:30] for cell in table.rows[0].cells] if rows > 0 else []
    tables_info.append(f"Table {i}: {rows}x{cols} | Header: {first_row}")

with open(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\doc_tables.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(tables_info))

print(f"Total tables: {len(doc.tables)}")
