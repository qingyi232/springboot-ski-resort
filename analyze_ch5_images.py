# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import zipfile
import os

doc = Document('5.26.2.docx')

# Get image relationships
print("=== 文档中图片关系映射 ===")
rels = doc.part.rels
for rid, rel in rels.items():
    if 'image' in rel.reltype:
        print(f"  {rid} -> {rel.target_ref}")

print("\n=== 第五章中每张图的 rId 和位置 ===\n")

ch5_start = False
img_idx = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    
    if '第5章' in text or '第五章' in text:
        ch5_start = True
    
    if ch5_start and style.startswith('Heading') and ('第6章' in text or '第六章' in text or '致谢' in text or '参考文献' in text):
        break
    
    if not ch5_start:
        continue
    
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            if embed and embed in rels:
                target = rels[embed].target_ref
                img_idx += 1
                # Get next paragraph text (figure caption)
                next_text = ""
                if i + 1 < len(doc.paragraphs):
                    next_text = doc.paragraphs[i + 1].text.strip()[:60]
                prev_text = ""
                if i > 0:
                    prev_text = doc.paragraphs[i - 1].text.strip()[:60]
                print(f"图{img_idx} [para {i}] rId={embed} -> {target}")
                print(f"  上文: {prev_text}")
                print(f"  下文(图题): {next_text}")
                print()

# Also check backup file
print("\n=== 备份文件(5.26.2_backup) 第五章图片 ===\n")
try:
    doc2 = Document('5.26.2_backup.docx')
    rels2 = doc2.part.rels
    ch5_start2 = False
    img_idx2 = 0
    
    for i, p in enumerate(doc2.paragraphs):
        text = p.text.strip()
        style = p.style.name
        
        if '第5章' in text or '第五章' in text:
            ch5_start2 = True
        
        if ch5_start2 and style.startswith('Heading') and ('第6章' in text or '第六章' in text or '致谢' in text or '参考文献' in text):
            break
        
        if not ch5_start2:
            continue
        
        drawings = p._element.findall('.//' + qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed and embed in rels2:
                    target = rels2[embed].target_ref
                    img_idx2 += 1
                    next_text = ""
                    if i + 1 < len(doc2.paragraphs):
                        next_text = doc2.paragraphs[i + 1].text.strip()[:60]
                    prev_text = ""
                    if i > 0:
                        prev_text = doc2.paragraphs[i - 1].text.strip()[:60]
                    print(f"图{img_idx2} [para {i}] rId={embed} -> {target}")
                    print(f"  上文: {prev_text}")
                    print(f"  下文(图题): {next_text}")
                    print()
except:
    print("备份文件不存在或无法打开")
