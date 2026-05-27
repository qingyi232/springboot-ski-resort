# -*- coding: utf-8 -*-
"""
修复 5.26.2.docx 第五章图片错位问题。
只交换第五章内图片引用的 rId，不修改其他章节。
"""
import shutil
from docx import Document
from docx.oxml.ns import qn

INPUT = '5.26.2.docx'
OUTPUT = '5.26.2.docx'
BACKUP = '5.26.2_pre_fix.docx'

shutil.copy2(INPUT, BACKUP)
print(f"已备份 -> {BACKUP}")

doc = Document(INPUT)
rels = doc.part.rels

rid_map = {}
for rid, rel in rels.items():
    if 'image' in rel.reltype:
        rid_map[rel.target_ref] = rid

rId_img11 = rid_map.get('media/image11.png')  # 用户登录
rId_img12 = rid_map.get('media/image12.png')  # 商品管理
rId_img13 = rid_map.get('media/image13.png')  # 教练预约
rId_img14 = rid_map.get('media/image14.png')  # 工作人员后台
rId_img15 = rid_map.get('media/image15.png')  # 工作人员模块
rId_img16 = rid_map.get('media/image16.png')  # 数据统计

print(f"rId_img11(用户登录) = {rId_img11}")
print(f"rId_img12(商品管理) = {rId_img12}")
print(f"rId_img13(教练预约) = {rId_img13}")
print(f"rId_img14(工作人员后台) = {rId_img14}")
print(f"rId_img15(工作人员模块) = {rId_img15}")
print(f"rId_img16(数据统计) = {rId_img16}")

ch5_start_para = -1
ch5_end_para = len(doc.paragraphs)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and ('第5章' in text or '第五章' in text):
        ch5_start_para = i
        print(f"\n第五章起始: para[{i}] {text}")
    elif ch5_start_para > 0 and style == 'Heading 1' and i > ch5_start_para:
        ch5_end_para = i
        print(f"第五章结束: para[{i}] {text}")
        break

ch5_images = []
for i in range(ch5_start_para, ch5_end_para):
    p = doc.paragraphs[i]
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            if embed and embed in rels:
                target = rels[embed].target_ref
                next_text = ""
                if i + 1 < len(doc.paragraphs):
                    next_text = doc.paragraphs[i + 1].text.strip()[:60]
                ch5_images.append({
                    'para_idx': i,
                    'blip': blip,
                    'current_rId': embed,
                    'current_target': target,
                    'caption': next_text,
                })

print(f"\n第五章找到 {len(ch5_images)} 张图片:")
for idx, img in enumerate(ch5_images):
    print(f"  图{idx+1}: para[{img['para_idx']}] {img['current_rId']} -> {img['current_target']}  图题: {img['caption']}")

correct_rids = [
    rId_img12,  # 图1: 5.1.1 商品管理 -> image12
    rId_img16,  # 图2: 5.1.2 数据统计 -> image16
    rId_img11,  # 图3: 5.2 用户登录 -> image11
    rId_img13,  # 图4: 5.3 教练预约 -> image13
    rId_img14,  # 图5: 5.4 工作人员后台 -> image14
    rId_img15,  # 图6: 5.4 工作人员模块 -> image15
]

assert len(ch5_images) == len(correct_rids), \
    f"期望 {len(correct_rids)} 张图片, 实际找到 {len(ch5_images)} 张"

print("\n=== 开始修复图片引用 ===")
changes = 0
for idx, img in enumerate(ch5_images):
    old_rid = img['current_rId']
    new_rid = correct_rids[idx]
    old_target = img['current_target']
    new_target = rels[new_rid].target_ref if new_rid in rels else '?'
    
    if old_rid != new_rid:
        img['blip'].set(qn('r:embed'), new_rid)
        changes += 1
        print(f"  图{idx+1} para[{img['para_idx']}]: {old_rid}({old_target}) -> {new_rid}({new_target})")
    else:
        print(f"  图{idx+1} para[{img['para_idx']}]: 无需修改 ({old_rid} -> {old_target})")

if changes > 0:
    doc.save(OUTPUT)
    print(f"\n修复完成! 共修改 {changes} 张图片引用, 已保存到 {OUTPUT}")
else:
    print("\n无需修改")

print("\n=== 验证修复结果 ===")
doc2 = Document(OUTPUT)
rels2 = doc2.part.rels
ch5_start2 = False
img_idx2 = 0

expected_content = [
    "image12 = 商品管理",
    "image16 = 数据统计", 
    "image11 = 用户登录",
    "image13 = 教练预约",
    "image14 = 工作人员后台",
    "image15 = 工作人员模块",
]

for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    
    if style == 'Heading 1' and ('第5章' in text or '第五章' in text):
        ch5_start2 = True
    elif ch5_start2 and style == 'Heading 1':
        break
    
    if not ch5_start2:
        continue
    
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            if embed and embed in rels2:
                target = rels2[embed].target_ref
                next_text = ""
                if i + 1 < len(doc2.paragraphs):
                    next_text = doc2.paragraphs[i + 1].text.strip()[:50]
                expected = expected_content[img_idx2] if img_idx2 < len(expected_content) else "?"
                status = "✓" if expected.split(" = ")[0].replace("image", "media/image") + ".png" == target else "?"
                print(f"  图{img_idx2+1}: {embed} -> {target}  图题: {next_text}  期望: {expected}")
                img_idx2 += 1
