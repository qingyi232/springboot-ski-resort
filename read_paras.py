# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx")

# Read paragraphs around 6.4 and 7.1
print("=== 6.4 测试总结 (P406-P410) ===")
for i in range(406, 411):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} [{p.style.name}]: {p.text}")
        print("---")

print("\n=== 7.1 总结 (P411-P416) ===")
for i in range(411, 423):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} [{p.style.name}]: {p.text}")
        print("---")

# Also check 3.2 and 3.4 structure
print("\n=== 3.2 功能需求 (P98-P109) ===")
for i in range(98, 110):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} [{p.style.name}]: {p.text[:100]}")
        print("---")

print("\n=== 3.4 用例分析 (P116-P135) ===")
for i in range(116, 135):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} [{p.style.name}]: {p.text[:100]}")
        print("---")

# Check chapter 5 structure
print("\n=== Chapter 5 headings ===")
for i in range(188, 382):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.style.name.startswith("Heading"):
            print(f"P{i} [{p.style.name}]: {p.text}")

# Check 4.4 and 4.5 structure
print("\n=== 4.4-4.5 area (P152-P184) ===")
for i in range(152, 184):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} [{p.style.name}]: {p.text[:120]}")
        print("---")
