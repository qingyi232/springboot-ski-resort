# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import re

doc = Document(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx')

# ============================================================
# STEP 1: 口语化 -> 书面语 替换
# ============================================================

phrase_replacements = [
    # 第1章
    ('就不太跟得上了', '已难以满足实际需求'),
    ('可以把各项业务流程规范起来，让管理者实时掌握运营情况，顾客也能获得更方便的服务体验。这就是本文选择做滑雪场管理系统的出发点。',
     '能够将各项业务流程规范化，使管理者实时掌握运营状况，同时为顾客提供更加便捷的服务体验。这是本文选择开发滑雪场管理系统的出发点。'),
    ('本文做的这个飞跃滑雪场管理系统，主要就是面向中小型雪场的日常运营。',
     '本文设计的飞跃滑雪场管理系统，主要面向中小型雪场的日常运营管理。'),
    ('各有各的操作界面和功能权限。',
     '分别拥有各自的操作界面和功能权限。'),
    ('开发这个系统也是把大学四年学的Java和前端知识综合用了一遍，希望能对雪场的信息化建设有一些参考。',
     '开发本系统也是对大学期间所学Java和前端知识的综合应用，以期为雪场的信息化建设提供参考。'),
    ('像瑞士、奥地利的一些老牌雪场', '如瑞士、奥地利的一些知名雪场'),
    ('价格比较贵，也不太适合', '价格较高，且不太适合'),
    ('都整合在了一起', '进行了整合'),
    # 第2章
    ('比较麻烦', '较为繁琐'),
    ('选这个版本主要是因为比较成熟稳定，网上的教程和文档也比较多，碰到问题好找答案。',
     '选择该版本的原因是其较为成熟稳定，且相关教程和文档较为丰富，便于解决开发中遇到的问题。'),
    ('在项目里，Spring Boot主要负责三件事', '在本项目中，Spring Boot主要承担三方面职责'),
    ('内嵌了Tomcat不用额外部署', '内嵌Tomcat无需额外部署'),
    ('不用写那些重复的代码', '避免编写重复的代码'),
    ('给了开发者更大的SQL控制权，可以自己写SQL语句', '赋予开发者更大的SQL控制权，允许自行编写SQL语句'),
    ('做条件分页查询的时候，可以根据', '进行条件分页查询时，可以根据'),
    ('性能不错，运行稳定，而且免费使用', '性能优良，运行稳定，且免费开源'),
    ('因为数据都存在内存里，读写速度比传统数据库快很多', '由于数据存储在内存中，其读写速度远高于传统数据库'),
    ('工作流程是这样的：', '其工作流程如下：'),
    ('前端收到Token后存在本地', '前端收到Token后将其存储在本地'),
    ('启动速度比传统的Webpack快很多', '启动速度显著优于传统的Webpack'),
    ('在国内Web开发中使用率很高', '在国内Web开发中应用广泛'),
    # 第3章
    ('在做需求调研的时候，主要参考了几个方面的资料：一是网上一些滑雪场的官方网站，看了它们提供的在线服务有哪些；二是查阅了一些关于体育场馆管理系统的论文，了解了这类系统通常需要具备什么功能；三是向身边去过滑雪场的同学了解了他们作为消费者的需求和体验。综合这些信息，确定了系统的主要功能模块和各角色的操作权限。',
     '在需求调研过程中，主要参考了以下几方面资料：一是浏览国内多家滑雪场的官方网站，了解其提供的在线服务内容；二是查阅体育场馆管理系统相关论文，了解此类系统通常应具备的功能；三是向有滑雪经历的同学了解其作为消费者的需求和体验。综合以上信息，确定了系统的主要功能模块和各角色的操作权限。'),
    ('避免一次性加载太多数据导致页面卡顿', '避免一次性加载过多数据导致页面响应缓慢'),
    # 第4章
    ('系统的请求处理流程大致是这样的：', '系统的请求处理流程如下：'),
    ('前端再将数据渲染到页面上', '前端再将数据渲染至页面'),
    ('选择前后端分离这种架构方式，主要有两个考虑。一个是开发效率，前端和后端可以各自独立开发，互不影响，只要约定好接口格式就行。另一个是部署灵活性，前端可以单独部署到Nginx上，后端也可以独立部署和扩展。',
     '选择前后端分离架构，主要基于两方面考虑。其一是开发效率，前端和后端可以各自独立开发，互不影响，只需约定好接口格式即可。其二是部署灵活性，前端可以单独部署至Nginx，后端也可以独立部署和横向扩展。'),
    ('遵循了几个原则', '遵循了以下设计原则'),
    ('这样设计的好处是', '这种设计方案的优势在于'),
    # 第5章
    ('添加课程的时候，前端会同时发两个请求去获取教练列表和场地列表作为下拉选项，管理员可以直接选择而不用手动输入ID',
     '添加课程时，前端会同时请求教练列表和场地列表作为下拉选项，管理员可以直接选择而无需手动输入ID'),
    # 第6章
    ('经过上面这些测试，系统的主要功能基本都跑通了，像用户登录、商品租赁、教练预约、课程报名、场地预约、数据统计这些核心流程都没有问题。接口的响应速度也还可以，都在200ms左右，满足性能要求。',
     '经过以上测试验证，系统的主要功能均已正常运行，包括用户登录、商品租赁、教练预约、课程报名、场地预约、数据统计等核心业务流程均通过测试。接口响应时间均在200ms以内，满足性能要求。'),
    ('测试的时候也发现了一些小问题，举几个例子：',
     '测试过程中也发现了若干细节问题，具体如下：'),
    ('用户看不懂是什么意思。后来在前端加了一个状态映射',
     '用户无法理解其含义。随后在前端增加了状态映射'),
    ('后来在前端用toFixed(2)处理了一下',
     '随后在前端使用toFixed(2)进行了格式化处理'),
    ('后来在前端加了一个日期选择器的', '随后在前端为日期选择器添加了'),
    ('这些都是比较小的问题，不影响核心功能，发现后很快就修好了。',
     '以上均为非核心问题，发现后已及时修复。'),
    # 第7章
    ('后端用Spring Boot 2.7加MyBatis来操作MySQL数据库，加了Redis做数据缓存，JWT做身份认证；前端用Vue 3和Element Plus来做页面，ECharts画数据图表。整个项目采用前后端分离的方式来开发。',
     '后端采用Spring Boot 2.7结合MyBatis操作MySQL数据库，引入Redis进行数据缓存，JWT实现身份认证；前端采用Vue 3和Element Plus构建页面，ECharts绑制数据图表。整个项目采用前后端分离的架构进行开发。'),
    ('数据库设计了9张表，后端编写了多个Controller接口，前端做了30多个页面组件。测试下来功能基本都没问题，可以满足中小型雪场的日常管理需要。',
     '数据库共设计9张数据表，后端编写了多个Controller接口，前端开发了30余个页面组件。经测试各项功能运行正常，能够满足中小型雪场的日常管理需求。'),
    ('做这个项目的过程中，我觉得收获还是比较大的。以前学Spring Boot和Vue的时候都是单独练习一些小例子，这次把它们放在一个完整的项目里，才真正体会到前后端联调、接口对接这些环节的复杂性。比如前端页面逻辑和后端数据格式不一致的时候要怎么处理，分页查询的参数怎么传递，跨域问题怎么解决等等，这些在做小练习的时候是碰不到的。当然了，系统肯定还有不少可以改进的地方。',
     '在本项目的开发过程中，笔者收获颇多。此前学习Spring Boot和Vue时多为独立的小型练习，本次将其整合于完整项目中，才深刻体会到前后端联调、接口对接等环节的复杂性。例如前端页面逻辑与后端数据格式不一致时的处理方式、分页查询参数的传递方式、跨域问题的解决方案等，这些在小型练习中难以遇到。当然，系统仍有诸多可改进之处。'),
    ('但因为时间和个人能力有限', '但受时间和个人能力所限'),
    ('有些地方以后还可以继续改进', '部分功能仍有优化空间'),
]

LQ = '\u201c'
RQ = '\u201d'

paragraph_replacements = {
    171: (
        '\u5546\u54c1\u79df\u8d41\u7684\u6574\u4f53\u4e1a\u52a1\u6d41\u7a0b\u5982\u4e0b\uff1a'
        '\u7528\u6237\u9996\u5148\u5728\u5546\u54c1\u5217\u8868\u9875\u9762\u6d4f\u89c8\u53ef\u79df\u8d41\u7684\u96ea\u5177\u548c\u88c5\u5907\uff0c'
        '\u9009\u62e9\u76ee\u6807\u5546\u54c1\u540e\u586b\u5199\u79df\u8d41\u65f6\u957f\u53ca\u6570\u91cf\uff0c'
        '\u968f\u540e\u63d0\u4ea4\u751f\u6210\u8ba2\u5355\u3002'
        '\u5728\u6b64\u8fc7\u7a0b\u4e2d\uff0c\u7cfb\u7edf\u4f1a\u81ea\u52a8\u6821\u9a8c\u8be5\u5546\u54c1\u7684\u5f53\u524d\u5e93\u5b58\u662f\u5426\u5145\u8db3\uff0c'
        '\u82e5\u5e93\u5b58\u6ee1\u8db3\u9700\u6c42\u5219\u6b63\u5e38\u751f\u6210\u8ba2\u5355\u5e76\u540c\u6b65\u6263\u51cf\u76f8\u5e94\u5e93\u5b58\u6570\u91cf\uff1b'
        '\u82e5\u5e93\u5b58\u4e0d\u8db3\uff0c\u5219\u63d0\u793a\u7528\u6237\u65e0\u6cd5\u5b8c\u6210\u79df\u8d41\u3002'
        '\u5f53\u7528\u6237\u4f7f\u7528\u5b8c\u6bd5\u6216\u5230\u8fbe\u7ea6\u5b9a\u5f52\u8fd8\u65f6\u95f4\u540e\uff0c'
        '\u5de5\u4f5c\u4eba\u5458\u5728\u540e\u53f0\u6267\u884c\u5f52\u8fd8\u786e\u8ba4\u64cd\u4f5c\uff0c'
        '\u7cfb\u7edf\u968f\u5373\u5c06\u8ba2\u5355\u72b6\u6001\u66f4\u65b0\u4e3a' + LQ + '\u5df2\u5f52\u8fd8' + RQ + '\uff0c'
        '\u5e76\u81ea\u52a8\u6062\u590d\u5bf9\u5e94\u5546\u54c1\u7684\u53ef\u7528\u5e93\u5b58\u6570\u91cf\u3002'
    ),
    172: (
        '\u7528\u6237\u767b\u5f55\u7684\u5904\u7406\u6d41\u7a0b\u5982\u4e0b\uff1a'
        '\u7528\u6237\u5728\u7cfb\u7edf\u767b\u5f55\u754c\u9762\u8f93\u5165\u8d26\u53f7\u548c\u5bc6\u7801\u540e\uff0c'
        '\u524d\u7aef\u9996\u5148\u8fdb\u884c\u8868\u5355\u6821\u9a8c\uff0c\u9a8c\u8bc1\u5404\u8f93\u5165\u9879\u662f\u5426\u4e3a\u7a7a\u3002'
        '\u6821\u9a8c\u901a\u8fc7\u540e\uff0c\u524d\u7aef\u901a\u8fc7HTTP\u8bf7\u6c42\u5c06\u767b\u5f55\u4fe1\u606f\u53d1\u9001\u81f3\u540e\u7aef\u8fdb\u884c\u8eab\u4efd\u9a8c\u8bc1\u3002'
        '\u540e\u7aef\u63a5\u6536\u5230\u8bf7\u6c42\u540e\uff0c\u67e5\u8be2\u6570\u636e\u5e93\u6838\u5b9e\u7528\u6237\u540d\u548c\u5bc6\u7801\u662f\u5426\u5339\u914d\uff0c'
        '\u9a8c\u8bc1\u901a\u8fc7\u540e\u751f\u6210JWT Token\u5e76\u8fd4\u56de\u7ed9\u524d\u7aef\u3002'
        '\u524d\u7aef\u6536\u5230Token\u540e\u5c06\u5176\u5b58\u50a8\u81f3\u6d4f\u89c8\u5668\u672c\u5730\u5b58\u50a8\uff0c'
        '\u968f\u540e\u6839\u636e\u7528\u6237\u89d2\u8272\u8df3\u8f6c\u81f3\u76f8\u5e94\u7684\u9996\u9875\u3002'
        '\u82e5\u540e\u7aef\u9a8c\u8bc1\u5931\u8d25\uff08\u5bc6\u7801\u9519\u8bef\u6216\u7528\u6237\u540d\u4e0d\u5b58\u5728\uff09\uff0c'
        '\u5219\u8fd4\u56de\u76f8\u5e94\u9519\u8bef\u4fe1\u606f\uff0c\u524d\u7aef\u5728\u9875\u9762\u4e0a\u5f39\u51fa\u63d0\u793a\u544a\u77e5\u7528\u6237\u3002'
    ),
    173: (
        '\u5728\u5b8c\u6210\u529f\u80fd\u6a21\u5757\u5212\u5206\u548c\u6570\u636e\u5e93\u8868\u7ed3\u6784\u8bbe\u8ba1\u4e4b\u540e\uff0c'
        '\u4e0b\u4e00\u6b65\u5de5\u4f5c\u662f\u68b3\u7406\u7cfb\u7edf\u4e2d\u6838\u5fc3\u4e1a\u52a1\u64cd\u4f5c\u7684\u6267\u884c\u6d41\u7a0b\uff0c'
        '\u4ee5\u4fbf\u5728\u7f16\u7801\u9636\u6bb5\u5177\u6709\u6e05\u6670\u7684\u5f00\u53d1\u601d\u8def\u3002'
        '\u672c\u8282\u4e3b\u8981\u9488\u5bf9\u7528\u6237\u767b\u5f55\u548c\u5546\u54c1\u79df\u8d41\u4e24\u4e2a\u6838\u5fc3\u4e1a\u52a1\u6d41\u7a0b\u8fdb\u884c\u4e86\u8be6\u7ec6\u7684\u5206\u6790\u4e0e\u8bbe\u8ba1\u3002'
    ),
    244: (
        '\u5728\u5de5\u4f5c\u4eba\u5458\u7684\u65e5\u5e38\u64cd\u4f5c\u4e2d\uff0c'
        '\u5f52\u8fd8\u5904\u7406\u662f\u6700\u4e3a\u9891\u7e41\u7684\u4e1a\u52a1\u64cd\u4f5c\u4e4b\u4e00\u3002'
        '\u5f53\u987e\u5ba2\u5f52\u8fd8\u79df\u8d41\u7684\u96ea\u5177\u65f6\uff0c'
        '\u5de5\u4f5c\u4eba\u5458\u9700\u5728\u7cfb\u7edf\u7684\u8ba2\u5355\u5217\u8868\u9875\u9762\u4e2d\u67e5\u627e\u5bf9\u5e94\u7684\u79df\u8d41\u8ba2\u5355\uff0c'
        '\u70b9\u51fb' + LQ + '\u786e\u8ba4\u5f52\u8fd8' + RQ + '\u6309\u94ae\u3002'
        '\u7cfb\u7edf\u63a5\u6536\u5230\u5f52\u8fd8\u786e\u8ba4\u64cd\u4f5c\u540e\uff0c'
        '\u5c06\u8be5\u8ba2\u5355\u72b6\u6001\u7531' + LQ + '\u4f7f\u7528\u4e2d' + RQ +
        '\u66f4\u65b0\u4e3a' + LQ + '\u5df2\u5f52\u8fd8' + RQ + '\uff0c'
        '\u540c\u65f6\u81ea\u52a8\u6062\u590d\u5bf9\u5e94\u5546\u54c1\u7684\u53ef\u7528\u5e93\u5b58\u6570\u91cf\u3002'
        '\u4ee5\u4e0b\u4e3a\u540e\u7aef\u5904\u7406\u5f52\u8fd8\u64cd\u4f5c\u7684\u6838\u5fc3\u4ee3\u7801\uff1a'
    ),
    247: (
        '\u5728\u9875\u9762\u5e03\u5c40\u65b9\u9762\uff0c'
        '\u5de5\u4f5c\u4eba\u5458\u540e\u53f0\u754c\u9762\u91c7\u7528\u4e0e\u7ba1\u7406\u5458\u76f8\u540c\u7684\u5de6\u4fa7\u5bfc\u822a\u83dc\u5355\u52a0\u53f3\u4fa7\u5185\u5bb9\u533a\u57df\u7684\u5e03\u5c40\u65b9\u5f0f\uff0c'
        '\u4f46\u83dc\u5355\u9879\u5df2\u6839\u636e\u89d2\u8272\u6743\u9650\u8fdb\u884c\u4e86\u7cbe\u7b80\uff0c'
        '\u4ec5\u4fdd\u7559\u4e0e\u65e5\u5e38\u5de5\u4f5c\u76f4\u63a5\u76f8\u5173\u7684\u529f\u80fd\u6a21\u5757\u3002'
        '\u9875\u9762\u9876\u90e8\u5c55\u793a\u5f53\u524d\u767b\u5f55\u7684\u5de5\u4f5c\u4eba\u5458\u59d3\u540d\u53ca\u89d2\u8272\u6807\u8bc6\uff0c'
        '\u5de6\u4fa7\u8fb9\u680f\u6392\u5217' + LQ + '\u5546\u54c1\u7ba1\u7406' + RQ +
        LQ + '\u79df\u8d41\u8ba2\u5355' + RQ +
        LQ + '\u573a\u5730\u4fe1\u606f' + RQ + '\u7b49\u529f\u80fd\u5165\u53e3\u3002'
        '\u5546\u54c1\u7ba1\u7406\u9875\u9762\u4ee5\u8868\u683c\u5f62\u5f0f\u5c55\u793a\u6240\u6709\u5546\u54c1\u4fe1\u606f\uff0c'
        '\u5e76\u63d0\u4f9b\u6309\u540d\u79f0\u548c\u5206\u7c7b\u8fdb\u884c\u68c0\u7d22\u7684\u529f\u80fd\uff1b'
        '\u79df\u8d41\u8ba2\u5355\u9875\u9762\u5c55\u793a\u7cfb\u7edf\u4e2d\u6240\u6709\u8ba2\u5355\u8bb0\u5f55\uff0c'
        '\u652f\u6301\u6309\u8ba2\u5355\u72b6\u6001\u8fdb\u884c\u7b5b\u9009\uff0c'
        '\u5de5\u4f5c\u4eba\u5458\u53ef\u5bf9' + LQ + '\u4f7f\u7528\u4e2d' + RQ +
        '\u72b6\u6001\u7684\u8ba2\u5355\u76f4\u63a5\u6267\u884c\u5f52\u8fd8\u64cd\u4f5c\u3002'
    ),
    248: (
        '\u5de5\u4f5c\u4eba\u5458\u5728\u7cfb\u7edf\u4e2d\u7684\u64cd\u4f5c\u6743\u9650\u4e3b\u8981\u5305\u62ec\u4ee5\u4e0b\u51e0\u4e2a\u65b9\u9762\uff1a'
        '\u4e00\u662f\u67e5\u9605\u6240\u6709\u79df\u8d41\u8ba2\u5355\u5217\u8868\uff0c'
        '\u5e76\u53ef\u6839\u636e\u8ba2\u5355\u72b6\u6001\uff08\u5f85\u652f\u4ed8\u3001\u4f7f\u7528\u4e2d\u3001\u5df2\u5f52\u8fd8\u7b49\uff09\u8fdb\u884c\u7b5b\u9009\u67e5\u8be2\uff1b'
        '\u4e8c\u662f\u4e3a\u5230\u5e97\u987e\u5ba2\u529e\u7406\u96ea\u5177\u6216\u88c5\u5907\u7684\u79df\u8d41\u624b\u7eed\uff0c'
        '\u9009\u62e9\u5546\u54c1\u540e\u5f55\u5165\u79df\u8d41\u4fe1\u606f\u5e76\u751f\u6210\u8ba2\u5355\uff1b'
        '\u4e09\u662f\u5728\u987e\u5ba2\u5f52\u8fd8\u5546\u54c1\u65f6\uff0c'
        '\u901a\u8fc7\u7cfb\u7edf\u786e\u8ba4\u5f52\u8fd8\u64cd\u4f5c\uff0c'
        '\u7cfb\u7edf\u81ea\u52a8\u66f4\u65b0\u8ba2\u5355\u72b6\u6001\u5e76\u6062\u590d\u5546\u54c1\u5e93\u5b58\u6570\u91cf\uff1b'
        '\u56db\u662f\u67e5\u770b\u5404\u7c7b\u5546\u54c1\u7684\u5b9e\u65f6\u5e93\u5b58\u72b6\u51b5\uff0c'
        '\u4ee5\u4fbf\u53ca\u65f6\u638c\u63e1\u53ef\u4f9b\u79df\u8d41\u7684\u96ea\u5177\u6570\u91cf\u3002'
    ),
    249: (
        '\u5728\u6ed1\u96ea\u573a\u7684\u65e5\u5e38\u8fd0\u8425\u7ba1\u7406\u4e2d\uff0c'
        '\u5de5\u4f5c\u4eba\u5458\u627f\u62c5\u7740\u91cd\u8981\u7684\u524d\u53f0\u4e1a\u52a1\u804c\u8d23\uff0c'
        '\u4e3b\u8981\u8d1f\u8d23\u4e3a\u987e\u5ba2\u529e\u7406\u96ea\u5177\u548c\u88c5\u5907\u7684\u79df\u8d41\u624b\u7eed\u4ee5\u53ca\u5904\u7406\u76f8\u5173\u8ba2\u5355\u4e8b\u52a1\u3002'
        '\u5de5\u4f5c\u4eba\u5458\u4f7f\u7528\u4e2a\u4eba\u8d26\u53f7\u767b\u5f55\u7cfb\u7edf\u540e\u53f0\u540e\uff0c'
        '\u7cfb\u7edf\u6839\u636e\u5176\u89d2\u8272\u6743\u9650\u5c55\u793a\u76f8\u5e94\u7684\u529f\u80fd\u83dc\u5355\uff0c'
        '\u5305\u62ec\u5546\u54c1\u4fe1\u606f\u7ba1\u7406\u3001\u79df\u8d41\u8ba2\u5355\u5904\u7406\u53ca\u573a\u5730\u72b6\u51b5\u67e5\u770b\u7b49\u4e3b\u8981\u529f\u80fd\u6a21\u5757\u3002'
    ),
}


def replace_in_paragraph(paragraph, old_text, new_text):
    full_text = paragraph.text
    if old_text not in full_text:
        return False
    if len(paragraph.runs) == 0:
        return False
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = paragraph.runs[0].text.replace(old_text, new_text)
        return True
    run_texts = [r.text for r in paragraph.runs]
    joined = ''.join(run_texts)
    if old_text not in joined:
        return False
    new_joined = joined.replace(old_text, new_text, 1)
    paragraph.runs[0].text = new_joined
    for r in paragraph.runs[1:]:
        r.text = ''
    return True


def replace_full_paragraph(paragraph, new_text):
    if len(paragraph.runs) == 0:
        return False
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ''
    return True


# Apply phrase replacements
phrase_count = 0
for old_t, new_t in phrase_replacements:
    for p in doc.paragraphs:
        if old_t in p.text:
            if replace_in_paragraph(p, old_t, new_t):
                phrase_count += 1

print(f'phrase replacements done: {phrase_count}')

# Apply full paragraph replacements
para_count = 0
for idx, new_text in paragraph_replacements.items():
    p = doc.paragraphs[idx]
    if replace_full_paragraph(p, new_text):
        para_count += 1

print(f'paragraph replacements done: {para_count}')


# ============================================================
# STEP 2: Reference reordering
# ============================================================

citation_pattern = re.compile(r'\[(\d+)\]')
first_appearance = {}

for i, p in enumerate(doc.paragraphs):
    text = p.text
    if p.style.name == 'Heading 1' and '\u53c2\u8003\u6587\u732e' in text:
        break
    matches = citation_pattern.findall(text)
    for m in matches:
        num = int(m)
        if num not in first_appearance:
            first_appearance[num] = i

ordered_refs = sorted(first_appearance.items(), key=lambda x: x[1])
old_to_new = {}
for new_num, (old_num, _) in enumerate(ordered_refs, 1):
    old_to_new[old_num] = new_num

print(f'ref mapping: {old_to_new}')

# Find reference list
ref_start = None
ref_paragraphs = {}
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and '\u53c2\u8003\u6587\u732e' in p.text:
        ref_start = i
        continue
    if ref_start is not None:
        match = re.match(r'^\[(\d+)\]', p.text)
        if match:
            ref_num = int(match.group(1))
            ref_paragraphs[ref_num] = p.text
        elif p.style.name == 'Heading 1':
            break

print(f'found {len(ref_paragraphs)} references')

# Rename citations in text: two-pass to avoid collision
# Pass 1: [N] -> <<TEMP_N>>
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and '\u53c2\u8003\u6587\u732e' in p.text:
        break
    for run in p.runs:
        for old_num in sorted(old_to_new.keys(), reverse=True):
            marker = f'[{old_num}]'
            if marker in run.text:
                run.text = run.text.replace(marker, f'<<TEMP_{old_num}>>')

# Pass 2: <<TEMP_N>> -> [new_N]
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and '\u53c2\u8003\u6587\u732e' in p.text:
        break
    for run in p.runs:
        for old_num, new_num in old_to_new.items():
            temp = f'<<TEMP_{old_num}>>'
            if temp in run.text:
                run.text = run.text.replace(temp, f'[{new_num}]')

print('text citations updated')

# Reorder reference list
new_ref_list = {}
for old_num, text in ref_paragraphs.items():
    new_num = old_to_new.get(old_num, old_num)
    new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', text)
    new_ref_list[new_num] = new_text

ref_idx = 0
for i, p in enumerate(doc.paragraphs):
    if ref_start is not None and i > ref_start:
        match = re.match(r'^\[(\d+)\]', p.text)
        if match:
            ref_idx += 1
            if ref_idx in new_ref_list:
                replace_full_paragraph(p, new_ref_list[ref_idx])

print('reference list reordered')


# ============================================================
# STEP 3: Make citations superscript
# ============================================================

def apply_superscript_to_citations(document):
    ref_section = False
    count = 0

    for p in document.paragraphs:
        if p.style.name == 'Heading 1' and '\u53c2\u8003\u6587\u732e' in p.text:
            ref_section = True
        if ref_section:
            continue

        full_text = p.text
        if not citation_pattern.search(full_text):
            continue

        if re.match(r'^\s*\[\d+\]', full_text):
            continue

        code_keywords = ['@PostMapping', '@GetMapping', '@PutMapping', '@DeleteMapping',
                         'public ', 'private ', 'return ', 'const ', 'router.',
                         'import ', 'claims.put', 'data.put', 'Map<', 'Result<',
                         'next(', 'userStore']
        if any(kw in full_text for kw in code_keywords):
            continue

        if not p.runs:
            continue

        first_run = p.runs[0]

        segments = []
        remaining = full_text
        while remaining:
            m = citation_pattern.search(remaining)
            if not m:
                if remaining:
                    segments.append((remaining, False))
                break

            start, end = m.span()
            if start > 0:
                segments.append((remaining[:start], False))

            cite_text = remaining[start:end]
            pos = end
            while pos < len(remaining):
                next_m = citation_pattern.match(remaining[pos:])
                if next_m:
                    cite_text += remaining[pos:pos + next_m.end()]
                    pos += next_m.end()
                else:
                    break

            segments.append((cite_text, True))
            remaining = remaining[pos:]

        if not any(is_cite for _, is_cite in segments):
            continue

        p_element = p._element
        for run in list(p.runs):
            p_element.remove(run._element)

        for text, is_citation in segments:
            if not text:
                continue
            new_run = p.add_run(text)
            if first_run.font.name:
                new_run.font.name = first_run.font.name
            if first_run.font.size:
                new_run.font.size = first_run.font.size
            if first_run.font.bold is not None:
                new_run.font.bold = first_run.font.bold
            if first_run.font.italic is not None:
                new_run.font.italic = first_run.font.italic

            src_rpr = first_run._element.find(qn('w:rPr'))
            if src_rpr is not None:
                src_rFonts = src_rpr.find(qn('w:rFonts'))
                if src_rFonts is not None:
                    eastAsia = src_rFonts.get(qn('w:eastAsia'))
                    if eastAsia:
                        new_rpr = new_run._element.get_or_add_rPr()
                        new_rFonts = new_rpr.find(qn('w:rFonts'))
                        if new_rFonts is None:
                            new_rFonts = docx.oxml.OxmlElement('w:rFonts')
                            new_rpr.insert(0, new_rFonts)
                        new_rFonts.set(qn('w:eastAsia'), eastAsia)

            if is_citation:
                new_run.font.superscript = True
                count += 1

    return count

sup_count = apply_superscript_to_citations(doc)
print(f'superscript done: {sup_count} citations')


# ============================================================
# SAVE
# ============================================================

output_path = r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx'
doc.save(output_path)
print(f'saved: {output_path}')
print('all done!')
