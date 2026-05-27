# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
import os

SCRIPT_DIR = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现"

# Re-generate from original
exec(open(os.path.join(SCRIPT_DIR, "fix_doc_v3.py"), encoding='utf-8').read())

doc = Document(os.path.join(SCRIPT_DIR, "5.26.2.docx"))

# Map caption text to image file
img_map = {
    "系统功能结构图": ("structure_chart.png", Cm(15.5)),
    "登录流程图": ("login_flow.png", Cm(12)),
    "租赁业务流程图": ("rental_flow.png", Cm(12)),
}

processed = set()

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for caption_key, (fname, width) in img_map.items():
        if caption_key not in text:
            continue
        
        img_path = os.path.join(SCRIPT_DIR, fname)
        
        # Search for adjacent image paragraph (prefer after, then before)
        for offset in [1, -1, 2, -2]:
            adj_idx = i + offset
            if adj_idx in processed or adj_idx < 0 or adj_idx >= len(doc.paragraphs):
                continue
            
            adj_p = doc.paragraphs[adj_idx]
            has_img = any(
                len(run._element.findall(qn('w:drawing'))) > 0 or
                len(run._element.findall(qn('w:pict'))) > 0
                for run in adj_p.runs
            )
            
            if has_img:
                for run in adj_p.runs:
                    run._element.getparent().remove(run._element)
                run = adj_p.add_run()
                run.add_picture(img_path, width=width)
                processed.add(adj_idx)
                print(f"  Replaced P{adj_idx} for '{caption_key}' ({fname})")
                break
        break

doc.save(os.path.join(SCRIPT_DIR, "5.26.2.docx"))
print("Saved!")
