# -*- coding: utf-8 -*-
"""
5.26.1.docx modification script - v3 (fixed element move logic)
Key fix: collect removal targets BEFORE insertions to avoid index shift
"""
import copy
import re
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

SRC = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx"
DST = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"

doc = Document(SRC)
body = doc.element.body


def para_replace(para, old, new):
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not para.runs:
        return False
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ""
    return True


def set_para_spacing(para, before_pt=0, after_pt=0):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))


def get_body_children_between(start_elem, end_elem):
    """Get all body children between start and end (inclusive)."""
    collecting = False
    elements = []
    for child in list(body):
        if child is start_elem:
            collecting = True
        if collecting:
            elements.append(child)
        if child is end_elem:
            break
    return elements


def find_heading(text_pattern, start_idx=0):
    """Find heading paragraph index containing text_pattern."""
    for i in range(start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style and p.style.name.startswith("Heading"):
            if text_pattern in p.text:
                return i
    return -1


# ============================================================
# STEP 1: Text fixes (6.4, 7.1)
# ============================================================
print("=== Step 1: Text fixes ===")

# 6.4
p408 = doc.paragraphs[408]
para_replace(p408, "还有一个问题是教练预约", "此外，教练预约")
para_replace(p408, "理论上不应该允许", "不应允许")

p409 = doc.paragraphs[409]
para_replace(p409, "还有一个在测试中发现的问题是关于并发的。", "此外，在并发控制方面也存在待优化之处。")
para_replace(p409, "理论上应该在数据库层面用乐观锁来处理", "应在数据库层面采用乐观锁机制加以处理")
para_replace(p409, "不过因为", "由于")
para_replace(p409, "所以这个问题在当前阶段影响不大", "该问题在当前阶段影响有限")
para_replace(p409, "但如果以后用户量增加的话，这个地方是需要优化的", "但随着用户规模扩大，该部分仍有优化的必要")

p410 = doc.paragraphs[410]
para_replace(p410, "都能正常走通", "均可正常运行")

# 7.1
p415 = doc.paragraphs[415]
para_replace(p415, "笔者收获颇多", "收获颇多")
para_replace(p415, "这些在小型练习中难以遇到", "这些在课程中难以遇到")

print("  done")

# ============================================================
# STEP 2: Reference reorder
# ============================================================
print("=== Step 2: References ===")

ref_map = {1:1, 2:2, 3:3, 4:4, 5:5, 6:15, 7:6, 8:7, 9:8, 10:9,
           11:10, 12:11, 13:12, 14:13, 15:14}

for i in range(56, 422):
    p = doc.paragraphs[i]
    text = p.text
    if not re.search(r'\[\d+\]', text):
        continue
    new_text = text
    for old_num in sorted(ref_map.keys(), reverse=True):
        new_text = new_text.replace(f'[{old_num}]', f'[###{ref_map[old_num]}###]')
    new_text = re.sub(r'\[###(\d+)###\]', r'[\1]', new_text)
    if new_text != text and p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""

doc.paragraphs[75]
para_replace(doc.paragraphs[75], "[6][4]", "[6]")
para_replace(doc.paragraphs[88], "[5][12]", "[12]")

ref_entries = {}
for i in range(423, 438):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    text = p.text
    match = re.match(r'\[(\d+)\]', text)
    if match:
        old_num = int(match.group(1))
        new_num = ref_map.get(old_num, old_num)
        new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', text)
        ref_entries[new_num] = new_text

for idx, new_num in enumerate(sorted(ref_entries.keys())):
    target_idx = 423 + idx
    if target_idx < len(doc.paragraphs):
        p = doc.paragraphs[target_idx]
        if p.runs:
            p.runs[0].text = ref_entries[new_num]
            for r in p.runs[1:]:
                r.text = ""

print("  done")

# ============================================================
# STEP 3: Table spacing
# ============================================================
print("=== Step 3: Table spacing ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if re.match(r'^(表[46]\.\d|续表[46]\.\d)', text):
        set_para_spacing(p, before_pt=3, after_pt=3)
print("  done")

# ============================================================
# STEP 4: Merge 3.4 into 3.2
# ============================================================
print("=== Step 4: Merge 3.4 into 3.2 ===")

# Strategy: use element references (not indices) to avoid shift issues
# Save all element references BEFORE any modifications

# 3.2.x target elements (after which to insert)
target_p103 = doc.paragraphs[103]._element  # after 3.2.1 管理员功能 text
target_p105 = doc.paragraphs[105]._element  # after 3.2.2 教练功能 text
target_p107 = doc.paragraphs[107]._element  # after 3.2.3 工作人员功能 text
target_p109 = doc.paragraphs[109]._element  # after 3.2.4 普通用户功能 text

# 3.4 section elements to remove (entire 3.4 section: P116 heading through P134)
sec34_heading = doc.paragraphs[116]._element  # "3.4 用例分析"
sec34_intro = doc.paragraphs[117]._element    # intro text
sec34_end = doc.paragraphs[134]._element      # "图3.5 普通用户用例图"

# Collect ALL elements in section 3.4
sec34_all = get_body_children_between(sec34_heading, sec34_end)

# Collect subsection elements (excluding heading of each subsection)
# 3.4.1 管理员用例: P118(heading), P119(表标题), [TABLE], P120(空/图), P121(图), P122(图标题)
h341 = doc.paragraphs[118]._element
h342 = doc.paragraphs[123]._element
h343 = doc.paragraphs[127]._element
h344 = doc.paragraphs[131]._element

# Collect elements for each subsection (between headings)
def get_subsection_content(heading_elem, next_heading_elem):
    """Get elements between heading and next heading, EXCLUDING the heading itself."""
    elements = []
    collecting = False
    for child in list(body):
        if child is heading_elem:
            collecting = True
            continue  # skip heading itself
        if collecting:
            if child is next_heading_elem:
                break
            elements.append(child)
    return elements

sub1_content = get_subsection_content(h341, h342)  # 管理员用例 content
sub2_content = get_subsection_content(h342, h343)  # 教练用例 content
sub3_content = get_subsection_content(h343, h344)  # 工作人员用例 content
# 普通用户: from h344 to end of section (sec34_end)
sub4_content = []
collecting = False
for child in list(body):
    if child is h344:
        collecting = True
        continue
    if collecting:
        sub4_content.append(child)
    if child is sec34_end:
        break

print(f"  sub1: {len(sub1_content)} elems, sub2: {len(sub2_content)}, sub3: {len(sub3_content)}, sub4: {len(sub4_content)}")

# Now insert copies at target positions (reverse order to preserve positions)
# 4. After P109 (普通用户), insert sub4
for elem in reversed(sub4_content):
    target_p109.addnext(copy.deepcopy(elem))

# 3. After P107 (工作人员), insert sub3
for elem in reversed(sub3_content):
    target_p107.addnext(copy.deepcopy(elem))

# 2. After P105 (教练), insert sub2
for elem in reversed(sub2_content):
    target_p105.addnext(copy.deepcopy(elem))

# 1. After P103 (管理员), insert sub1
for elem in reversed(sub1_content):
    target_p103.addnext(copy.deepcopy(elem))

print("  Inserted use case content after each 3.2.x section")

# Remove entire 3.4 section using saved element references
for elem in sec34_all:
    body.remove(elem)

print(f"  Removed 3.4 section ({len(sec34_all)} elements)")

# ============================================================
# STEP 5: Swap 4.4 and 4.5 (fresh index lookup after 3.4 changes)
# ============================================================
print("=== Step 5: Swap 4.4/4.5 ===")

h44 = find_heading("4.4")
h45 = find_heading("4.5")
h46 = find_heading("4.6")

print(f"  4.4 at P{h44}, 4.5 at P{h45}, 4.6 at P{h46}")

if h44 >= 0 and h45 >= 0 and h46 >= 0:
    # Save element references
    h44_elem = doc.paragraphs[h44]._element
    h45_elem = doc.paragraphs[h45]._element
    h46_elem = doc.paragraphs[h46]._element
    
    # Find flow paragraphs in 4.4 (the last 3 before 4.5)
    flow_keywords = ["商品租赁的整体业务流程", "用户登录的处理流程", "在完成功能模块划分"]
    flow_elems = []
    for i in range(h44, h45):
        p = doc.paragraphs[i]
        if any(kw in p.text for kw in flow_keywords):
            flow_elems.append(p._element)
    
    print(f"  Flow paragraphs found: {len(flow_elems)}")
    
    # Collect all 4.5 section elements (from 4.5 heading to just before 4.6)
    sec45 = get_body_children_between(h45_elem, h46_elem)
    sec45 = sec45[:-1]  # exclude 4.6 heading
    
    # Insert 4.5 content before 4.4
    for elem in reversed(sec45):
        h44_elem.addprevious(copy.deepcopy(elem))
    
    # Remove original 4.5 section
    for elem in sec45:
        body.remove(elem)
    
    print("  Moved 4.5 before 4.4")
    
    # Move flow paragraphs from 4.4 to after the new 4.5 heading (now before 4.4)
    # Find new 4.5 heading position
    new_h45 = find_heading("4.5")
    if new_h45 >= 0:
        new_h45_elem = doc.paragraphs[new_h45]._element
        # Move flow paragraphs after 4.5 heading
        for elem in reversed(flow_elems):
            new_h45_elem.addnext(copy.deepcopy(elem))
            body.remove(elem)
        print("  Moved flow paragraphs to 4.5")
    
    # Swap section numbers: what's now 4.5(业务流程) becomes 4.4, and 4.4(数据库) becomes 4.5
    h45_pos = find_heading("4.5")
    h44_pos = find_heading("4.4")
    if h45_pos >= 0 and h44_pos >= 0 and h45_pos < h44_pos:
        para_replace(doc.paragraphs[h45_pos], "4.5", "4.4")
        para_replace(doc.paragraphs[h44_pos], "4.4", "4.5")
        
        # Fix subsection numbers: 4.4.1->4.5.1, 4.4.2->4.5.2
        for i in range(h44_pos, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            if p.style and p.style.name == "Heading 2" and i != h44_pos:
                break
            if p.style and p.style.name == "Heading 3":
                if "4.4.1" in p.text:
                    para_replace(p, "4.4.1", "4.5.1")
                elif "4.4.2" in p.text:
                    para_replace(p, "4.4.2", "4.5.2")
        
        print("  Section numbers swapped")


# ============================================================
# STEP 6: Chapter 5 reorganization
# ============================================================
print("=== Step 6: Ch5 reorg ===")

ch5_start = find_heading("第5章")
ch6_start = find_heading("第6章")

if ch5_start >= 0 and ch6_start >= 0:
    # Find all H2 sections within Ch5
    ch5_h2 = {}
    for i in range(ch5_start, ch6_start):
        p = doc.paragraphs[i]
        if p.style and p.style.name == "Heading 2":
            for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
                if key in p.text and key not in ch5_h2:
                    ch5_h2[key] = i
                    break
    
    print(f"  Ch5 sections: {ch5_h2}")
    
    if all(k in ch5_h2 for k in ["5.1","5.2","5.3","5.4","5.5","5.6","5.7"]):
        # Save element references before any moves
        h51_elem = doc.paragraphs[ch5_h2["5.1"]]._element
        h52_elem = doc.paragraphs[ch5_h2["5.2"]]._element
        h53_elem = doc.paragraphs[ch5_h2["5.3"]]._element
        h54_elem = doc.paragraphs[ch5_h2["5.4"]]._element
        h55_elem = doc.paragraphs[ch5_h2["5.5"]]._element
        h56_elem = doc.paragraphs[ch5_h2["5.6"]]._element
        h57_elem = doc.paragraphs[ch5_h2["5.7"]]._element
        
        # Collect 5.3 section elements (from 5.3 to 5.4)
        sec53 = get_body_children_between(h53_elem, h54_elem)
        sec53 = sec53[:-1]
        
        # Collect 5.6 section elements (from 5.6 to 5.7)
        sec56 = get_body_children_between(h56_elem, h57_elem)
        sec56 = sec56[:-1]
        
        # Move 5.6 (data stats) before 5.2
        for elem in reversed(sec56):
            h52_elem.addprevious(copy.deepcopy(elem))
        for elem in sec56:
            body.remove(elem)
        print("  Moved 5.6 before 5.2")
        
        # Move 5.3 (product mgmt) before 5.6 copy (which is now before 5.2)
        # Need fresh reference - find 5.6's new position
        new_h56 = None
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name == "Heading 2" and "5.6" in p.text:
                new_h56 = p._element
                break
        
        if new_h56 is not None:
            for elem in reversed(sec53):
                new_h56.addprevious(copy.deepcopy(elem))
            for elem in sec53:
                body.remove(elem)
            print("  Moved 5.3 before 5.6")
        
        # Now rename all sections
        # Current order: 5.1(admin) → 5.3(product) → 5.6(data) → 5.2(user) → 5.4(coach) → 5.5(staff) → 5.7(security)
        # Target:        5.1(admin) → 5.1.1(product) → 5.1.2(data) → 5.2(user) → 5.3(coach) → 5.4(staff) → 5.5(security)
        
        ch5_new = find_heading("第5章")
        ch6_new = find_heading("第6章")
        
        rename_rules = [
            ("5.1", "5.1　管理员模块的实现", "Heading 2"),
            ("5.3", "5.1.1　商品管理功能", "Heading 3"),
            ("5.6", "5.1.2　数据统计功能", "Heading 3"),
            ("5.2", "5.2　用户模块的实现", "Heading 2"),
            ("5.4", "5.3　教练模块的实现", "Heading 2"),
            ("5.5", "5.4　工作人员模块的实现", "Heading 2"),
            ("5.7", "5.5　系统安全实现", "Heading 2"),
        ]
        
        for i in range(ch5_new, ch6_new):
            p = doc.paragraphs[i]
            if not (p.style and p.style.name.startswith("Heading")):
                continue
            for old_key, new_title, target_style in rename_rules:
                if old_key in p.text:
                    if p.runs:
                        p.runs[0].text = new_title
                        for r in p.runs[1:]:
                            r.text = ""
                    if target_style == "Heading 3" and p.style.name != "Heading 3":
                        p.style = doc.styles["Heading 3"]
                    print(f"  P{i}: {old_key} -> {new_title}")
                    break
        
        # Add intro text to 5.1
        h51_new = find_heading("管理员模块")
        if h51_new >= 0 and h51_new + 1 < len(doc.paragraphs):
            intro_p = doc.paragraphs[h51_new + 1]
            old_text = intro_p.text
            new_text = "管理员是系统的最高权限角色，负责系统整体的数据管理和运营配置。在介绍管理员的具体功能之前，先简要说明系统的开发环境。" + old_text
            if intro_p.runs:
                intro_p.runs[0].text = new_text
                for r in intro_p.runs[1:]:
                    r.text = ""
        
        # Add intro text to 5.2
        h52_new = find_heading("用户模块")
        if h52_new >= 0 and h52_new + 1 < len(doc.paragraphs):
            intro_p = doc.paragraphs[h52_new + 1]
            old_text = intro_p.text
            new_text = "普通用户是系统的最终使用者，通过前台页面访问各项服务功能。" + old_text
            if intro_p.runs:
                intro_p.runs[0].text = new_text
                for r in intro_p.runs[1:]:
                    r.text = ""


# ============================================================
# STEP 7: Reduce Ch5 text
# ============================================================
print("=== Step 7: Ch5 text reduction ===")

for i, p in enumerate(doc.paragraphs):
    text = p.text
    if "工作人员在系统中的操作权限主要包括以下几个方面" in text:
        if p.runs:
            p.runs[0].text = "工作人员的主要操作权限涵盖租赁订单管理、商品信息查看、场地状态管理等方面。"
            for r in p.runs[1:]:
                r.text = ""

    if "在滑雪场的日常运营管理中，工作人员承担着重要的前台业务职责" in text:
        if p.runs:
            p.runs[0].text = "工作人员在系统中主要负责前台业务操作，登录后台后可使用商品信息管理、租赁订单处理及场地状态查看等功能。"
            for r in p.runs[1:]:
                r.text = ""

    if "在页面布局方面，工作人员后台界面采用" in text:
        if p.runs:
            p.runs[0].text = "工作人员后台采用左侧导航加右侧内容区域的布局方式，菜单项根据角色权限进行了精简。商品管理页面提供按名称和分类检索的功能，租赁订单页面支持按状态筛选并可直接执行归还操作。"
            for r in p.runs[1:]:
                r.text = ""

print("  done")

# ============================================================
# Save
# ============================================================
print("=== Saving ===")
doc.save(DST)
print(f"Saved to {DST}")
print("DONE")
