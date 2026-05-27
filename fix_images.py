# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
import os
import copy

DOC_PATH = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx"
SCRIPT_DIR = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现"

# Re-run the full modification on the original file
# Then insert images correctly

# First, run the v3 script to get the modified doc without image issues
exec(open(os.path.join(SCRIPT_DIR, "fix_doc_v3.py"), encoding='utf-8').read())

# Now work on the already-saved 5.26.2.docx
doc = Document(os.path.join(SCRIPT_DIR, "5.26.2.docx"))
body = doc.element.body

# Find all image paragraphs and their captions
replacements = {
    "系统功能结构图": os.path.join(SCRIPT_DIR, "structure_chart.png"),
    "登录流程图": os.path.join(SCRIPT_DIR, "login_flow.png"),
    "租赁业务流程图": os.path.join(SCRIPT_DIR, "rental_flow.png"),
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for caption_key, img_path in replacements.items():
        if caption_key in text:
            # Check adjacent paragraphs for image
            for offset in [-1, 1, -2, 2]:
                adj_idx = i + offset
                if 0 <= adj_idx < len(doc.paragraphs):
                    adj_p = doc.paragraphs[adj_idx]
                    has_img = any(
                        run._element.findall(qn('w:drawing')) or
                        run._element.findall(qn('w:pict'))
                        for run in adj_p.runs
                    )
                    if has_img:
                        # Clear and replace
                        for run in adj_p.runs:
                            run._element.getparent().remove(run._element)
                        run = adj_p.add_run()
                        
                        if "结构图" in caption_key:
                            run.add_picture(img_path, width=Cm(15.5))
                        else:
                            run.add_picture(img_path, width=Cm(12))
                        
                        print(f"  Replaced P{adj_idx} for '{caption_key}'")
                        break
            break

OUT = os.path.join(SCRIPT_DIR, "5.26.2.docx")
doc.save(OUT)
print(f"Saved: {OUT}")
