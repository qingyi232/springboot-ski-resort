# -*- coding: utf-8 -*-
import subprocess,sys,os,re
try:
    from docx import Document
    from docx.shared import Pt,Cm,Mm,RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable,"-m","pip","install","python-docx"])
    from docx import Document
    from docx.shared import Pt,Cm,Mm,RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

doc = Document()
SP = '\u3000'
AL_C = WD_ALIGN_PARAGRAPH.CENTER
AL_J = WD_ALIGN_PARAGRAPH.JUSTIFY
AL_L = WD_ALIGN_PARAGRAPH.LEFT

def srun(r, cn='宋体', en='Times New Roman', sz=Pt(12), b=False):
    r.font.name = en; r.font.size = sz; r.font.bold = b
    r.font.color.rgb = RGBColor(0,0,0); r.font.italic = False; r.font.underline = False
    rPr = r.element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None: rF = OxmlElement('w:rFonts'); rPr.append(rF)
    rF.set(qn('w:eastAsia'), cn); rF.set(qn('w:ascii'), en); rF.set(qn('w:hAnsi'), en)

def spf(p, al=None, sb=None, sa=None, ls=None, fi=None, sgl=False, kn=False):
    pf = p.paragraph_format
    if al: pf.alignment = al
    if sb is not None: pf.space_before = sb
    if sa is not None: pf.space_after = sa
    if ls: pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.line_spacing = ls
    if sgl: pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if fi is not None: pf.first_line_indent = fi
    if kn: pf.keep_with_next = True

def B(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    srun(r); spf(p, AL_J, Pt(0), Pt(0), Pt(20), Pt(24))
    return p

def BR(text):
    p = doc.add_paragraph()
    parts = re.split(r'(\[\d+\])', text)
    for part in parts:
        if re.match(r'\[\d+\]', part):
            r = p.add_run(part); srun(r,'Times New Roman','Times New Roman',Pt(9)); r.font.superscript = True
        elif part:
            r = p.add_run(part); srun(r)
    spf(p, AL_J, Pt(0), Pt(0), Pt(20), Pt(24))
    return p

def CB(code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(); r = p.add_run(line if line.strip() else ' ')
        srun(r, 'Consolas', 'Consolas', Pt(9)); spf(p, AL_L, Pt(0), Pt(0), Pt(14))

def stitle(text, cn='黑体', en='黑体', sz=Pt(16)):
    p = doc.add_paragraph(); r = p.add_run(text)
    srun(r, cn, en, sz, False); spf(p, AL_C, Pt(40), Pt(20), Pt(20))

def tcap(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    srun(r, '黑体', '黑体', Pt(10.5), False); spf(p, AL_C, Pt(6), Pt(3), Pt(20), kn=True)

def fcap(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    srun(r, '黑体', '黑体', Pt(10.5), False); spf(p, AL_C, Pt(6), Pt(12), sgl=True)

def fimg(hint='此处插入图片'):
    p = doc.add_paragraph(); r = p.add_run(f'\n\n（{hint}）\n\n')
    srun(r, '宋体', 'Times New Roman', Pt(10.5)); r.font.color.rgb = RGBColor(150,150,150)
    spf(p, AL_C, Pt(6), Pt(0), sgl=True)

def three_line(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    for c in list(tblPr):
        if c.tag == qn('w:tblBorders'): tblPr.remove(c)
    bds = OxmlElement('w:tblBorders')
    for nm, sz in [('top','12'),('bottom','12')]:
        e = OxmlElement(f'w:{nm}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); bds.append(e)
    for nm in ['left','right','insideH','insideV']:
        e = OxmlElement(f'w:{nm}'); e.set(qn('w:val'),'none'); e.set(qn('w:sz'),'0'); e.set(qn('w:space'),'0'); bds.append(e)
    tblPr.append(bds)
    for cell in table.rows[0].cells:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders'); bb = OxmlElement('w:bottom')
        bb.set(qn('w:val'),'single'); bb.set(qn('w:sz'),'8'); bb.set(qn('w:space'),'0'); bb.set(qn('w:color'),'000000')
        tcB.append(bb); tcPr.append(tcB)

def CT(hds, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hds)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c = t.rows[0].cells[i]; c.text = ''; r = c.paragraphs[0].add_run(h)
        srun(r,'宋体','Times New Roman',Pt(10.5),True); spf(c.paragraphs[0],AL_C,Pt(3),Pt(3),sgl=True)
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; r = c.paragraphs[0].add_run(str(v))
            srun(r,'宋体','Times New Roman',Pt(10.5)); spf(c.paragraphs[0],AL_C,Pt(3),Pt(3),sgl=True)
    three_line(t); doc.add_paragraph()
    return t

def add_toc():
    p = doc.add_paragraph()
    r1 = p.add_run(); fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'),'begin'); fc1.set(qn('w:dirty'),'true'); r1._r.append(fc1)
    r2 = p.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text = r' TOC \o "1-3" \h \z \u '; r2._r.append(it)
    r3 = p.add_run(); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'separate'); r3._r.append(fc2)
    r4 = p.add_run('（在Word中右键此处→更新域→更新整个目录）'); srun(r4,'宋体','Times New Roman',Pt(12))
    r5 = p.add_run(); fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'),'end'); r5._r.append(fc3)

for sec in doc.sections:
    sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(3.17); sec.right_margin=Cm(3.17)

ns = doc.styles['Normal']; ns.font.name = 'Times New Roman'; ns.font.size = Pt(12); ns.font.color.rgb = RGBColor(0,0,0)
nrPr = ns.element.get_or_add_rPr(); nrF = nrPr.find(qn('w:rFonts'))
if nrF is None: nrF = OxmlElement('w:rFonts'); nrPr.append(nrF)
nrF.set(qn('w:eastAsia'),'宋体')

for lvl,sz,al,sb,sa in [(1,16,AL_C,40,20),(2,14,AL_L,24,6),(3,13,AL_L,12,6)]:
    hs = doc.styles[f'Heading {lvl}']; hs.font.name='黑体'; hs.font.size=Pt(sz); hs.font.bold=False
    hs.font.color.rgb=RGBColor(0,0,0); hs.font.italic=False; hs.font.underline=False
    hrPr = hs.element.get_or_add_rPr(); hrF = hrPr.find(qn('w:rFonts'))
    if hrF is None: hrF = OxmlElement('w:rFonts'); hrPr.append(hrF)
    hrF.set(qn('w:eastAsia'),'黑体'); hrF.set(qn('w:ascii'),'黑体'); hrF.set(qn('w:hAnsi'),'黑体')
    for attr in ['w:asciiTheme','w:hAnsiTheme','w:eastAsiaTheme']:
        if hrF.get(qn(attr)): del hrF.attrib[qn(attr)]
    hpf = hs.paragraph_format; hpf.alignment=al; hpf.space_before=Pt(sb); hpf.space_after=Pt(sa)
    hpf.line_spacing_rule=WD_LINE_SPACING.EXACTLY; hpf.line_spacing=Pt(20); hpf.keep_with_next=True
    if lvl==1: hpf.page_break_before=True
    hsPr = hs.element.find(qn('w:pPr'))
    if hsPr:
        pBdr = hsPr.find(qn('w:pBdr'))
        if pBdr is not None: hsPr.remove(pBdr)

# ======================== 摘要 ========================
p0 = doc.add_paragraph(); spf(p0, sb=Pt(0), sa=Pt(0))
stitle(f'摘{SP}{SP}要')
BR('近些年来，随着冰雪运动的推广和普及，国内滑雪场数量不断增加，但很多中小型滑雪场在管理方面还比较传统，存在效率低、信息不通畅等问题[7][8]。本文设计并实现了一套基于Spring Boot的飞跃滑雪场管理系统。')
BR('系统采用前后端分离的B/S架构。后端使用Spring Boot框架[6]，数据持久化采用MyBatis操作MySQL数据库[14]，引入Redis做数据缓存[3]，身份认证使用JWT实现[15]；前端使用Vue框架[10]搭配Element Plus组件库[12]进行页面开发。')
B('系统按照角色划分为管理员、教练、工作人员和普通用户四个部分。管理员负责系统整体的数据管理，包括用户管理、教练管理、课程管理、场地管理、商品管理和数据统计等功能。教练可以查看自己负责的课程和学员预约信息，并进行确认操作。工作人员主要负责商品租赁的日常办理、订单处理和库存管理等前台业务。普通用户可以在线浏览课程和教练信息，进行商品租赁、教练预约、课程报名和场地预约等操作。经测试，系统各项功能运行正常，能够满足中小型滑雪场的日常运营管理需求。')
pk = doc.add_paragraph()
rk1 = pk.add_run('关键词：'); srun(rk1,'黑体','黑体',Pt(12),True)
rk2 = pk.add_run('Spring Boot；滑雪场管理；前后端分离；Vue；MySQL'); srun(rk2); spf(pk,AL_J,Pt(0),Pt(0),Pt(20))
doc.add_page_break()

# ======================== ABSTRACT ========================
stitle('ABSTRACT', 'Times New Roman', 'Times New Roman', Pt(16))
pa1 = doc.add_paragraph()
ra1 = pa1.add_run('In recent years, with the promotion of ice and snow sports, the number of ski resorts in China has been increasing. However, many small and medium-sized ski resorts still rely on traditional management methods. This paper designs and implements a ski resort management system based on the Spring Boot framework.')
srun(ra1,'Times New Roman','Times New Roman',Pt(12)); spf(pa1,AL_J,Pt(0),Pt(0),Pt(20),Pt(24))
pa2 = doc.add_paragraph()
ra2 = pa2.add_run('The system adopts a front-end and back-end separated B/S architecture. The back-end uses Spring Boot, MyBatis with MySQL, Redis for caching, and JWT for authentication. The front-end is built with Vue and Element Plus.')
srun(ra2,'Times New Roman','Times New Roman',Pt(12)); spf(pa2,AL_J,Pt(0),Pt(0),Pt(20),Pt(24))
pa3 = doc.add_paragraph()
ra3 = pa3.add_run('The system is divided into four roles: administrator, coach, staff, and ordinary user. The administrator manages all business data. The coach views courses and manages student appointments. The staff handles daily rental operations and order processing. Ordinary users can browse information and perform rental, booking, and registration operations. After testing, all functions run normally and meet the daily management needs of small and medium-sized ski resorts.')
srun(ra3,'Times New Roman','Times New Roman',Pt(12)); spf(pa3,AL_J,Pt(0),Pt(0),Pt(20),Pt(24))
pkw = doc.add_paragraph()
rkw1 = pkw.add_run('Key words: '); srun(rkw1,'Times New Roman','Times New Roman',Pt(12),True)
rkw2 = pkw.add_run('Spring Boot; Ski Resort Management; Front-end and Back-end Separation; Vue; MySQL')
srun(rkw2,'Times New Roman','Times New Roman',Pt(12)); spf(pkw,AL_J,Pt(0),Pt(0),Pt(20))
doc.add_page_break()

# ======================== 目录 ========================
stitle(f'目{SP}{SP}录')
add_toc()

# ======================== 第1章 绪论 ========================
doc.add_heading(f'第1章{SP}绪论', level=1)
doc.add_heading(f'1.1{SP}研究背景与意义', level=2)
BR('冰雪运动在我国有比较长的历史，尤其在东北和华北地区，滑雪一直是冬季比较受欢迎的运动项目。2015年北京申办冬奥会成功以后，国家提出了"带动三亿人参与冰雪运动"的目标[7]，各地的滑雪场建设明显加快。根据行业数据，2023年全国滑雪场已超过700家[8]。')
B('不过从实际情况来看，很多中小型滑雪场在管理上还比较传统。有的雪场租赁雪具还在用纸质登记本，有的教练预约靠电话或微信沟通，收入数据统计也不方便。随着客流量增大，这些做法就不太跟得上了。')
B('开发一套针对滑雪场业务的管理系统，可以把各项业务流程规范起来，让管理者实时掌握运营情况，顾客也能获得更方便的服务体验。这就是本文选择做滑雪场管理系统的出发点。')
B('本文做的这个飞跃滑雪场管理系统，主要就是面向中小型雪场的日常运营。系统里涵盖了雪具租赁、教练预约、课程报名、场地预约等滑雪场最常见的业务场景，管理员、教练、工作人员和普通用户各有各的操作界面和功能权限。开发这个系统也是把大学四年学的Java和前端知识综合用了一遍，希望能对雪场的信息化建设有一些参考。')
doc.add_heading(f'1.2{SP}国内外研究现状', level=2)
B('国外方面，欧洲和北美的大型滑雪度假村较早就开始使用信息化管理工具了，像瑞士、奥地利的一些老牌雪场基本都有完善的票务和租赁管理系统。美国的Vail Resorts集团甚至开发了自己的EpicMix平台，把滑雪数据追踪和社交分享功能都整合在了一起。但这些系统多是商业化的定制产品，价格比较贵，也不太适合国内中小型雪场直接使用。')
BR('国内方面，滑雪场管理系统发展起步较晚，早期多是基于C/S架构开发的桌面应用，功能上以票务管理和财务管理为主，对租赁、预约等业务支持不够。随着Web技术发展，越来越多的开发者开始采用B/S架构来设计滑雪场管理系统，这样用户通过浏览器就可以访问。目前主流的技术方案是后端用Spring Boot[6]，前端用Vue[5]，这也是本文系统采用的技术路线。')
doc.add_heading(f'1.3{SP}主要研究内容', level=2)
B('本文围绕飞跃滑雪场管理系统，主要做了以下工作：')
B('（1）对滑雪场的实际业务进行了调研和分析，通过查阅资料、浏览同类网站和访谈等方式，了解了滑雪场在雪具租赁、教练预约、课程管理、场地管理等方面的业务需求，确定了系统应当具备的功能模块和管理员、教练、工作人员、普通用户四种使用角色。')
B('（2）完成了系统的总体设计工作，包括技术架构选型（前后端分离的B/S架构）、前后端框架选择（Spring Boot + Vue 3）、功能模块划分、数据库表结构设计（9张业务表）和API接口设计等。')
B('（3）按照设计方案编码实现了各项功能。后端使用Spring Boot框架编写了多个Controller和Service层代码，前端使用Vue 3开发了30多个页面组件，涵盖了登录注册、商品管理、教练预约、课程报名、场地管理、数据统计等全部业务功能。')
B('（4）对系统进行了功能测试和性能测试，编写了测试用例验证各模块的功能正确性，并对主要接口的响应时间进行了测量，确认系统运行稳定。')
doc.add_heading(f'1.4{SP}论文组织结构', level=2)
B('本论文共七章，各章内容安排如下：第一章是绪论部分，介绍课题的研究背景和意义，梳理了国内外在滑雪场管理信息化方面的现状。第二章对开发中用到的核心技术做了介绍，包括Spring Boot、MyBatis、MySQL、Redis、JWT、Vue等。第三章是需求分析，详细分析了系统应具备的功能模块和非功能需求，并按角色绘制了用例图。第四章是系统设计，介绍了总体架构、前后端技术方案和数据库表结构设计。第五章是系统实现，展示了各模块的实现过程和核心代码。第六章是系统测试，给出了功能测试和性能测试的结果。第七章对整个毕设工作做了总结，并提出了后续改进方向。')

# ======================== 第2章 相关技术介绍 ========================
doc.add_heading(f'第2章{SP}相关技术介绍', level=1)
doc.add_heading(f'2.1{SP}Spring Boot框架', level=2)
BR('Spring Boot是Pivotal团队提供的用于简化Spring应用开发的框架[9]。在传统的Spring项目中，开发者需要写大量的XML配置文件，比较麻烦。Spring Boot通过"约定优于配置"的理念，大幅减少了配置工作量，开发者只需要引入相应的Starter依赖，框架就会自动完成大部分配置[6]。')
B('本系统使用的是Spring Boot 2.7版本。选这个版本主要是因为比较成熟稳定，网上的教程和文档也比较多，碰到问题好找答案。在项目里，Spring Boot主要负责三件事：一是作为Web服务器，内嵌了Tomcat不用额外部署；二是管理各个组件之间的依赖关系；三是通过application.yml文件集中管理各种配置参数。')
doc.add_heading(f'2.2{SP}MyBatis持久层框架', level=2)
BR('MyBatis是一个比较流行的Java持久层框架，它的定位介于JDBC和Hibernate之间[11]。跟直接写JDBC相比，MyBatis封装了底层的数据库连接和结果集映射，不用写那些重复的代码；跟Hibernate相比，MyBatis给了开发者更大的SQL控制权，可以自己写SQL语句。')
B('在本项目中，每个数据库表对应一个Mapper接口和一个XML文件，接口里定义方法，XML里写具体的SQL。比如用户表对应的就是UserMapper接口和UserMapper.xml。MyBatis支持动态SQL，做条件分页查询的时候，可以根据实际传入的参数灵活拼接查询语句，比如商品列表可以按名称、分类任意组合筛选。')
doc.add_heading(f'2.3{SP}MySQL数据库', level=2)
BR('MySQL是目前全球使用很广泛的开源关系型数据库[14]。它性能不错，运行稳定，而且免费使用，比较适合中小型项目[1]。本系统使用MySQL 8.0版本，这个版本默认采用utf8mb4编码，可以很好地存储中文字符。')
B('系统的数据库名为ski_resort，共创建了9张业务表，使用InnoDB作为存储引擎，支持事务操作。所有表都采用了BIGINT类型的自增主键，并且设置了create_time和update_time字段来记录数据的创建和更新时间。')
doc.add_heading(f'2.4{SP}Redis缓存', level=2)
BR('Redis是一个开源的内存数据存储系统，因为数据都存在内存里，读写速度比传统数据库快很多[3]。在本系统中，Redis主要用于缓存统计概览数据和热门排行榜等不经常变化的数据。通过引入spring-boot-starter-data-redis依赖，可以很方便地在Spring Boot项目中使用Redis。项目中还对Redis的序列化方式做了配置，用Jackson替代默认的JDK序列化，这样存在Redis里的数据看起来更直观。')
doc.add_heading(f'2.5{SP}JWT身份认证', level=2)
BR('JWT全称是JSON Web Token，是一种用于在各方之间安全传递信息的开放标准[15]。在本系统中，JWT用于实现用户的身份认证。工作流程是这样的：用户登录成功后，服务端生成一个包含用户ID、用户名和角色信息的Token返回给前端；前端收到Token后存在本地，之后每次发请求都在HTTP头部带上这个Token；服务端收到请求后验证Token的有效性，确认用户身份。本项目使用的JWT库是jjwt，Token有效期设为7天。')
doc.add_heading(f'2.6{SP}Vue前端框架', level=2)
BR('Vue是一个构建用户界面的渐进式JavaScript框架，在国内Web开发中使用率很高[5][10]。本系统使用Vue 3版本，相比Vue 2，它在性能和TypeScript支持方面都有改进。前端项目使用Vite作为构建工具，启动速度比传统的Webpack快很多。UI组件库用的Element Plus[12]，提供了表格、表单、对话框、分页器等组件，开发效率比较高。状态管理用Pinia，用来在不同组件之间共享用户登录状态等数据。数据可视化部分用ECharts[13]来画统计图表。')
doc.add_heading(f'2.7{SP}其他相关技术', level=2)
B('（1）Lombok：一个Java工具库，通过注解的方式自动生成实体类的getter、setter、构造方法等代码，省去了手工编写的工作量。')
B('（2）Knife4j：基于Swagger的API文档生成工具，可以自动根据Controller中的注解生成在线接口文档，方便前后端开发人员联调。')
B('（3）Spring Security：Spring生态中的安全框架，本项目主要使用它提供的BCryptPasswordEncoder来对用户密码进行加密存储。')
BR('（4）Axios：基于Promise的HTTP客户端库，用于前端向后端发送请求[2]。项目中对Axios进行了二次封装，请求拦截器会自动从Pinia Store中取出Token加到请求头里，响应拦截器统一处理后端返回的错误码，如果Token过期就自动跳转到登录页面。')

# ======================== 第3章 系统需求分析 ========================
doc.add_heading(f'第3章{SP}系统需求分析', level=1)
doc.add_heading(f'3.1{SP}系统概述', level=2)
B('飞跃滑雪场管理系统是一个面向中小型滑雪场运营管理需求的Web应用系统。系统的使用者按角色分为四类：管理员、教练、工作人员和普通用户。管理员拥有最高权限，可以管理所有业务数据；教练可以查看和管理自己的课程及学员预约；工作人员负责日常的雪具租赁办理和订单处理等前台业务；普通用户则可以浏览滑雪场的各项服务信息，进行商品租赁、教练预约、课程报名和场地预约等操作。')
B('在做需求调研的时候，主要参考了几个方面的资料：一是网上一些滑雪场的官方网站，看了它们提供的在线服务有哪些；二是查阅了一些关于体育场馆管理系统的论文，了解了这类系统通常需要具备什么功能；三是向身边去过滑雪场的同学了解了他们作为消费者的需求和体验。综合这些信息，确定了系统的主要功能模块和各角色的操作权限。')
doc.add_heading(f'3.2{SP}功能需求分析', level=2)
B('根据对滑雪场实际业务的调研，本系统按照四种角色组织功能模块，系统总体功能结构如图3.1所示：')
fimg('此处插入系统功能结构图（3级）：一级"飞跃滑雪场管理系统"，二级"管理员/教练/工作人员/普通用户"，三级各角色下的具体功能')
fcap(f'图3.1{SP}系统功能结构图')
doc.add_heading(f'3.2.1{SP}管理员功能', level=3)
B('管理员是系统的最高权限角色，登录后台管理界面后可以进行所有管理操作。用户管理模块可以查看所有注册用户的列表，支持按用户名、手机号等条件搜索，可以新增用户、编辑用户信息、禁用或启用用户账号、为用户分配不同的角色。教练管理模块可以录入教练的基本信息（姓名、等级、专长、课时费等），更新教练的在职/休假状态。课程管理模块可以发布新课程、编辑现有课程信息、更新课程状态。场地管理模块可以录入和编辑场地信息，切换场地的开放/维护状态。商品管理模块可以管理雪具和装备的基本信息及库存数量。数据统计模块展示系统的运营数据概览和各类可视化图表。')
doc.add_heading(f'3.2.2{SP}教练功能', level=3)
B('教练使用自己的账号登录后台，看到的是教练专属的功能菜单。"我的课程"页面展示教练自己负责教授的所有课程列表，可以查看每门课程的详细信息和报名学员名单。"学员预约"页面列出学员对该教练的所有预约记录，教练可以对待确认的预约进行确认或取消操作。在个人中心页面，教练可以查看与自己相关的统计数据（如本月课时数、预约次数等），也可以修改自己的个人资料和登录密码。')
doc.add_heading(f'3.2.3{SP}工作人员功能', level=3)
B('工作人员是滑雪场前台业务的日常操作人员，相当于柜台工作人员。工作人员登录后台后，主要的工作内容包括：为到场的顾客办理雪具和装备的租赁手续，登记租赁信息并生成订单；当顾客归还商品时，确认归还并更新订单状态；查看当前所有的租赁订单列表，了解在租和待归还的商品情况；管理商品库存，及时更新商品的可用数量和状态（如把损坏的雪具标记为维修中）；还可以查看各场地的实时使用情况。')
doc.add_heading(f'3.2.4{SP}普通用户功能', level=3)
B('普通用户是系统的最终使用者，通过前台页面来使用系统。首先需要注册一个账号并登录，登录后可以浏览系统中的各种信息：在商品租赁页面查看可租赁的雪具和装备列表，选择需要的商品提交租赁订单；在教练预约页面浏览各位教练的等级、专长和课时费信息，选择合适的教练提交预约；在滑雪课程页面浏览所有开放报名的课程，选择感兴趣的课程报名参加；在场地信息页面查看各场地的状态和详情，对需要预约的场地进行预约。此外在个人中心可以查看自己的所有订单和预约记录，也可以修改个人资料和登录密码。')
doc.add_heading(f'3.3{SP}非功能需求', level=2)
B('（1）性能需求：系统页面的平均响应时间应控制在2秒以内，对于数据量较大的列表查询应支持分页功能，避免一次性加载太多数据导致页面卡顿。')
B('（2）安全需求：用户密码必须加密存储，不能以明文形式保存在数据库中；接口调用需要携带有效的Token进行身份验证，防止未授权的访问。')
B('（3）兼容性需求：系统前端应兼容主流浏览器，包括Chrome、Firefox、Edge等。由于前端使用了Vue 3和Element Plus，这些框架本身对现代浏览器有比较好的兼容性。需要注意的是IE浏览器不支持，因为Vue 3已经放弃了对IE的支持。')
B('（4）可维护性需求：代码结构要清晰，采用分层架构，各层之间职责明确。后端严格按照Controller-Service-Mapper三层架构来组织代码，每层只做自己应该做的事，不越级调用。前端按功能模块分目录放置页面组件，API调用统一封装在api目录下。')
B('（5）易用性需求：系统界面应简洁明了，操作流程符合用户习惯。列表页面提供搜索和筛选功能，帮助用户快速定位数据。重要操作（如删除）需要弹出确认对话框进行二次确认，防止误操作。操作完成后应有明确的成功或失败提示，使用Element Plus的ElMessage组件在页面顶部展示提示信息。表单提交前应进行前端验证，对必填项、格式等进行检查，减少无效请求。')
doc.add_heading(f'3.4{SP}用例分析', level=2)
B('根据功能需求分析，按照四种用户角色分别进行用例分析。')
doc.add_heading(f'3.4.1{SP}管理员用例', level=3)
tcap(f'表3.1{SP}管理员用例描述')
CT(['用例编号','用例名称','用例描述'],[['UC-A01','登录系统','输入用户名和密码登录后台'],['UC-A02','用户管理','查看、新增、编辑、删除用户'],['UC-A03','教练管理','新增、编辑教练信息'],['UC-A04','课程管理','发布、编辑课程'],['UC-A05','场地管理','新增、编辑场地'],['UC-A06','商品管理','管理商品信息和库存'],['UC-A07','数据统计','查看运营数据和图表']])
fimg('此处插入管理员用例图'); fcap(f'图3.2{SP}管理员用例图')
doc.add_heading(f'3.4.2{SP}教练用例', level=3)
tcap(f'表3.2{SP}教练用例描述')
CT(['用例编号','用例名称','用例描述'],[['UC-C01','登录系统','教练登录后台'],['UC-C02','查看我的课程','查看自己教授的课程'],['UC-C03','管理学员预约','确认或取消预约'],['UC-C04','查看统计','查看个人数据'],['UC-C05','编辑个人信息','修改资料和密码']])
fimg('此处插入教练用例图'); fcap(f'图3.3{SP}教练用例图')
doc.add_heading(f'3.4.3{SP}工作人员用例', level=3)
tcap(f'表3.3{SP}工作人员用例描述')
CT(['用例编号','用例名称','用例描述'],[['UC-S01','登录系统','工作人员登录后台'],['UC-S02','办理租赁','为用户办理商品租赁'],['UC-S03','办理归还','确认归还并更新订单'],['UC-S04','订单管理','查看和处理租赁订单'],['UC-S05','库存管理','查看和更新商品库存']])
fimg('此处插入工作人员用例图'); fcap(f'图3.4{SP}工作人员用例图')
doc.add_heading(f'3.4.4{SP}普通用户用例', level=3)
tcap(f'表3.4{SP}普通用户用例描述')
CT(['用例编号','用例名称','用例描述'],[['UC-U01','注册/登录','注册新账号或登录'],['UC-U02','租赁商品','浏览商品并提交租赁订单'],['UC-U03','预约教练','浏览教练并提交预约'],['UC-U04','报名课程','浏览课程并报名'],['UC-U05','预约场地','查看场地并预约'],['UC-U06','查看订单','查看租赁和预约记录'],['UC-U07','管理个人信息','修改个人资料和密码']])
fimg('此处插入普通用户用例图'); fcap(f'图3.5{SP}普通用户用例图')

# ======================== 第4章 系统设计 ========================
doc.add_heading(f'第4章{SP}系统设计', level=1)
doc.add_heading(f'4.1{SP}系统总体架构设计', level=2)
B('系统架构的选择对整个项目的开发效率和后续维护都有很大影响。经过对比分析，本系统采用了前后端分离的B/S（Browser/Server）架构。B/S架构的优势是用户不需要安装客户端软件，只要有浏览器就可以使用系统，方便部署和更新。系统在逻辑上分为三个层次：表现层（前端Vue应用）、业务逻辑层（后端Spring Boot服务）和数据层（MySQL数据库和Redis缓存）。')
fimg('此处插入系统总体架构图，分三层：表现层(Vue+Element Plus)、业务层(Spring Boot+MyBatis)、数据层(MySQL+Redis)')
fcap(f'图4.1{SP}系统总体架构图')
B('系统的请求处理流程大致是这样的：用户在浏览器中进行操作，前端Vue应用将操作转化为HTTP请求，通过Axios发送到后端的Controller层；Controller接收到请求后调用Service层来处理业务逻辑，Service层根据需要调用Mapper层来操作数据库；处理完成后将结果逐层返回，最终以JSON格式响应给前端，前端再将数据渲染到页面上。')
B('选择前后端分离这种架构方式，主要有两个考虑。一个是开发效率，前端和后端可以各自独立开发，互不影响，只要约定好接口格式就行。另一个是部署灵活性，前端可以单独部署到Nginx上，后端也可以独立部署和扩展。')
doc.add_heading(f'4.2{SP}后端架构设计', level=2)
B('后端项目的包结构遵循标准的MVC分层模式，各包的职责如下表所示：')
tcap(f'表4.1{SP}后端项目包结构说明')
CT(['包名','说明'],[['com.ski.resort.controller','控制层，处理HTTP请求'],['com.ski.resort.service','业务层，处理业务逻辑'],['com.ski.resort.mapper','数据层，操作数据库'],['com.ski.resort.entity','实体层，数据库表对应的Java类'],['com.ski.resort.common','公共类，Result统一返回类等'],['com.ski.resort.config','配置类，安全、缓存、跨域配置'],['com.ski.resort.util','工具类，JWT、Redis工具等']])
B('所有接口采用RESTful风格，返回结果用Result类统一包装，包含code（状态码，200表示成功，其他表示各种错误）、message（提示信息）和data（返回数据）三个字段。接口之间的调用关系遵循严格的分层规范：Controller层不直接操作数据库，而是调用Service层来完成业务逻辑；Service层不处理HTTP请求参数，只关注业务逻辑本身；Mapper层不包含业务逻辑，只负责执行SQL语句。这样的分层方式使得代码职责清晰，后续如果要修改某个模块的功能，只需要改对应层的代码就行，不会影响到其他层。')
doc.add_heading(f'4.3{SP}前端架构设计', level=2)
B('前端项目基于Vue 3 + Vite搭建。前端的路由设计分为两大部分：以"/"开头的路由组用于普通用户的前台页面，以"/admin"开头的路由组用于管理员、教练和工作人员的后台页面。路由守卫中会检查用户是否已登录以及角色权限是否匹配，如果未登录会自动跳转到登录页面，如果角色不匹配则跳转到对应角色的首页。')
B('前台和后台使用了不同的布局组件。前台用的是FrontLayout，包含顶部导航栏和内容区域，导航栏上有首页、商品租赁、教练预约、滑雪课程、场地信息等菜单项。后台用的是MainLayout，采用了左侧边栏加右侧内容区的经典管理后台布局，左侧边栏显示功能菜单，右侧是具体的页面内容。')
doc.add_heading(f'4.4{SP}数据库设计', level=2)
doc.add_heading(f'4.4.1{SP}概念结构设计', level=3)
B('根据需求分析阶段确定的业务功能，对系统涉及的数据进行抽象，识别出主要的实体对象。系统中的核心实体包括：用户（系统的所有使用者，包括管理员、教练、工作人员和普通用户）、商品（统一管理的雪具和装备，既可用于租赁也可用于销售）、教练（滑雪场的教练人员）、课程（教练开设的滑雪培训课程）、场地（滑雪场内的各类场地）。此外还有几个关联实体：租赁订单（用户租赁商品时产生的订单记录）、教练预约（用户预约教练的记录）、场地预订（用户预约场地的记录）和课程报名（用户报名课程的记录）。')
B('这些实体之间存在以下主要关系：一个用户可以有多个租赁订单、多个教练预约、多个课程报名记录，即用户与这些订单/预约记录之间是一对多的关系。一个教练可以教授多门课程，同时也可以被多个用户预约，教练与课程之间是一对多关系，教练与用户之间通过教练预约表形成多对多关系。一门课程由一个教练教授，在一个场地上课，课程与教练是多对一关系，课程与场地也是多对一关系。一个租赁订单关联一件商品。')
fimg('此处插入E-R图，包含用户、商品、教练、课程、场地、租赁订单、教练预约、课程报名、场地预订等实体及关系')
fcap(f'图4.2{SP}系统E-R关系简图')
doc.add_heading(f'4.4.2{SP}逻辑结构设计', level=3)
B('将概念模型转化为关系模型，系统共设计了9张数据库表。在设计表结构的时候，遵循了几个原则：一是每张表都设置了自增的BIGINT类型主键，方便进行关联查询；二是需要记录时间的字段统一使用DATETIME类型，并在create_time字段设置默认值为CURRENT_TIMESTAMP；三是涉及到金额的字段使用DECIMAL(10,2)类型，避免浮点数精度问题；四是用TINYINT类型来表示状态字段，通过不同的数值表示不同的状态。下面列出主要表的结构：')
H4 = ['字段名','数据类型','可否为空','说明']
tcap(f'表4.2{SP}用户表（t_user）结构')
CT(H4,[['id','BIGINT','否（主键）','用户ID，自增'],['username','VARCHAR(50)','否','用户名，唯一'],['password','VARCHAR(255)','否','密码（BCrypt加密）'],['real_name','VARCHAR(100)','是','真实姓名'],['phone','VARCHAR(20)','是','手机号'],['role','VARCHAR(20)','否','角色：admin/user/coach/staff'],['status','TINYINT','是','状态：0禁用，1正常'],['create_time','DATETIME','是','创建时间']])
tcap(f'表4.3{SP}商品表（t_product）结构')
CT(H4,[['id','BIGINT','否（主键）','商品ID，自增'],['product_code','VARCHAR(50)','否','商品编号，唯一'],['product_name','VARCHAR(100)','否','商品名称'],['category','VARCHAR(50)','否','分类：雪具/装备/配件'],['brand','VARCHAR(100)','是','品牌'],['rental_price','DECIMAL(10,2)','是','租赁价格（元/小时）'],['sale_price','DECIMAL(10,2)','是','销售价格'],['stock_quantity','INT','是','库存数量'],['available_quantity','INT','是','可用数量'],['status','TINYINT','是','状态：0下架，1上架']])
tcap(f'表4.4{SP}教练表（t_coach）结构')
CT(H4,[['id','BIGINT','否（主键）','教练ID，自增'],['user_id','BIGINT','是','关联用户ID'],['coach_name','VARCHAR(50)','否','教练姓名'],['coach_level','VARCHAR(20)','否','等级：初级/中级/高级'],['specialty','VARCHAR(200)','是','专长'],['hourly_rate','DECIMAL(10,2)','否','课时费'],['rating','DECIMAL(3,2)','是','评分'],['status','TINYINT','是','状态：0休假，1在职']])
tcap(f'表4.5{SP}课程表（t_course）结构')
CT(H4,[['id','BIGINT','否（主键）','课程ID，自增'],['course_name','VARCHAR(100)','否','课程名称'],['course_type','VARCHAR(50)','否','类型：初级/中级/高级'],['coach_id','BIGINT','否','教练ID'],['max_students','INT','否','最大学员数'],['current_students','INT','是','当前报名人数'],['course_price','DECIMAL(10,2)','否','课程价格'],['venue_id','BIGINT','是','场地ID'],['status','TINYINT','是','状态：0未开始~3已结束']])
tcap(f'表4.6{SP}场地表（t_venue）结构')
CT(H4,[['id','BIGINT','否（主键）','场地ID，自增'],['venue_name','VARCHAR(100)','否','场地名称'],['venue_type','VARCHAR(50)','否','类型：雪道/练习场'],['difficulty_level','VARCHAR(20)','是','难度等级'],['max_capacity','INT','是','最大容纳人数'],['rental_price','DECIMAL(10,2)','是','租赁价格'],['status','TINYINT','是','状态：0维护中，1开放']])
tcap(f'表4.7{SP}租赁订单表（t_rental_order）结构')
CT(H4,[['id','BIGINT','否（主键）','订单ID，自增'],['order_no','VARCHAR(50)','否','订单号，唯一'],['user_id','BIGINT','否','用户ID'],['product_id','BIGINT','否','商品ID'],['rental_start_time','DATETIME','否','租赁开始时间'],['rental_end_time','DATETIME','是','归还时间'],['total_amount','DECIMAL(10,2)','否','总金额'],['order_status','TINYINT','是','状态：0待支付~3已取消']])
B('除上述6张主要表外，系统还包括3张关联业务表。教练预约表（t_coach_booking）记录用户对教练的预约信息，包含用户ID、教练ID、预约日期、时间段、费用金额和预约状态等字段。场地预订表（t_venue_booking）记录用户对场地的预订信息，结构与教练预约表类似。课程报名表（t_course_enrollment）记录学员的课程报名数据，包含用户ID、课程ID、报名时间和报名状态等字段。这三张表都通过外键关联到用户表和对应的业务主表上。')
doc.add_heading(f'4.5{SP}接口设计', level=2)
B('系统后端采用RESTful风格设计API接口。RESTful的核心思想是用HTTP方法来表示操作类型：GET用于查询、POST用于新增、PUT用于修改、DELETE用于删除。每个接口的URL路径对应一个资源，通过不同的HTTP方法对同一资源进行不同的操作。后端共定义了约50个API接口，覆盖了用户、商品、教练、课程、场地、租赁订单和统计等各模块的业务需求。主要接口列表如下：')
tcap(f'表4.8{SP}主要API接口列表')
CT(['模块','请求方式','接口路径','功能说明'],[['用户','POST','/api/v1/users/login','用户登录'],['用户','POST','/api/v1/users/register','用户注册'],['商品','GET','/api/v1/products/list','查询商品列表'],['教练','GET','/coach/list','查询教练列表'],['课程','GET','/api/v1/courses/list','查询课程列表'],['租赁','POST','/rental-order','创建租赁订单'],['租赁','PUT','/rental-order/{id}/return','归还商品'],['统计','GET','/api/v1/dashboard/overview','获取统计概览']])

# ======================== 第5章 系统实现 ========================
doc.add_heading(f'第5章{SP}系统实现', level=1)
doc.add_heading(f'5.1{SP}开发环境', level=2)
B('在正式开始编码之前，需要先搭建好开发环境。后端使用IntelliJ IDEA作为开发工具，通过Maven管理项目依赖。项目创建时使用Spring Initializr初始化，勾选了Web、MyBatis、MySQL Driver、Redis等Starter依赖。前端使用Visual Studio Code编辑器，通过npm管理前端依赖包。具体的开发环境配置如下表所示：')
tcap(f'表5.1{SP}开发环境配置')
CT(['名称','版本/说明'],[['操作系统','Windows 10/11'],['开发工具（后端）','IntelliJ IDEA 2023'],['开发工具（前端）','Visual Studio Code'],['JDK','17'],['Maven','3.9'],['Node.js','18.x'],['MySQL','8.0'],['Redis','5.0+'],['浏览器','Chrome 120']])
doc.add_heading(f'5.2{SP}用户登录模块的实现', level=2)
B('用户登录是系统的入口功能。前端登录页面采用了左右分栏的布局，左侧展示系统名称和特色功能介绍，右侧是登录和注册的表单区域，支持在登录和注册两个Tab页之间切换。登录表单中用到了Element Plus的el-form组件，配合el-input来收集用户名和密码，并使用了表单验证规则来限制输入。')
B('登录的具体流程如下：用户输入用户名和密码后点击登录按钮，前端先对表单进行校验（用户名和密码不能为空），校验通过后通过Axios向后端发送POST请求。后端UserController接收到请求后，调用UserService验证用户名和密码，密码验证使用BCryptPasswordEncoder进行比对。验证通过后生成JWT Token，连同用户信息一起返回给前端。前端收到数据后将Token和用户信息保存到Pinia Store和LocalStorage中，然后根据用户角色进行页面跳转：普通用户跳转到前台首页"/home"，管理员、教练和工作人员跳转到后台首页"/admin/dashboard"。后端登录接口核心代码如下：')
fimg('此处插入用户登录页面截图')
fcap(f'图5.1{SP}用户登录页面')
CB('@PostMapping("/login")\npublic Result<Map<String, Object>> login(@RequestBody LoginDTO loginDTO) {\n    User user = userService.login(\n        loginDTO.getUsername(), loginDTO.getPassword());\n    String token = JwtUtil.generateToken(\n        user.getId(), user.getUsername(), user.getRole());\n    Map<String, Object> data = new HashMap<>();\n    data.put("token", token);\n    data.put("userInfo", user);\n    return Result.success("登录成功", data);\n}')
B('JWT Token生成方法，将用户信息写入Token并设置7天有效期：')
CB('public static String generateToken(Long userId, String username, String role) {\n    Map<String, Object> claims = new HashMap<>();\n    claims.put("userId", userId);\n    claims.put("username", username);\n    claims.put("role", role);\n    return Jwts.builder().setClaims(claims).setSubject(username)\n        .setExpiration(new Date(System.currentTimeMillis() + 7*24*3600*1000))\n        .signWith(getSecretKey(), SignatureAlgorithm.HS256).compact();\n}')
fimg('此处插入用户登录流程图')
fcap(f'图5.2{SP}用户登录流程图')
doc.add_heading(f'5.3{SP}商品管理模块的实现', level=2)
B('商品管理模块是系统的核心业务模块之一，统一管理滑雪场的雪具和装备信息。在本系统的设计中，雪具（如滑雪板、雪鞋、头盔等）和装备（如护具、雪镜等）统一放在一张商品表中管理，通过category字段区分是雪具、装备还是配件。这样设计的好处是数据结构更简洁，前后端的代码也不用分开写两套。')
B('管理员和工作人员可以在后台的商品管理页面对商品进行增删改查操作。页面顶部有搜索区域，可以按商品名称和分类进行筛选查询。列表使用了Element Plus的el-table组件来展示数据，支持分页功能，每页显示10条记录。点击"新增"按钮会弹出对话框让用户填写商品信息，包括商品编号、名称、分类、品牌、租赁价格、销售价格、库存数量等字段。编辑操作也是弹出同样的对话框，只是自动填入了已有的数据。删除操作采用逻辑删除，通过设置is_deleted字段为1来标记。后端分页查询接口代码如下：')
fimg('此处插入商品管理页面截图')
fcap(f'图5.3{SP}商品管理页面')
CB('@GetMapping("/page/condition")\npublic Result<PageResult<Product>> getByCondition(\n        @RequestParam(defaultValue = "1") Integer page,\n        @RequestParam(defaultValue = "10") Integer pageSize,\n        @RequestParam(required = false) String name,\n        @RequestParam(required = false) String category) {\n    return Result.success(\n        productService.pageByCondition(page, pageSize, name, category));\n}')
B('用户端以卡片形式展示可租赁的商品，每张卡片显示商品图片、名称、品牌、租赁价格等信息。用户选择商品后指定租赁时间和数量，系统自动计算租赁费用并生成订单。订单创建时系统自动扣减商品的可用数量，归还时再加回来。租赁订单有待支付、使用中、已归还、已取消四种状态，工作人员可以在后台管理所有的租赁订单，进行归还确认和取消操作。')
fimg('此处插入商品租赁业务流程图')
fcap(f'图5.4{SP}商品租赁业务流程图')
doc.add_heading(f'5.4{SP}教练预约模块的实现', level=2)
B('教练预约涉及管理员、教练和用户三个角色。管理员在后台维护教练的基本信息，包括姓名、等级、专长、课时费等。教练列表页面支持按姓名、等级和状态进行筛选查询。')
B('用户端的教练预约页面展示了所有在职教练的信息。用户选择教练后，需要指定预约日期、开始时间和结束时间，系统根据教练的课时费和预约时长自动计算费用。提交预约后状态为"待确认"。教练登录后可在"学员预约"页面查看并确认或取消预约。后端通过coach_id进行数据隔离，确保每个教练只能看到自己的预约数据：')
fimg('此处插入教练预约页面截图')
fcap(f'图5.5{SP}教练预约页面')
CB('@GetMapping("/my-bookings")\npublic Result<List<CoachBooking>> getMyBookings() {\n    Long userId = SecurityUtils.getCurrentUserId();\n    Coach coach = coachService.getByUserId(userId);\n    return Result.success(\n        coachBookingService.getByCoachId(coach.getId()));\n}')
doc.add_heading(f'5.5{SP}课程与场地模块的实现', level=2)
B('课程管理模块包含课程的发布管理和学员的报名功能。管理员在后台可以发布新课程，需要填写课程名称、类型（初级入门/中级进阶/高级提升）、授课教练、开课日期、上课时间、场地、最大学员数和价格等信息。添加课程的时候，前端会同时发两个请求去获取教练列表和场地列表作为下拉选项，管理员可以直接选择而不用手动输入ID。用户端的课程页面展示所有"报名中"的课程，用户点击报名后系统会检查报名人数是否已满，未满则允许报名并将当前报名人数加1。')
B('场地管理比较简单。管理员可以在后台录入和编辑场地信息，包括场地名称、类型、难度等级和容纳人数。系统预置了初级雪道、中级雪道、高级雪道和自由式公园等场地数据。雪道类型的场地包含在门票里不用另外收费，自由式公园则需要额外收费。用户可以查看各场地的状态信息，对需要付费的场地进行在线预约。')
doc.add_heading(f'5.6{SP}数据统计模块的实现', level=2)
B('数据统计模块是管理员了解雪场整体运营状况的重要工具，也是做经营决策的数据支撑。这个模块包含两个主要页面：Dashboard看板页面和数据可视化分析页面。管理员登录后默认进入的就是Dashboard看板页面。')
B('Dashboard看板页面展示了系统的核心运营数据概览，包括总用户数、今日新增用户数、总订单数、今日订单数、总收入金额、今日收入等关键指标。这些数据通过后端的DashboardController接口获取，接口内部对多张业务表进行聚合查询和统计。为了避免每次访问都执行复杂的数据库查询，统计结果会缓存到Redis中，设置了5分钟的过期时间。')
B('数据可视化分析页面使用了ECharts图表库来展示各种统计图表。主要包括：收入趋势折线图，展示近30天的每日收入变化趋势；订单类型分布饼图，展示租赁订单、教练预约和课程报名各占的比例；热门商品排行柱状图，展示租赁次数最多的商品。前端使用Vue的onMounted生命周期钩子在页面加载时请求数据，然后调用ECharts的setOption方法渲染图表。')
fimg('此处插入数据统计页面截图')
fcap(f'图5.6{SP}数据统计页面')
doc.add_heading(f'5.7{SP}系统安全实现', level=2)
B('系统的安全控制主要体现在以下几个方面：')
B('密码加密存储：用户密码在注册时通过BCryptPasswordEncoder进行加密后才保存到数据库中，BCrypt是一种单向加密算法，即使数据库被泄露也无法直接获取用户的明文密码。')
B('JWT Token认证：系统配置了JwtAuthenticationFilter过滤器，对每个HTTP请求进行Token的解析和验证。前端在每次请求时会在Authorization请求头中携带Bearer Token，后端过滤器从中提取Token并验证，验证通过后将用户信息设置到Security上下文中。')
B('前端路由守卫：前端在路由配置中使用了beforeEach全局守卫，检查目标路由是否需要登录权限以及当前用户的角色是否有权限访问。如果用户未登录会被重定向到登录页面；如果角色不匹配会被重定向到自己角色对应的首页。')
B('跨域配置：由于前后端分别运行在不同的端口（前端5173，后端8080），需要处理跨域请求的问题。后端通过CorsConfig配置类来允许跨域访问。')
B('前端路由守卫核心逻辑：')
CB('router.beforeEach((to, from, next) => {\n  const userStore = useUserStore()\n  if (to.meta.requiresAuth && !userStore.token) {\n    next(\'/login\')\n  } else if (to.meta.roles &&\n      !to.meta.roles.includes(userStore.userInfo?.role)) {\n    next(userStore.userInfo?.role === \'user\' ? \'/home\' : \'/admin/dashboard\')\n  } else { next() }\n})')

# ======================== 第6章 系统测试 ========================
doc.add_heading(f'第6章{SP}系统测试', level=1)
doc.add_heading(f'6.1{SP}测试环境', level=2)
B('系统测试环境如下：')
tcap(f'表6.1{SP}测试环境')
CT(['项目','配置'],[['操作系统','Windows 11'],['处理器','Intel Core i5'],['内存','16GB'],['浏览器','Chrome 120'],['后端环境','JDK 17 + Tomcat(内嵌)'],['数据库','MySQL 8.0'],['缓存','Redis 7.0']])
B('测试账号：管理员admin，教练coach1-coach5，工作人员staff1，普通用户testuser（密码均为123456）。测试数据方面，系统中预置了20条商品数据（包含单板、双板、雪鞋、头盔、护具等类型）、5名教练数据、8门课程数据和6个场地数据。测试时还额外注册了若干普通用户账号来模拟多用户场景。')
doc.add_heading(f'6.2{SP}功能测试', level=2)
B('功能测试采用手工测试的方式，按各模块的业务流程，使用不同角色的账号分别登录系统，模拟用户实际操作来验证各项功能是否符合需求规格。测试过程中对每个功能点编写了测试用例，记录测试步骤、预期结果和实际结果。')
doc.add_heading(f'6.2.1{SP}用户模块测试', level=3)
tcap(f'表6.2{SP}用户模块测试用例')
CT(['编号','测试项','预期结果','实际结果','是否通过'],[['T-U01','管理员登录','跳转后台首页','跳转/admin/dashboard','通过'],['T-U02','错误密码登录','提示密码错误','提示错误','通过'],['T-U03','普通用户登录','跳转前台首页','跳转/home','通过'],['T-U04','用户注册','注册成功','提示成功','通过'],['T-U05','重复用户名注册','提示已存在','提示失败','通过'],['T-U06','工作人员登录','跳转后台首页','跳转后台','通过']])
doc.add_heading(f'6.2.2{SP}商品租赁模块测试', level=3)
tcap(f'表6.3{SP}商品租赁模块测试用例')
CT(['编号','测试项','预期结果','实际结果','是否通过'],[['T-P01','查看商品列表','显示分页列表','正确显示','通过'],['T-P02','新增商品','新增成功','成功','通过'],['T-P03','条件搜索','筛选结果正确','正确','通过'],['T-P04','用户租赁商品','创建订单并扣库存','正确','通过'],['T-P05','归还商品','状态变更库存恢复','正确','通过']])
doc.add_heading(f'6.2.3{SP}教练预约与课程模块测试', level=3)
tcap(f'表6.4{SP}教练预约与课程模块测试用例')
CT(['编号','测试项','预期结果','实际结果','是否通过'],[['T-C01','查看教练列表','显示在职教练','正确','通过'],['T-C02','提交预约','创建成功','成功','通过'],['T-C03','教练确认预约','状态变为已确认','正确','通过'],['T-CR01','查看课程列表','显示可报名课程','正确','通过'],['T-CR02','报名课程','报名成功人数+1','成功','通过'],['T-V01','预约场地','预约创建成功','成功','通过']])
doc.add_heading(f'6.3{SP}性能测试', level=2)
B('除了功能测试之外，还对系统的主要接口进行了简单的性能测试。测试方法是在单用户场景下，使用Chrome浏览器的开发者工具中的Network面板来记录每个接口从发起请求到收到响应的耗时，每个接口测试5次取平均值。测试结果如下表所示：')
tcap(f'表6.5{SP}接口响应时间测试')
CT(['接口','说明','平均响应时间','是否达标(<2s)'],[['POST /users/login','用户登录','156ms','是'],['GET /products/page','商品列表','72ms','是'],['GET /coach/list','教练列表','65ms','是'],['GET /courses/list','课程列表','48ms','是'],['GET /dashboard/overview','统计概览','210ms','是']])
B('从测试结果来看，所有接口的响应时间均在200ms左右，远低于2秒的标准，系统性能表现良好。统计概览接口由于需要做多表聚合查询，响应时间相对较长，但通过Redis缓存可以进一步优化。')
doc.add_heading(f'6.4{SP}测试总结', level=2)
B('经过上面这些测试，系统的主要功能基本都跑通了，像用户登录、商品租赁、教练预约、课程报名、场地预约、数据统计这些核心流程都没有问题。接口的响应速度也还可以，都在200ms左右，满足性能要求。')
B('测试的时候也发现了一些小问题，举几个例子：一个是商品管理页面的状态字段，后端返回的是数字（0表示下架、1表示上架），前端没有做转换直接显示数字，用户看不懂是什么意思。后来在前端加了一个状态映射，把0显示为"已下架"、把1显示为"上架中"就解决了。另一个是租赁订单的金额显示，有时候会显示很多位小数，比如显示"25.000000001"这种，后来在前端用toFixed(2)处理了一下。还有一个问题是教练预约的时间段校验不够严格，理论上不应该允许预约过去的时间，但系统没有做这个检查，后来在前端加了一个日期选择器的disabledDate属性来禁止选择今天之前的日期。这些都是比较小的问题，不影响核心功能，发现后很快就修好了。')
B('还有一个在测试中发现的问题是关于并发的。当两个用户同时租赁同一件商品的时候，可能会出现库存变成负数的情况。这个问题的根源是在扣减库存的时候没有做并发控制。理论上应该在数据库层面用乐观锁来处理，不过因为本系统定位为中小型雪场使用，并发量不会太大，所以这个问题在当前阶段影响不大，但如果以后用户量增加的话，这个地方是需要优化的。')
B('总的来说，系统基本达到了设计的功能目标，主要的业务流程都能正常走通。')

# ======================== 第7章 总结与展望 ========================
doc.add_heading(f'第7章{SP}总结与展望', level=1)
doc.add_heading(f'7.1{SP}总结', level=2)
B('本文设计并实现了一个基于Spring Boot的飞跃滑雪场管理系统。后端用Spring Boot 2.7加MyBatis来操作MySQL数据库，加了Redis做数据缓存，JWT做身份认证；前端用Vue 3和Element Plus来做页面，ECharts画数据图表。整个项目采用前后端分离的方式来开发。')
B('系统按角色分为管理员、教练、工作人员和普通用户四种，实现了用户管理、商品管理、教练预约、课程管理、场地管理和数据统计等功能模块。数据库设计了9张表，后端编写了多个Controller接口，前端做了30多个页面组件。测试下来功能基本都没问题，可以满足中小型雪场的日常管理需要。')
B('做这个项目的过程中，我觉得收获还是比较大的。以前学Spring Boot和Vue的时候都是单独练习一些小例子，这次把它们放在一个完整的项目里，才真正体会到前后端联调、接口对接这些环节的复杂性。比如前端页面逻辑和后端数据格式不一致的时候要怎么处理，分页查询的参数怎么传递，跨域问题怎么解决等等，这些在做小练习的时候是碰不到的。当然了，系统肯定还有不少可以改进的地方。')
doc.add_heading(f'7.2{SP}展望', level=2)
B('系统当前的功能基本都实现了，但因为时间和个人能力有限，有些地方以后还可以继续改进：')
B('（1）移动端适配：目前系统虽然使用了响应式设计，但在手机上的使用体验还不够理想，后续可以考虑开发专门的微信小程序或移动端App，让用户随时随地都能查看和管理。')
B('（2）支付功能集成：当前系统的支付功能只做了状态模拟，没有对接真实的支付接口。后续可以集成微信支付或支付宝等第三方支付平台，实现线上支付。')
B('（3）消息通知功能：可以增加短信通知或站内信功能，在预约确认、订单状态变更等时机自动通知用户，提升用户体验。')
B('（4）权限细化：目前的权限控制比较粗，只做到了角色级别，后续可以引入更细粒度的权限管理，比如按菜单和按钮级别进行控制。')

# ======================== 参考文献 ========================
doc.add_heading('参考文献', level=1)
refs = [
    '[1] 王珊, 萨师煊. 数据库系统概论(第5版)[M]. 北京: 高等教育出版社, 2014.',
    '[2] 谢希仁. 计算机网络(第7版)[M]. 北京: 电子工业出版社, 2017.',
    '[3] 黄健宏. Redis设计与实现[M]. 北京: 机械工业出版社, 2014.',
    '[4] Craig Walls. Spring实战(第5版)[M]. 张卫滨, 译. 北京: 人民邮电出版社, 2020.',
    '[5] 霍春阳. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022.',
    '[6] 杨开振. 深入浅出Spring Boot 2.x[M]. 北京: 人民邮电出版社, 2019.',
    '[7] 国务院办公厅. 关于促进全民健身和体育消费推动体育产业高质量发展的意见[Z]. 国办发〔2019〕43号, 2019.',
    '[8] 国家体育总局. 冰雪运动发展规划(2016-2025年)[Z]. 2016.',
    '[9] Spring Boot Reference Documentation[EB/OL]. https://spring.io/projects/spring-boot, 2024.',
    '[10] Vue.js Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
    '[11] MyBatis Official Documentation[EB/OL]. https://mybatis.org/mybatis-3/, 2024.',
    '[12] Element Plus Official Documentation[EB/OL]. https://element-plus.org/, 2024.',
    '[13] Apache ECharts Documentation[EB/OL]. https://echarts.apache.org/, 2024.',
    '[14] MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/, 2024.',
    '[15] JSON Web Token Introduction[EB/OL]. https://jwt.io/introduction, 2024.',
]
for ref in refs:
    p = doc.add_paragraph(); r = p.add_run(ref)
    srun(r, '黑体', 'Times New Roman', Pt(10.5)); spf(p, AL_J, Pt(3), Pt(0), Pt(16))

# ======================== 致谢 ========================
doc.add_heading(f'致{SP}{SP}谢', level=1)
B('不知不觉大学四年就这么过去了，论文写到这里也快结束了，心里还是有不少想要感谢的人。')
B('首先要感谢我的指导老师，从选题开始就一直在帮我理清思路，开发过程中遇到技术难题也会耐心地给我指导和建议。论文修改的时候老师更是一版一版地仔细审阅，从内容到格式都帮我把关，提出了很多宝贵的修改意见。没有老师的悉心指导，这篇论文是无法顺利完成的，真的很感谢老师这一年来的辛苦付出。')
B('感谢大学四年来所有教过我的老师们，正是你们在课堂上的教导，让我从一个对编程一窍不通的小白，逐渐掌握了Java编程、数据库设计、Web开发等专业知识和技能。特别是数据库原理、软件工程和Web开发这几门课的老师，你们讲的知识在这次毕设中都用上了。')
B('还有我的同学和朋友们，写代码卡住的时候经常找他们讨论，有时候自己调了半天的Bug，别人看一眼就发现问题所在了。大家一起熬夜赶项目的日子虽然辛苦，但现在回想起来也挺怀念的。')
B('最后要感谢我的家人，是你们一直以来在背后的支持和鼓励，让我能够安心地完成学业。谢谢你们。')

# ======================== 保存 ========================
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '论文_排版版v3.docx')
doc.save(out)
print(f'论文docx已生成: {out}')
