# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
import re

doc = Document('5.26.2.docx')

print("=" * 70)
print("要求1: 3.4要和3.2的内容结合起来")
print("=" * 70)
has_34 = False
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '3.4' in p.text:
        has_34 = True
        print(f"  [X] 3.4 仍然存在: [{i}] {p.text.strip()}")
if not has_34:
    print(f"  [OK] 3.4 已不存在（已合并到3.2）")

# Check 3.2 has activity diagrams
print("\n  3.2 子节:")
in_32 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '3.2' in text and style.startswith('Heading 2'):
        in_32 = True
    if in_32 and style == 'Heading 2' and '3.3' in text:
        break
    if in_32 and style.startswith('Heading 3'):
        print(f"    {text}")
    if in_32 and ('图3' in text) and not style.startswith('Heading'):
        print(f"    -> 图: {text}")

print("\n" + "=" * 70)
print("要求2: 第5章按角色分类(管理员-用户-工作人员-教练)")
print("=" * 70)
in_ch5 = False
ch5_sections = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and ('第5章' in text):
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    elif in_ch5 and style.startswith('Heading'):
        ch5_sections.append(text)
        print(f"  {text}")

# Check order: should be 管理员-用户-工作人员-教练
section_order = []
for s in ch5_sections:
    if '管理员' in s and '工作' not in s:
        section_order.append('管理员')
    elif '用户' in s:
        section_order.append('用户')
    elif '工作人员' in s:
        section_order.append('工作人员')
    elif '教练' in s:
        section_order.append('教练')

expected_order = ['管理员', '用户', '工作人员', '教练']
actual_main = [x for x in section_order if x in expected_order]
if actual_main == expected_order:
    print(f"\n  [OK] 角色顺序正确: {actual_main}")
else:
    print(f"\n  [X] 角色顺序不对!")
    print(f"  期望: {expected_order}")
    print(f"  实际: {actual_main}")

# Check 5.1 开发环境
first_section_text = ""
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '5.1' in text:
        first_section_text = text
        break
print(f"\n  5.1节标题: {first_section_text}")
if '开发环境' in first_section_text or '开发' in first_section_text:
    print(f"  [X] 5.1 仍然是单独的开发环境节")
else:
    print(f"  [OK] 5.1 不是单独的开发环境节")

print("\n" + "=" * 70)
print("要求3: 文献按顺序")
print("=" * 70)
refs_found = set()
for p in doc.paragraphs:
    citations = re.findall(r'\[(\d+)\]', p.text)
    for c in citations:
        refs_found.add(int(c))
refs_sorted = sorted(refs_found)
print(f"  引用编号: {refs_sorted}")
missing = [n for n in range(1, max(refs_sorted) + 1) if n not in refs_found] if refs_sorted else []
if missing:
    print(f"  [X] 缺失: {missing}")
else:
    print(f"  [OK] 编号连续")

print("\n" + "=" * 70)
print("要求4: 功能结构图/流程图字体大小 (需人工查看图片)")
print("=" * 70)
print("  前一个会话(MCP-3)已将功能结构图和流程图字号改为24pt")
print("  老师要求流程图9-10.5pt，用户后续要求改为24pt")
print("  [需人工确认] 请打开docx查看图片字体是否合适")

print("\n" + "=" * 70)
print("要求5: 4.5(数据库设计)应在4.4(业务流程设计)前边")
print("=" * 70)
in_ch4 = False
ch4_order = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第4章' in text:
        in_ch4 = True
    elif in_ch4 and style == 'Heading 1':
        break
    elif in_ch4 and style == 'Heading 2':
        ch4_order.append(text)
        print(f"  {text}")

# Check: 数据库 should come before 业务流程
db_idx = -1
biz_idx = -1
for idx, s in enumerate(ch4_order):
    if '数据库' in s:
        db_idx = idx
    if '业务' in s:
        biz_idx = idx
if db_idx >= 0 and biz_idx >= 0:
    if db_idx < biz_idx:
        print(f"\n  [OK] 数据库设计(位置{db_idx}) 在 业务流程(位置{biz_idx}) 前")
    else:
        print(f"\n  [X] 数据库设计(位置{db_idx}) 在 业务流程(位置{biz_idx}) 后!")
        print(f"  需要交换 4.4 和 4.5 的顺序")

print("\n" + "=" * 70)
print("要求6: 第5章字数检查")
print("=" * 70)
print("  (此项无法精确判断是否'太多'，需人工审阅)")

print("\n" + "=" * 70)
print("要求7: 三线表间距检查")
print("=" * 70)
print("  (三线表格式已由MCP-3统一修改为1.5pt/1pt)")
print("  [需人工确认] 表间空格大小需打开docx查看")

print("\n" + "=" * 70)
print("要求8: 6.4 口语化检查")
print("=" * 70)
in_64 = False
oral_issues = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
    elif in_64 and style.startswith('Heading') and '6.4' not in text:
        break
    elif in_64 and text:
        oral_words = ['还有一个问题', '理论上', '都能', '其实', '挺', '可能会']
        found = [w for w in oral_words if w in text]
        if found:
            oral_issues.append((i, found, text[:120]))

if oral_issues:
    print(f"  [X] 发现 {len(oral_issues)} 处口语化表述:")
    for idx, words, txt in oral_issues:
        print(f"    [{idx}] 词: {words}")
        print(f"         {txt}")
else:
    print(f"  [OK] 未发现明显口语化表述")

# Also check for "都能" and the last sentence issue
in_64 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
    elif in_64 and style.startswith('Heading'):
        break
    elif in_64 and text:
        if '都能' in text:
            pos = text.find('都能')
            context = text[max(0, pos-20):pos+20]
            print(f"\n  [X] 发现'都能': ...{context}...")

print("\n" + "=" * 70)
print("要求9: 7.1 最后一段 - '笔者' 和 '小型练习'")
print("=" * 70)
in_71 = False
paras_71 = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if '7.1' in text and style.startswith('Heading'):
        in_71 = True
    elif in_71 and style.startswith('Heading'):
        break
    elif in_71 and text:
        paras_71.append((i, text))

if paras_71:
    # Check first sentence of last paragraph
    last_idx, last_text = paras_71[-1]
    print(f"  7.1最后一段 [{last_idx}]:")
    print(f"  {last_text[:200]}")
    
    if '笔者' in last_text:
        print(f"\n  [X] 仍包含'笔者'")
    else:
        print(f"\n  [OK] 不包含'笔者'")
    
    if '小型练习' in last_text:
        print(f"  [X] 仍包含'小型练习'，应改为'课程中'")
    else:
        print(f"  [OK] 不包含'小型练习'")
    
    if '在课程中' in last_text or '在课程' in last_text:
        print(f"  [OK] 已改为'课程'相关表述")

    # Also check first paragraph for 笔者
    for idx, txt in paras_71:
        if '笔者' in txt:
            print(f"\n  [X] [{idx}] 段中包含'笔者': ...{txt[max(0, txt.find('笔者')-15):txt.find('笔者')+15]}...")

print("\n" + "=" * 70)
print("总结")
print("=" * 70)
