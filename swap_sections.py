# -*- coding: utf-8 -*-
"""
整段交换文档章节：
1. 第5章: 5.3(教练) ↔ 5.4(工作人员) 
2. 第4章: 4.4(业务流程) ↔ 4.5(数据库设计)
3. 修复6.4口语化
"""
import sys, io, copy, shutil, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

INPUT = '5.26.2.docx'
OUTPUT = '5.26.2.docx'

shutil.copy2(INPUT, '5.26.2_before_swap.docx')
print("已备份 -> 5.26.2_before_swap.docx")

doc = Document(INPUT)
body = doc.element.body

def find_heading(body, heading_text_contains, heading_level=None):
    """找到包含指定文字的标题段落的XML元素"""
    for elem in body:
        if elem.tag == qn('w:p'):
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'))
                    if style_val and style_val.startswith('Heading'):
                        text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
                        if heading_text_contains in text:
                            if heading_level is None or style_val == f'Heading {heading_level}':
                                return elem
    return None

def get_section_elements(body, start_heading, end_heading):
    """获取两个标题之间的所有XML元素（不含标题本身）"""
    elements = []
    collecting = False
    for elem in list(body):
        if elem is start_heading:
            collecting = True
            continue
        if elem is end_heading:
            break
        if collecting:
            elements.append(elem)
    return elements

def get_heading_text(elem):
    """获取段落的文字"""
    return ''.join(r.text or '' for r in elem.findall(qn('w:r')))

def set_heading_text(elem, new_text):
    """设置段落的文字（保留格式，只改第一个run的文字）"""
    runs = elem.findall(qn('w:r'))
    if runs:
        all_text_parts = []
        for r in runs:
            t = r.find(qn('w:t'))
            if t is not None and t.text:
                all_text_parts.append(t.text)
        
        first_run_set = False
        for r in runs:
            t = r.find(qn('w:t'))
            if t is not None:
                if not first_run_set:
                    t.text = new_text
                    t.set(qn('xml:space'), 'preserve')
                    first_run_set = True
                else:
                    t.text = ''

def swap_sections(body, h1_elem, h2_elem, h_end_elem, 
                  new_h1_text, new_h2_text):
    """
    交换两个连续章节的内容。
    h1_elem: 第一个章节的标题
    h2_elem: 第二个章节的标题
    h_end_elem: 第二个章节后的下一个标题（边界）
    """
    content1 = get_section_elements(body, h1_elem, h2_elem)
    content2 = get_section_elements(body, h2_elem, h_end_elem)
    
    print(f"  章节1内容: {len(content1)} 个元素")
    print(f"  章节2内容: {len(content2)} 个元素")
    
    content1_copy = [copy.deepcopy(e) for e in content1]
    content2_copy = [copy.deepcopy(e) for e in content2]
    
    for e in content1:
        body.remove(e)
    for e in content2:
        body.remove(e)
    
    insert_point = h1_elem
    for e in content2_copy:
        insert_point.addnext(e)
        insert_point = e
    
    insert_point = h2_elem
    for e in content1_copy:
        insert_point.addnext(e)
        insert_point = e
    
    set_heading_text(h1_elem, new_h1_text)
    set_heading_text(h2_elem, new_h2_text)

# ============================================================
# 第5章: 交换 5.3(教练) 和 5.4(工作人员)
# ============================================================
print("\n" + "=" * 60)
print("交换第5章 5.3 和 5.4")
print("=" * 60)

h53 = find_heading(body, '5.3', 2)
h54 = find_heading(body, '5.4', 2)
h55 = find_heading(body, '5.5', 2)

if h53 and h54 and h55:
    old_53 = get_heading_text(h53)
    old_54 = get_heading_text(h54)
    print(f"  找到 5.3: {old_53}")
    print(f"  找到 5.4: {old_54}")
    print(f"  找到 5.5 边界: {get_heading_text(h55)}")
    
    swap_sections(body, h53, h54, h55,
                  old_53.replace('教练', '工作人员'),
                  old_54.replace('工作人员', '教练'))
    print("  交换完成!")
else:
    print("  [ERROR] 未找到所有标题")

# ============================================================
# 第4章: 交换 4.4(业务流程) 和 4.5(数据库设计)
# ============================================================
print("\n" + "=" * 60)
print("交换第4章 4.4 和 4.5")
print("=" * 60)

h44 = find_heading(body, '4.4', 2)
h45 = find_heading(body, '4.5', 2)
h46 = find_heading(body, '4.6', 2)

if h44 and h45 and h46:
    old_44 = get_heading_text(h44)
    old_45 = get_heading_text(h45)
    print(f"  找到 4.4: {old_44}")
    print(f"  找到 4.5: {old_45}")
    print(f"  找到 4.6 边界: {get_heading_text(h46)}")
    
    content_44 = get_section_elements(body, h44, h45)
    content_45 = get_section_elements(body, h45, h46)
    
    print(f"  4.4 内容: {len(content_44)} 个元素")
    print(f"  4.5 内容: {len(content_45)} 个元素")
    
    content_44_copy = [copy.deepcopy(e) for e in content_44]
    content_45_copy = [copy.deepcopy(e) for e in content_45]
    
    for e in content_44:
        body.remove(e)
    for e in content_45:
        body.remove(e)
    
    insert_point = h44
    for e in content_45_copy:
        insert_point.addnext(e)
        insert_point = e
    
    insert_point = h45
    for e in content_44_copy:
        insert_point.addnext(e)
        insert_point = e
    
    new_44_text = old_44.replace('业务流程设计', '数据库设计').replace('4.4', '4.4')
    new_45_text = old_45.replace('数据库设计', '业务流程设计').replace('4.5', '4.5')
    set_heading_text(h44, new_44_text)
    set_heading_text(h45, new_45_text)
    
    # 还需要处理 4.5 原来的子标题 4.5.1 和 4.5.2 -> 改为 4.4.1 和 4.4.2
    # 这些子标题现在在新的 4.4 内容区域中
    for elem in body:
        if elem.tag == qn('w:p'):
            text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading 3':
                    if '4.5.1' in text:
                        set_heading_text(elem, text.replace('4.5.1', '4.4.1'))
                        print(f"  子标题: {text} -> {text.replace('4.5.1', '4.4.1')}")
                    elif '4.5.2' in text:
                        set_heading_text(elem, text.replace('4.5.2', '4.4.2'))
                        print(f"  子标题: {text} -> {text.replace('4.5.2', '4.4.2')}")
    
    print("  交换完成!")
else:
    print(f"  [ERROR] 未找到所有标题: h44={h44 is not None}, h45={h45 is not None}, h46={h46 is not None}")

# ============================================================
# 更新第5章图编号: 
# 交换后 5.3=工作人员(原5.4的内容), 5.4=教练(原5.3的内容)
# 原图编号: 图5.4教练 | 图5.5工作人员后台 | 图5.6工作人员模块
# 新图编号: 图5.4工作人员后台 | 图5.5工作人员模块 | 图5.6教练预约
# ============================================================
print("\n" + "=" * 60)
print("更新第5章图编号")
print("=" * 60)

fig_renames_ch5 = {
    '图5.4　教练预约页面': '图5.6　教练预约页面',
    '图5.5　工作人员后台页面': '图5.4　工作人员后台页面',
    '图5.6': None,  # 标记用
}

for elem in body:
    if elem.tag == qn('w:p'):
        text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
        if '图5.4' in text and '教练预约' in text:
            set_heading_text(elem, text.replace('图5.4', '图5.6'))
            print(f"  {text.strip()} -> 图5.6")
        elif '图5.5' in text and '工作人员后台' in text:
            set_heading_text(elem, text.replace('图5.5', '图5.4'))
            print(f"  {text.strip()} -> 图5.4")
        elif '图5.6' in text and '工作人员' in text:
            set_heading_text(elem, text.replace('图5.6', '图5.5'))
            print(f"  {text.strip()} -> 图5.5")

# ============================================================
# 更新第4章引用编号
# 原: 4.4业务流程中有图和表，4.5数据库中有图和表
# 交换后需更新相关表/图编号
# ============================================================
print("\n" + "=" * 60)
print("更新第4章表/图编号")
print("=" * 60)

# 找到所有在4.4和4.5范围内的表/图标题，更新编号
h44_new = find_heading(body, '4.4', 2)
h45_new = find_heading(body, '4.5', 2)
h46_new = find_heading(body, '4.6', 2)

if h44_new and h45_new and h46_new:
    in_44 = False
    in_45 = False
    for elem in body:
        if elem is h44_new:
            in_44 = True
            in_45 = False
            continue
        if elem is h45_new:
            in_44 = False
            in_45 = True
            continue
        if elem is h46_new:
            break
        
        if elem.tag == qn('w:p'):
            text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
            
            if in_44:
                # 原4.5的内容现在在4.4中, 表编号 4.5.x -> 暂不改，因为表编号可能不跟章节号
                pass
            if in_45:
                pass

# ============================================================
# 修复6.4口语化
# ============================================================
print("\n" + "=" * 60)
print("修复6.4口语化")
print("=" * 60)

oral_fixes = {
    '一个是商品管理页面的状态字段，后端返回的是数字（0表示下架、1表示上架），前端没有做转换直接显示数字，用户无法理解其含义。随后在前端增加了状态映射，把0显示为"下架"、把1显示为"架中"解决了。另一个是租赁订单的金额显示，有时候会显示很多位小数，比如显示"25.000000001"种，随后在前端使用toFixed(2)进行了格式化处理。此外，教练预约的时间段校验不够严格，不应允许预约过去的时间，但系统没有做这个检查，随后在前端为日期选择器添加了disabledDate属性来禁止选择今天之前的日期。以上均为非核心问题，发现后已及时修复。':
    '其一，商品管理页面的状态字段中，后端返回的是数字编码（0表示下架、1表示上架），前端未做转换而直接显示原始数字，导致用户无法理解其含义，后通过在前端增加状态映射予以修复。其二，租赁订单金额存在浮点精度问题，偶尔出现多位小数的情况，后在前端使用toFixed(2)进行格式化处理。其三，教练预约的时间段校验不够严格，未限制用户选择已过去的时间，后在前端为日期选择器添加disabledDate属性加以约束。以上均属非核心缺陷，发现后已及时修复。',
    
    '此外，在并发控制方面也存在待优化之处。当两个用户同时租赁同一件商品的时候，可能会出现库存变成负数的情况。这个问题的根源是在扣减库存的时候没有做并发控制。应在数据库层面采用乐观锁机制加以处理，由于本系统定位为中小型雪场使用，并发量不会太大，该问题在当前阶段影响有限，但随着用户规模扩大，该部分仍有优化的必要。':
    '此外，在并发控制方面仍存在待优化之处。当多个用户同时租赁同一件商品时，存在库存出现负数的风险，其根本原因在于扣减库存时未引入并发控制机制。应在数据库层面采用乐观锁等手段加以处理。鉴于本系统定位为中小型雪场使用，并发量相对有限，该问题在当前阶段影响较小，但随着用户规模扩大，该部分仍有优化的必要。',

    '总的来说，系统基本达到了设计的功能目标，主要的业务流程均可正常运行。':
    '综上所述，系统基本达到了预期的功能目标，各项核心业务流程均可正常运行。',
}

fix_count = 0
for elem in body:
    if elem.tag == qn('w:p'):
        text = ''.join(r.text or '' for r in elem.findall(qn('w:r')))
        for old_text, new_text in oral_fixes.items():
            if old_text[:30] in text:
                set_heading_text(elem, new_text)
                fix_count += 1
                print(f"  已修复第{fix_count}处: {old_text[:40]}... -> {new_text[:40]}...")

print(f"\n  共修复 {fix_count} 处口语化表述")

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT)
print(f"\n已保存到 {OUTPUT}")

# ============================================================
# 验证
# ============================================================
print("\n" + "=" * 60)
print("验证结果")
print("=" * 60)

doc2 = Document(OUTPUT)
print("\n第4章结构:")
in_ch4 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第4章' in text:
        in_ch4 = True
    elif in_ch4 and style == 'Heading 1':
        break
    if in_ch4 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n第5章结构:")
in_ch5 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style.startswith('Heading'):
        print(f"  [{i}] [{style}] {text}")

print("\n第5章图片验证:")
in_ch5 = False
rels = doc2.part.rels
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and '图5' in text:
        print(f"  [{i}] {text}")
    if in_ch5:
        drawings = p._element.findall('.//' + qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed and embed in rels:
                    print(f"  [{i}] IMG -> {rels[embed].target_ref}")
