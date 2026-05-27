# -*- coding: utf-8 -*-
import docx
doc = docx.Document(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx')
targets = [57, 58, 59, 74, 171, 172, 173, 244, 296, 302, 304]
with open(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\verify.txt', 'w', encoding='utf-8') as f:
    for i in targets:
        p = doc.paragraphs[i]
        f.write(f'=== [{i}] ===\n{p.text}\n\n')
    f.write('=== REFERENCES ===\n')
    in_ref = False
    for p in doc.paragraphs:
        if '\u53c2\u8003\u6587\u732e' in p.text and p.style.name == 'Heading 1':
            in_ref = True
            continue
        if in_ref and p.text.strip().startswith('['):
            f.write(p.text[:120] + '\n')
        elif in_ref and p.style.name == 'Heading 1':
            break
    f.write('\n=== SUPERSCRIPT CHECK ===\n')
    for i in [56, 57, 73, 79, 82]:
        p = doc.paragraphs[i]
        for r in p.runs:
            if r.font.superscript:
                f.write(f'Para [{i}]: superscript = "{r.text}"\n')
print('done')
