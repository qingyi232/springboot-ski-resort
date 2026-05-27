# -*- coding: utf-8 -*-
"""
更新目录文本以匹配当前文档标题结构
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document('5.26.2.docx')
body = doc.element.body

# 获取sdt中的TOC段落
sdts = body.findall(qn('w:sdt'))
sdt = sdts[0]
content = sdt.find(qn('w:sdtContent'))
toc_paras = content.findall(qn('w:p'))

def get_toc_text(p):
    return ''.join(r.text or '' for r in p.findall('.//' + qn('w:r')))

def set_toc_text(p, new_text):
    """设置TOC段落的文字，保留格式和制表符+页码"""
    runs = p.findall('.//' + qn('w:r'))
    if not runs:
        return
    
    # 收集所有run的文字
    old_full = ''
    for r in runs:
        t = r.find(qn('w:t'))
        if t is not None and t.text:
            old_full += t.text
    
    # 提取页码（制表符后的数字）
    page_num = ''
    if '\t' in old_full:
        parts = old_full.rsplit('\t', 1)
        if len(parts) > 1:
            page_num = parts[-1].strip()
    
    # 合成新文本
    if page_num:
        final_text = new_text + '\t' + page_num
    else:
        final_text = new_text
    
    # 写入第一个有文字的run，清空其他
    first_set = False
    for r in runs:
        t = r.find(qn('w:t'))
        if t is not None:
            if not first_set:
                t.text = final_text
                t.set(qn('xml:space'), 'preserve')
                first_set = True
            else:
                t.text = ''

# 打印当前TOC
print("=== 当前目录 ===")
for j, p in enumerate(toc_paras):
    text = get_toc_text(p).strip()
    if text:
        print(f"  [{j}] {text[:80]}")

# 修改计划
print("\n=== 开始修改 ===")

changes = 0
for j, p in enumerate(toc_paras):
    text = get_toc_text(p).strip()
    
    # 去掉页码部分来匹配
    text_no_page = text.rsplit('\t', 1)[0].strip() if '\t' in text else text.strip()
    
    # Chapter 3: 去掉 3.4 相关条目
    if text_no_page.startswith('3.4'):
        # 清空这些段落的内容（设为空）
        for r in p.findall('.//' + qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None:
                t.text = ''
        changes += 1
        print(f"  [{j}] 删除: {text_no_page}")
    
    # Chapter 5 重建
    elif '5.1' in text_no_page and '开发环境' in text_no_page:
        set_toc_text(p, '5.1　管理员模块的实现')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.1　管理员模块的实现")
    
    elif '5.2' in text_no_page and '用户登录' in text_no_page:
        set_toc_text(p, '5.1.1　商品管理功能')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.1.1　商品管理功能")
    
    elif '5.3' in text_no_page and '商品管理' in text_no_page:
        set_toc_text(p, '5.1.2　数据统计功能')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.1.2　数据统计功能")
    
    elif '5.4' in text_no_page and '教练预约' in text_no_page:
        set_toc_text(p, '5.2　用户模块的实现')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.2　用户模块的实现")
    
    elif '5.5' in text_no_page and '课程' in text_no_page:
        set_toc_text(p, '5.3　工作人员模块的实现')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.3　工作人员模块的实现")
    
    elif '5.6' in text_no_page and '数据统计' in text_no_page:
        set_toc_text(p, '5.4　教练模块的实现')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.4　教练模块的实现")
    
    elif '5.7' in text_no_page and '系统安全' in text_no_page:
        set_toc_text(p, '5.5　系统安全实现')
        changes += 1
        print(f"  [{j}] {text_no_page} -> 5.5　系统安全实现")

doc.save('5.26.2.docx')
print(f"\n共修改 {changes} 处, 已保存")

# 删除空的3.4条目的段落（清空文字后的空段落）
# 需要从sdt content中移除
doc2 = Document('5.26.2.docx')
body2 = doc2.element.body
sdts2 = body2.findall(qn('w:sdt'))
content2 = sdts2[0].find(qn('w:sdtContent'))
toc_paras2 = content2.findall(qn('w:p'))

removed = 0
for p in list(toc_paras2):
    text = get_toc_text(p).strip()
    if not text:
        # Check if this was a 3.4 entry we cleared
        content2.remove(p)
        removed += 1

if removed > 0:
    doc2.save('5.26.2.docx')
    print(f"已移除 {removed} 个空段落")

# 验证
print("\n=== 修改后目录 ===")
doc3 = Document('5.26.2.docx')
body3 = doc3.element.body
sdts3 = body3.findall(qn('w:sdt'))
content3 = sdts3[0].find(qn('w:sdtContent'))
toc_paras3 = content3.findall(qn('w:p'))

for j, p in enumerate(toc_paras3):
    text = get_toc_text(p).strip()
    if text:
        print(f"  [{j}] {text[:80]}")
