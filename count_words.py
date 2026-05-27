from docx import Document
d = Document('论文_排版版v3.docx')
text = ''.join([p.text for p in d.paragraphs])
t = ''.join([c.text for tbl in d.tables for r in tbl.rows for c in r.cells])
total = text + t
chn = sum(1 for c in total if '\u4e00' <= c <= '\u9fff')
clean = total.replace('\n','').replace(' ','')
print(f'段落字符数: {len(text)}')
print(f'表格字符数: {len(t)}')
print(f'中文字数: {chn}')
print(f'去空格总字符: {len(clean)}')
