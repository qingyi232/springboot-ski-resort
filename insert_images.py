# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
import os

DOC_PATH = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"
OUT_PATH = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"
IMG_DIR = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现"

doc = Document(DOC_PATH)

# Map: caption text -> (new image path, width)
replacements = {
    "系统功能结构图": (os.path.join(IMG_DIR, "structure_chart.png"), Cm(16)),
    "用户登录流程图": (os.path.join(IMG_DIR, "login_flow.png"), Cm(11)),
    "商品租赁业务流程图": (os.path.join(IMG_DIR, "rental_flow.png"), Cm(11)),
}

# Find image paragraphs by looking for caption text in surrounding paragraphs
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    for caption_key, (img_path, img_width) in replacements.items():
        if caption_key in text and "图" in text:
            # The image is in the paragraph BEFORE the caption
            if i > 0:
                img_para = doc.paragraphs[i - 1]
                # Check if it has an image
                has_img = any(
                    run._element.findall(qn('w:drawing')) or 
                    run._element.findall(qn('w:pict'))
                    for run in img_para.runs
                )
                if has_img:
                    # Clear existing runs (remove old image)
                    for run in img_para.runs:
                        run._element.getparent().remove(run._element)
                    
                    # Add new image
                    run = img_para.add_run()
                    run.add_picture(img_path, width=img_width)
                    print(f"  Replaced image at P{i-1} for '{caption_key}'")
                    break

doc.save(OUT_PATH)
print(f"\nSaved to {OUT_PATH}")
