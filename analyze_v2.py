import docx
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import sys

doc = docx.Document(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx')

result = {
    "total_paragraphs": len(doc.paragraphs),
    "total_tables": len(doc.tables),
    "tables": [],
    "sections": [],
    "images": [],
    "code_blocks": []
}

# Analyze tables
for i, t in enumerate(doc.tables):
    table_info = {
        "index": i + 1,
        "rows": len(t.rows),
        "cols": len(t.columns),
        "first_cell": t.cell(0, 0).text[:60],
        "content": []
    }
    for ri, row in enumerate(t.rows):
        cells = [cell.text[:40] for cell in row.cells]
        table_info["content"].append(cells)
    result["tables"].append(table_info)

# Analyze document body order (paragraphs, tables, images)
from docx.oxml.ns import qn
body = doc.element.body
para_idx = 0
table_idx = 0
element_order = []

for element in body:
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    if tag == 'p':
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
            style = p.style.name if p.style else 'None'
            text = p.text[:100]
            
            # Check for images in paragraph
            has_image = False
            for run in p.runs:
                if run._element.findall(qn('w:drawing')):
                    has_image = True
                    break
                if run._element.findall(qn('w:pict')):
                    has_image = True
                    break
            
            # Check if it looks like code
            is_code = False
            if any(kw in text for kw in ['@PostMapping', '@GetMapping', '@PutMapping', 'public ', 'private ', 'return ', 'router.', 'const ', 'import ', 'String ', 'Long ', 'Map<', 'List<', 'Result<', 'void ', '.put(', '.get(', 'next(']):
                is_code = True
            
            if text.strip() or has_image:
                entry = {
                    "type": "paragraph",
                    "index": para_idx,
                    "style": style,
                    "text": text,
                    "has_image": has_image,
                    "is_code": is_code
                }
                element_order.append(entry)
            para_idx += 1
    elif tag == 'tbl':
        if table_idx < len(doc.tables):
            t = doc.tables[table_idx]
            entry = {
                "type": "table",
                "table_index": table_idx + 1,
                "rows": len(t.rows),
                "cols": len(t.columns),
                "first_cell": t.cell(0, 0).text[:60]
            }
            element_order.append(entry)
            table_idx += 1

# Find chapter 5 content (system implementation)
chapter5_elements = []
in_chapter5 = False
for elem in element_order:
    if elem["type"] == "paragraph":
        if "第5章" in elem.get("text", "") or "5.1" in elem.get("text", ""):
            in_chapter5 = True
        elif "第6章" in elem.get("text", ""):
            in_chapter5 = False
    if in_chapter5:
        chapter5_elements.append(elem)

result["chapter5_elements"] = chapter5_elements
result["element_order_count"] = len(element_order)

# Write output as JSON
with open(r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\doc_analysis.json', 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print("Analysis complete. Output written to doc_analysis.json")
