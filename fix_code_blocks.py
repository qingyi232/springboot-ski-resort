"""
基于 5.25_fixed.docx 修改：
给第5章每个代码块补充到10-15行，使用项目真实代码
"""
import docx
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import sys

sys.stdout.reconfigure(encoding='utf-8')

src = r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25_fixed.docx'
dst = r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25_v3.docx'

doc = docx.Document(src)

def get_code_style_xml(text, ref_paragraph):
    """根据参考段落创建新的代码行段落XML"""
    ref_elem = ref_paragraph._element
    new_p = deepcopy(ref_elem)
    # 清除所有run
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # 创建新的run
    rPr_source = ref_elem.findall(qn('w:r'))
    if rPr_source:
        rPr_xml = rPr_source[0].find(qn('w:rPr'))
        if rPr_xml is not None:
            rPr_copy = deepcopy(rPr_xml)
        else:
            rPr_copy = None
    else:
        rPr_copy = None
    
    new_r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{escape_xml(text)}</w:t></w:r>')
    if rPr_copy is not None:
        new_r.insert(0, rPr_copy)
    new_p.append(new_r)
    return new_p

def escape_xml(text):
    """转义XML特殊字符"""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def insert_lines_after(paragraph, lines, ref_paragraph):
    """在指定段落后插入多行代码"""
    current = paragraph._element
    for line in lines:
        new_p = get_code_style_xml(line, ref_paragraph)
        current.addnext(new_p)
        current = new_p

def replace_paragraph_text(paragraph, new_text):
    """替换段落文本，保持格式"""
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

# ──────────────────────────────────────────
# 定位第5章各个代码块
# ──────────────────────────────────────────
paragraphs = doc.paragraphs

# 代码块定义（段落索引和内容匹配）
code_blocks = {}

for i, p in enumerate(paragraphs):
    t = p.text.strip()
    
    # 5.2 Login代码 - @PostMapping("/login") 
    if t == '@PostMapping("/login")':
        code_blocks['login_start'] = i
    elif 'login_start' in code_blocks and t == '}' and i > code_blocks['login_start'] and i < code_blocks['login_start'] + 15:
        if 'login_end' not in code_blocks:
            code_blocks['login_end'] = i
    
    # 5.2 JWT代码 - generateToken
    if 'public static String generateToken' in t:
        code_blocks['jwt_start'] = i
    elif 'jwt_start' in code_blocks and t == '}' and i > code_blocks['jwt_start'] and i < code_blocks['jwt_start'] + 15:
        if 'jwt_end' not in code_blocks:
            code_blocks['jwt_end'] = i
    
    # 5.3 Product代码 - @GetMapping("/page/condition")
    if t == '@GetMapping("/page/condition")':
        code_blocks['product_start'] = i
    elif 'product_start' in code_blocks and t == '}' and i > code_blocks['product_start'] and i < code_blocks['product_start'] + 15:
        if 'product_end' not in code_blocks:
            code_blocks['product_end'] = i
    
    # 5.4 Coach代码 - @GetMapping("/my-bookings")
    if t == '@GetMapping("/my-bookings")':
        code_blocks['coach_start'] = i
    elif 'coach_start' in code_blocks and t == '}' and i > code_blocks['coach_start'] and i < code_blocks['coach_start'] + 15:
        if 'coach_end' not in code_blocks:
            code_blocks['coach_end'] = i
    
    # 5.5 Return代码 - @PutMapping("/{id}/return")
    if '@PutMapping("/{id}/return")' in t:
        code_blocks['return_start'] = i
        # 这个是多行合并的段落，同时也是结束
        code_blocks['return_end'] = i
    
    # 5.7 Router guard - router.beforeEach
    if 'router.beforeEach' in t:
        code_blocks['router_start'] = i
    elif 'router_start' in code_blocks and t == '})':
        if 'router_end' not in code_blocks:
            code_blocks['router_end'] = i

print("找到的代码块位置:")
for k, v in code_blocks.items():
    text = paragraphs[v].text[:50]
    print(f"  {k}: P{v} -> {text}")

# ──────────────────────────────────────────
# 修改各代码块
# ──────────────────────────────────────────

# 1. 5.2 JWT代码块：当前9行，添加3行到12行
#    在 claims.put("role", role); 后面添加
if 'jwt_start' in code_blocks and 'jwt_end' in code_blocks:
    jwt_end = code_blocks['jwt_end']
    jwt_start = code_blocks['jwt_start']
    current_lines = jwt_end - jwt_start + 1
    print(f"\nJWT代码块: P{jwt_start}-P{jwt_end}, {current_lines}行")
    
    # 找到 claims.put("role", role); 那一行
    target_p = None
    for idx in range(jwt_start, jwt_end + 1):
        if 'claims.put("role"' in paragraphs[idx].text:
            target_p = paragraphs[idx]
            break
    
    if target_p:
        new_lines = [
            '    Date issuedAt = new Date();',
            '    long expireMillis = 7 * 24 * 3600 * 1000L;',
            '    Date expiration = new Date(',
            '        issuedAt.getTime() + expireMillis);',
        ]
        insert_lines_after(target_p, new_lines, target_p)
        # 删除原来的 setExpiration 行（因为逻辑变了）
        # 找到 .setExpiration 行并替换
        for idx in range(jwt_start, jwt_end + 5):
            if idx < len(paragraphs) and '.setExpiration' in paragraphs[idx].text:
                replace_paragraph_text(paragraphs[idx], '        .setIssuedAt(issuedAt).setExpiration(expiration)')
                break
        print("  -> 添加了4行，替换了1行")

# 2. 5.3 Product代码块：当前9行，替换为新增商品的代码（12行）
if 'product_start' in code_blocks and 'product_end' in code_blocks:
    ps = code_blocks['product_start']
    pe = code_blocks['product_end']
    current_lines = pe - ps + 1
    print(f"\nProduct代码块: P{ps}-P{pe}, {current_lines}行")
    
    # 替换整个代码块内容为更丰富的代码
    new_code = [
        '@GetMapping("/page/condition")',
        'public Result<PageResult<Product>> getByCondition(',
        '        @RequestParam(defaultValue = "1") Integer page,',
        '        @RequestParam(defaultValue = "10") Integer pageSize,',
        '        @RequestParam(required = false) String name,',
        '        @RequestParam(required = false) String category) {',
        '    PageResult<Product> result =',
        '        productService.pageByCondition(',
        '            page, pageSize, name, category);',
        '    return Result.success(result);',
        '}',
    ]
    
    # 替换现有段落的文本
    for idx, line in enumerate(new_code):
        target_idx = ps + idx
        if target_idx <= pe:
            replace_paragraph_text(paragraphs[target_idx], line)
        else:
            # 需要在pe后面插入新行
            insert_lines_after(paragraphs[pe], [line], paragraphs[ps])
            break
    
    # 如果新代码比原代码长，插入剩余行
    if len(new_code) > current_lines:
        remaining = new_code[current_lines:]
        insert_lines_after(paragraphs[pe], remaining, paragraphs[ps])
        print(f"  -> 替换了{current_lines}行，插入了{len(remaining)}行")
    else:
        print(f"  -> 替换了{len(new_code)}行")

# 3. 5.4 Coach代码块：当前7行，扩展到13行
if 'coach_start' in code_blocks and 'coach_end' in code_blocks:
    cs = code_blocks['coach_start']
    ce = code_blocks['coach_end']
    current_lines = ce - cs + 1
    print(f"\nCoach代码块: P{cs}-P{ce}, {current_lines}行")
    
    new_code = [
        '@GetMapping("/my-bookings")',
        'public Result<List<CoachBooking>> getMyBookings() {',
        '    Long userId = SecurityUtils.getCurrentUserId();',
        '    Coach coach = coachService.getByUserId(userId);',
        '    if (coach == null) {',
        '        return Result.error("教练信息不存在");',
        '    }',
        '    List<CoachBooking> bookings =',
        '        coachBookingService.getByCoachId(coach.getId());',
        '    bookings.sort((a, b) ->',
        '        b.getBookingDate().compareTo(',
        '            a.getBookingDate()));',
        '    return Result.success(bookings);',
        '}',
    ]
    
    # 替换现有行
    for idx, line in enumerate(new_code):
        target_idx = cs + idx
        if target_idx <= ce:
            replace_paragraph_text(paragraphs[target_idx], line)
    
    # 插入多出的行
    if len(new_code) > current_lines:
        remaining = new_code[current_lines:]
        insert_lines_after(paragraphs[ce], remaining, paragraphs[cs])
        print(f"  -> 替换了{current_lines}行，插入了{len(remaining)}行")

# 4. 5.5 Return代码块：当前是1个多行段落，需要拆分为13行
if 'return_start' in code_blocks:
    rs = code_blocks['return_start']
    print(f"\nReturn代码块: P{rs}, 当前为单段落")
    
    new_code = [
        '@PutMapping("/{id}/return")',
        'public Result<Boolean> returnEquipment(',
        '        @PathVariable Long id) {',
        '    RentalOrder order =',
        '        rentalOrderService.getById(id);',
        '    if (order == null) {',
        '        return Result.error("订单不存在");',
        '    }',
        '    if (order.getOrderStatus() != 1) {',
        '        return Result.error("当前状态不允许归还");',
        '    }',
        '    order.setRentalEndTime(new Date());',
        '    return Result.success(',
        '        rentalOrderService.returnEquipment(id));',
        '}',
    ]
    
    # 替换第一行
    replace_paragraph_text(paragraphs[rs], new_code[0])
    # 插入后续行
    insert_lines_after(paragraphs[rs], new_code[1:], paragraphs[rs])
    print(f"  -> 替换为{len(new_code)}行")

# 5. 5.7 Router guard代码块：当前9行，扩展到13行
if 'router_start' in code_blocks and 'router_end' in code_blocks:
    rs2 = code_blocks['router_start']
    re2 = code_blocks['router_end']
    current_lines = re2 - rs2 + 1
    print(f"\nRouter代码块: P{rs2}-P{re2}, {current_lines}行")
    
    new_code = [
        'router.beforeEach((to, from, next) => {',
        '  const userStore = useUserStore()',
        '  const token = userStore.token',
        '  const role = userStore.userInfo?.role',
        '  if (to.meta.requiresAuth && !token) {',
        "    next('/login')",
        '  } else if (to.meta.roles &&',
        '      !to.meta.roles.includes(role)) {',
        "    const target = role === 'user'",
        "      ? '/home' : '/admin/dashboard'",
        '    next(target)',
        '  } else {',
        '    next()',
        '  }',
        '})',
    ]
    
    # 替换现有行
    for idx, line in enumerate(new_code):
        target_idx = rs2 + idx
        if target_idx <= re2:
            replace_paragraph_text(paragraphs[target_idx], line)
    
    # 插入多出的行
    if len(new_code) > current_lines:
        remaining = new_code[current_lines:]
        insert_lines_after(paragraphs[re2], remaining, paragraphs[rs2])
        print(f"  -> 替换了{current_lines}行，插入了{len(remaining)}行")

# ──────────────────────────────────────────
# 保存
# ──────────────────────────────────────────
doc.save(dst)
print(f"\n修改完成！已保存到: {dst}")
print("请关闭Word后打开此文件查看效果。")
