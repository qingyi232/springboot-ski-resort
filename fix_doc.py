# -*- coding: utf-8 -*-
"""
5.26.1.docx modification script
"""
import copy
import re
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips
from lxml import etree

SRC = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx"
DST = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"

doc = Document(SRC)
body = doc.element.body

# ============================================================
# Helper: replace text in a paragraph while preserving formatting
# ============================================================
def para_replace(para, old, new):
    """Replace text in paragraph preserving run formatting."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # Rebuild runs: put all text in first run, clear others
    if not para.runs:
        return False
    # Collect all run formatting
    runs = para.runs
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ""
    return True

def para_replace_re(para, pattern, replacement):
    """Regex replace in paragraph text, preserving first run format."""
    full = para.text
    new_full = re.sub(pattern, replacement, full)
    if new_full == full:
        return False
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ""
    return True


# ============================================================
# 1. 修改 6.4 口语化表达
# ============================================================
print("=== Step 1: 修改 6.4 口语化 ===")

p408 = doc.paragraphs[408]
# "还有一个问题是教练预约" → "此外，教练预约"
para_replace(p408, "还有一个问题是教练预约", "此外，教练预约")
# "理论上不应该允许" → "不应允许"
para_replace(p408, "理论上不应该允许", "不应允许")
print(f"  P408 fixed")

p409 = doc.paragraphs[409]
# "还有一个在测试中发现的问题是关于并发的。" → "此外，在并发控制方面也存在待优化之处。"
para_replace(p409, "还有一个在测试中发现的问题是关于并发的。", "此外，在并发控制方面也存在待优化之处。")
# "理论上应该在数据库层面用乐观锁来处理" → "应在数据库层面采用乐观锁机制加以处理"
para_replace(p409, "理论上应该在数据库层面用乐观锁来处理", "应在数据库层面采用乐观锁机制加以处理")
# "不过因为" → "由于"
para_replace(p409, "不过因为", "由于")
# "所以这个问题在当前阶段影响不大" → "该问题在当前阶段影响有限"
para_replace(p409, "所以这个问题在当前阶段影响不大", "该问题在当前阶段影响有限")
# "但如果以后用户量增加的话，这个地方是需要优化的" → "但随着用户规模扩大，该部分仍有优化的必要"
para_replace(p409, "但如果以后用户量增加的话，这个地方是需要优化的", "但随着用户规模扩大，该部分仍有优化的必要")
print(f"  P409 fixed")

p410 = doc.paragraphs[410]
# "都能正常走通" → "均可正常运行"
para_replace(p410, "都能正常走通", "均可正常运行")
print(f"  P410 fixed")


# ============================================================
# 2. 修改 7.1 笔者和最后一句
# ============================================================
print("\n=== Step 2: 修改 7.1 ===")

p415 = doc.paragraphs[415]
# "笔者收获颇多" → "收获颇多"
para_replace(p415, "笔者收获颇多", "收获颇多")
# "这些在小型练习中难以遇到" → "这些在课程中难以遇到"
para_replace(p415, "这些在小型练习中难以遇到", "这些在课程中难以遇到")
print(f"  P415 fixed")


# ============================================================
# 3. 参考文献重新排序
# ============================================================
print("\n=== Step 3: 参考文献重新排序 ===")

# 当前引用顺序: [1]-[5]正确, [6]在P93(最后), [7]在P75(应为6)
# 映射: old → new
ref_map = {
    1: 1, 2: 2, 3: 3, 4: 4, 5: 5,
    6: 15,   # 谢希仁 计算机网络 → 最后引用
    7: 6,    # R.Shah → P75首次出现
    8: 7,    # J.Troya MyBatis
    9: 8,    # Z.Wang MySQL
    10: 9,   # 王珊 数据库
    11: 10,  # Redis
    12: 11,  # JWT
    13: 12,  # Vue+ECharts
    14: 13,  # Vue.js Virtual DOM
    15: 14,  # ECharts
}

# 3a: 替换正文中的引用编号 (P56-P421)
# 先用临时占位符避免冲突
for i in range(56, 422):
    p = doc.paragraphs[i]
    text = p.text
    if not re.search(r'\[\d+\]', text):
        continue
    # 先替换为临时占位符 ##N##
    new_text = text
    for old_num in sorted(ref_map.keys(), reverse=True):
        new_text = new_text.replace(f'[{old_num}]', f'[###{ref_map[old_num]}###]')
    # 再替换回正式格式
    new_text = re.sub(r'\[###(\d+)###\]', r'[\1]', new_text)
    
    if new_text != text:
        runs = p.runs
        if runs:
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""
            print(f"  P{i}: refs updated")

# 3b: 去除重复引用
# P75: 原来有[7][4] → 新[6][4], 去掉[4]因为已出现
p75 = doc.paragraphs[75]
para_replace(p75, "[6][4]", "[6]")
print(f"  P75: removed duplicate [4]")

# P88: 原来有[5][13][14][15] → 新[5][12][13][14], 去掉[5]因为已出现
p88 = doc.paragraphs[88]
para_replace(p88, "[5][12]", "[12]")
print(f"  P88: removed duplicate [5]")

# 3c: 参考文献列表重新排序 (P423-P437)
ref_paras_start = 423
ref_paras_end = 437

# 读取所有参考文献段落内容和格式
ref_entries = {}
for i in range(ref_paras_start, ref_paras_end + 1):
    p = doc.paragraphs[i]
    text = p.text
    match = re.match(r'\[(\d+)\]', text)
    if match:
        old_num = int(match.group(1))
        new_num = ref_map.get(old_num, old_num)
        # 替换编号
        new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', text)
        ref_entries[new_num] = {
            'text': new_text,
            'para_idx': i,
            'element': copy.deepcopy(p._element)
        }

# 按新编号排序并写回
sorted_nums = sorted(ref_entries.keys())
for idx, new_num in enumerate(sorted_nums):
    target_para_idx = ref_paras_start + idx
    p = doc.paragraphs[target_para_idx]
    entry = ref_entries[new_num]
    
    # 保留段落格式但替换内容
    runs = p.runs
    if runs:
        runs[0].text = entry['text']
        for r in runs[1:]:
            r.text = ""
    
    print(f"  Ref [{new_num}] → P{target_para_idx}")

print("  References reordered.")


# ============================================================
# 4. 第4章/第6章三线表间距调整
# ============================================================
print("\n=== Step 4: 三线表间距调整 ===")

def set_para_spacing(para, before_pt=0, after_pt=0):
    """设置段落前后间距"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))

# 找到表标题段落（如"表4.2", "表6.3"等）并减小其前后间距
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # 第4章表标题
    if re.match(r'^(表4\.\d|续表4\.\d)', text):
        set_para_spacing(p, before_pt=3, after_pt=3)
        print(f"  P{i}: ch4 spacing reduced")
    # 第6章表标题
    if re.match(r'^(表6\.\d|续表6\.\d)', text):
        set_para_spacing(p, before_pt=3, after_pt=3)
        print(f"  P{i}: ch6 spacing reduced")

print("  Table spacing adjusted.")


# ============================================================
# 5. 3.4 合并到 3.2
# ============================================================
print("\n=== Step 5: 3.4 合并到 3.2 ===")

# 策略: 把3.4各小节的表格和图片移到3.2对应小节的文字后面
# 3.2.1 管理员功能 P102-P103 → 后面添加3.4.1的表和图(P119表, P121图, P122图标题)
# 3.2.2 教练功能 P104-P105 → 后面添加3.4.2的表和图(P124表, P125图, P126图标题)
# 3.2.3 工作人员功能 P106-P107 → 后面添加3.4.3的表和图
# 3.2.4 普通用户功能 P108-P109 → 后面添加3.4.4的表和图

# 需要在XML层面操作
# 先找到body中所有子元素的索引
def get_body_element_index(elem):
    """获取元素在body中的索引"""
    for idx, child in enumerate(body):
        if child is elem:
            return idx
    return -1

def find_para_body_idx(para_idx):
    """找到指定段落在body中的元素索引"""
    return get_body_element_index(doc.paragraphs[para_idx]._element)

# 收集3.4各小节中需要移动的元素
# 3.4.1: P118(heading) → P119(表标题) → Table → P120(空) → P121(图) → P122(图标题)
# 3.4.2: P123(heading) → P124(表标题) → Table → P125(图) → P126(图标题)
# 3.4.3: P127(heading) → P128(表标题) → Table → P129(图) → P130(图标题)
# 3.4.4: P131(heading) → P132(表标题) → Table → P133(图) → P134(图标题)

# 映射: (3.2.x末尾段落, 3.4.x中需要移动的元素范围)
move_map = [
    (103, 118, 122),  # 管理员: 3.2.1(P103)后 ← 3.4.1(P118-P122)
    (105, 123, 126),  # 教练: 3.2.2(P105)后 ← 3.4.2(P123-P126)
    (107, 127, 130),  # 工作人员: 3.2.3(P107)后 ← 3.4.3(P127-P130)
    (109, 131, 134),  # 普通用户: 3.2.4(P109)后 ← 3.4.4(P131-P134)
]

# 收集要移动的元素(从body中)
# 注意: 段落之间可能有table元素, 需要收集段落和它们之间的tables
def collect_elements_between_paras(start_para_idx, end_para_idx):
    """收集start和end段落(含)之间的所有body子元素(含表格)"""
    start_elem = doc.paragraphs[start_para_idx]._element
    end_elem = doc.paragraphs[end_para_idx]._element
    
    collecting = False
    elements = []
    for child in body:
        if child is start_elem:
            collecting = True
        if collecting:
            elements.append(child)
        if child is end_elem:
            break
    return elements

# 先收集所有要移动的元素(不含heading), 然后移动
# 注意: 移动时需要跳过3.4.x的heading段落(不需要heading了)
moves_data = []
for target_para_idx, src_start, src_end in move_map:
    # 收集3.4.x中heading之后的元素(表标题+表格+图+图标题)
    elems = collect_elements_between_paras(src_start, src_end)
    # 跳过第一个元素(heading)
    elems_to_move = elems[1:]  # 不包含heading
    moves_data.append((target_para_idx, elems_to_move))

# 从后往前移动,避免索引偏移
for target_para_idx, elems_to_move in reversed(moves_data):
    target_elem = doc.paragraphs[target_para_idx]._element
    # 在target_elem之后插入元素(从后往前,保持顺序)
    for elem in reversed(elems_to_move):
        target_elem.addnext(copy.deepcopy(elem))
    print(f"  Moved {len(elems_to_move)} elements after P{target_para_idx}")

# 删除3.4整个小节(P116-P134): heading + 所有内容
# P116 = 3.4 heading, P117 = 文字, P118-P134 = 各小节
elems_to_remove = collect_elements_between_paras(116, 134)
for elem in elems_to_remove:
    body.remove(elem)
print(f"  Removed 3.4 section ({len(elems_to_remove)} elements)")


# ============================================================
# 6. 4.5 移到 4.4 前面
# ============================================================
print("\n=== Step 6: 4.5 移到 4.4 前面 ===")

# 移除3.4后段落索引变了，需要重新定位
# 重新找4.4和4.5的位置
def find_heading_para(text_pattern):
    """查找匹配文本的heading段落索引"""
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name.startswith("Heading"):
            if text_pattern in p.text:
                return i
    return -1

h44_idx = find_heading_para("4.4")
h45_idx = find_heading_para("4.5")
h46_idx = find_heading_para("4.6")

if h44_idx >= 0 and h45_idx >= 0 and h46_idx >= 0:
    print(f"  Found: 4.4 at P{h44_idx}, 4.5 at P{h45_idx}, 4.6 at P{h46_idx}")
    
    # 4.4最后三段(商品租赁流程、用户登录流程、过渡段)需要移到4.5
    # 这三段在4.4的数据库表之后、4.5 heading之前
    # 找到4.5 heading前的最后3个段落
    
    # 收集4.4末尾的流程描述段落(P176-P178位置，但现在索引变了)
    # 找法: 搜索"商品租赁的整体业务流程"和"用户登录的处理流程"
    flow_paras = []
    for i in range(h44_idx, h45_idx):
        p = doc.paragraphs[i]
        if "商品租赁的整体业务流程" in p.text or "用户登录的处理流程" in p.text or "在完成功能模块划分" in p.text:
            flow_paras.append(i)
    
    print(f"  Flow paragraphs to move: {flow_paras}")
    
    # 收集4.5全部元素(heading到4.6 heading前)
    sec45_elems = collect_elements_between_paras(h45_idx, h46_idx - 1)
    
    # 找到4.4 heading的位置，在其前面插入4.5的内容
    h44_elem = doc.paragraphs[h44_idx]._element
    
    # 先将4.5内容复制到4.4前面
    for elem in reversed(sec45_elems):
        h44_elem.addprevious(copy.deepcopy(elem))
    
    # 删除原4.5内容
    for elem in sec45_elems:
        body.remove(elem)
    
    # 然后将4.4末尾的流程段落移到新4.5(现在在4.4内)的图片后面
    # 由于结构变了，4.5的图现在在4.4前面，流程文字应该在4.5的图之前
    # 实际上，用户要求是: 4.4最后三段跟4.5的业务流程图对应起来
    # 即把流程描述段落从4.4末尾移到4.5的流程图前面
    
    # 重新搜索流程段落
    h45_new = find_heading_para("4.5")
    h44_new = find_heading_para("4.4")
    
    if h45_new >= 0 and h44_new >= 0:
        # 找到4.4中的流程描述段落
        flow_para_idxs = []
        for i in range(h44_new, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            if p.style and p.style.name.startswith("Heading") and "4.5" not in p.text and "4.4" not in p.text:
                break
            if "商品租赁的整体业务流程" in p.text:
                flow_para_idxs.append(i)
            elif "用户登录的处理流程" in p.text:
                flow_para_idxs.append(i)
            elif "在完成功能模块划分" in p.text:
                flow_para_idxs.append(i)
        
        if flow_para_idxs:
            # 将这三段移到4.5 heading之后、图之前
            h45_elem_new = doc.paragraphs[h45_new]._element
            for idx in reversed(flow_para_idxs):
                elem = doc.paragraphs[idx]._element
                h45_elem_new.addnext(copy.deepcopy(elem))
                body.remove(elem)
            print(f"  Moved {len(flow_para_idxs)} flow paragraphs to section 4.5")
    
    # 更新标题编号: 4.5→4.4, 4.4→4.5
    h45_final = find_heading_para("4.5")
    h44_final = find_heading_para("4.4")
    
    if h45_final >= 0 and h44_final >= 0 and h45_final < h44_final:
        p45 = doc.paragraphs[h45_final]
        p44 = doc.paragraphs[h44_final]
        # 交换编号
        para_replace(p45, "4.5", "4.4")
        para_replace(p44, "4.4", "4.5")
        print(f"  Swapped section numbers: 4.4<->4.5")
    
    print("  Section 4.5 moved before 4.4")
else:
    print(f"  WARNING: Could not find headings. h44={h44_idx}, h45={h45_idx}, h46={h46_idx}")


# ============================================================
# 7. 第5章按角色分类重组
# ============================================================
print("\n=== Step 7: 第5章按角色分类重组 ===")

# 当前结构: 5.1开发环境, 5.2用户登录, 5.3商品管理, 5.4教练预约, 5.5课程与场地, 5.6数据统计, 5.7系统安全
# 目标结构: 
#   5.1 管理员模块的实现 (含开发环境简述 + 商品管理 + 数据统计)
#   5.2 用户模块的实现 (用户登录 + 教练预约用户端)
#   5.3 工作人员模块的实现 (5.5中工作人员相关)
#   5.4 教练模块的实现 (5.4中教练端)
#   保留5.7

# 策略: 修改heading文本, 重新排列section内容
# 由于大规模移动段落容易出问题, 采用更安全的方式:
# 修改标题文字, 适当调整段落位置

# 找到各节heading
ch5_headings = {}
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 2":
        if "5.1" in p.text:
            ch5_headings["5.1"] = i
        elif "5.2" in p.text:
            ch5_headings["5.2"] = i
        elif "5.3" in p.text:
            ch5_headings["5.3"] = i
        elif "5.4" in p.text:
            ch5_headings["5.4"] = i
        elif "5.5" in p.text:
            ch5_headings["5.5"] = i
        elif "5.6" in p.text:
            ch5_headings["5.6"] = i
        elif "5.7" in p.text:
            ch5_headings["5.7"] = i

print(f"  Chapter 5 headings: {ch5_headings}")

# 修改标题文字
if "5.1" in ch5_headings:
    p = doc.paragraphs[ch5_headings["5.1"]]
    para_replace(p, p.text, "5.1　管理员模块的实现")
    # 在开发环境内容前添加一段说明
    dev_env_para = doc.paragraphs[ch5_headings["5.1"] + 1]
    old_text = dev_env_para.text
    new_intro = "管理员是系统的最高权限角色，负责整体的数据管理和系统配置。在开始各模块的介绍之前，先简要说明开发环境。" + old_text
    if dev_env_para.runs:
        dev_env_para.runs[0].text = new_intro
        for r in dev_env_para.runs[1:]:
            r.text = ""

if "5.2" in ch5_headings:
    p = doc.paragraphs[ch5_headings["5.2"]]
    para_replace(p, p.text, "5.2　用户模块的实现")
    # 添加用户角色说明
    first_para = doc.paragraphs[ch5_headings["5.2"] + 1]
    old_text = first_para.text
    new_intro = "普通用户是系统的最终使用者，通过前台页面访问各项服务功能。本节主要介绍用户端的登录注册和商品租赁功能。" + old_text
    if first_para.runs:
        first_para.runs[0].text = new_intro
        for r in first_para.runs[1:]:
            r.text = ""

if "5.3" in ch5_headings:
    p = doc.paragraphs[ch5_headings["5.3"]]
    # 5.3 原来是商品管理, 但需要放在管理员下
    # 我们需要重新分配: 5.3→管理员的一部分(后续处理)
    # 先改标题
    pass

# 实际上，按照用户要求，需要按角色重新组织内容
# 由于直接移动大量段落风险太高，采用修改标题的方式:
# 5.1 管理员模块 → 保持原5.1(开发环境) + 原5.3(商品管理) + 原5.6(数据统计)
# 5.2 用户模块 → 原5.2(用户登录) + 原5.4部分内容
# 5.3 工作人员模块 → 原5.5部分内容
# 5.4 教练模块 → 原5.4部分内容
# 5.5→5.6 删除/合并
# 5.7 保留

# 当前顺序: 5.1 5.2 5.3 5.4 5.5 5.6 5.7
# 目标顺序: 管理员(5.1开发环境+5.3商品+5.6数据) → 用户(5.2登录) → 工作人员(5.5部分) → 教练(5.4) → 5.7安全

# 需要做的移动:
# 1. 5.3(商品管理)移到5.1后面 → 作为5.1的子内容
# 2. 5.6(数据统计)也移到5.1后面
# 3. 5.5(课程与场地中工作人员部分)改标题为5.3
# 4. 5.4(教练预约)改标题

# 简化方案: 将节标题改为以角色为主题的标题
# 5.1 → "5.1 管理员模块的实现" (已改)
# 5.2 → "5.2 用户模块的实现" (已改) 
# 5.3 → 商品管理(这是管理员功能，但保持位置,改标题为管理员商品管理)
# 5.4 → 教练预约(教练端)
# 5.5 → 工作人员模块
# 5.6 → 数据统计(管理员)

# 更合理的方案: 移动5.3和5.6紧跟5.1后面

# 收集各section的元素
def get_section_elements(heading_idx, next_heading_idx):
    """获取一个section的所有元素(从heading到下一个heading前)"""
    start_elem = doc.paragraphs[heading_idx]._element
    end_limit = doc.paragraphs[next_heading_idx]._element if next_heading_idx < len(doc.paragraphs) else None
    
    elements = []
    collecting = False
    for child in list(body):
        if child is start_elem:
            collecting = True
        if collecting:
            if end_limit and child is end_limit:
                break
            elements.append(child)
    return elements

# 重新查找heading位置(索引可能已变)
ch5_h = {}
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 2":
        for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
            if key in p.text and key not in ch5_h:
                ch5_h[key] = i
                break

print(f"  Updated headings: {ch5_h}")

if all(k in ch5_h for k in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]):
    # 将5.3内容移到5.2前面(紧跟5.1管理员)
    sec53 = get_section_elements(ch5_h["5.3"], ch5_h["5.4"])
    sec56 = get_section_elements(ch5_h["5.6"], ch5_h["5.7"])
    
    # 找到5.2 heading的位置
    sec52_elem = doc.paragraphs[ch5_h["5.2"]]._element
    
    # 在5.2前面插入5.3(商品管理)
    for elem in reversed(sec53):
        sec52_elem.addprevious(copy.deepcopy(elem))
    # 删除原5.3
    for elem in sec53:
        body.remove(elem)
    
    # 重新查找heading
    ch5_h2 = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 2":
            for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
                if key in p.text and ("管理员" in p.text or "用户" in p.text or "商品" in p.text or "教练" in p.text or "课程" in p.text or "数据" in p.text or "安全" in p.text or "开发" in p.text):
                    if key not in ch5_h2:
                        ch5_h2[key] = i
    
    # 更简单的查找
    ch5_h2 = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 2":
            text = p.text
            if "管理员" in text and "5.1" in text:
                ch5_h2["admin"] = i
            elif "商品" in text:
                ch5_h2["product"] = i
            elif "用户" in text and "5.2" in text:
                ch5_h2["user"] = i
            elif "教练" in text and "5.4" in text:
                ch5_h2["coach"] = i
            elif "课程" in text and "5.5" in text:
                ch5_h2["course"] = i
            elif "数据" in text and "5.6" in text:
                ch5_h2["data"] = i
            elif "安全" in text and "5.7" in text:
                ch5_h2["security"] = i
    
    print(f"  After moving 5.3: {ch5_h2}")
    
    # 现在移动5.6(数据统计)到商品管理后面、用户模块前面
    if "data" in ch5_h2 and "user" in ch5_h2 and "security" in ch5_h2:
        sec56_new = get_section_elements(ch5_h2["data"], ch5_h2["security"])
        user_elem = doc.paragraphs[ch5_h2["user"]]._element
        for elem in reversed(sec56_new):
            user_elem.addprevious(copy.deepcopy(elem))
        for elem in sec56_new:
            body.remove(elem)
        print("  Moved 5.6 (data stats) after 5.3 (product)")
    
    # 现在重新编号所有heading
    # 新顺序: 管理员(5.1) → 商品管理(→5.1下) → 数据统计(→5.1下) → 用户(5.2) → 教练(5.3) → 课程与场地(5.4→工作人员) → 安全(5.5→5.7保留)
    
    # 重新查找并重编号
    ch5_sections = []
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 2":
            text = p.text
            if any(kw in text for kw in ["管理员", "商品", "用户", "教练", "课程", "数据", "安全"]):
                ch5_sections.append((i, text))
    
    print(f"  All Ch5 sections: {ch5_sections}")
    
    # 重新命名
    section_counter = 1
    for idx, (pi, text) in enumerate(ch5_sections):
        p = doc.paragraphs[pi]
        if "管理员" in text:
            para_replace(p, p.text, "5.1　管理员模块的实现")
            section_counter = 1
        elif "商品" in text:
            # 作为5.1的子内容，改为不带5.x编号或用5.1.1
            para_replace(p, p.text, "5.1.1　商品管理功能")
            p.style = doc.styles["Heading 3"]
        elif "数据" in text:
            para_replace(p, p.text, "5.1.2　数据统计功能")
            p.style = doc.styles["Heading 3"]
        elif "用户" in text:
            para_replace(p, p.text, "5.2　用户模块的实现")
        elif "教练" in text:
            para_replace(p, p.text, "5.3　教练模块的实现")
        elif "课程" in text:
            para_replace(p, p.text, "5.4　工作人员模块的实现")
        elif "安全" in text:
            para_replace(p, p.text, "5.5　系统安全实现")
    
    print("  Chapter 5 reorganized by roles")
else:
    print(f"  WARNING: Missing headings, skipping Ch5 reorg. Found: {ch5_h}")


# ============================================================
# 8. 第5章精简文字 (5.1-5.6)
# ============================================================
print("\n=== Step 8: 第5章精简文字 ===")

# 找到5.4(原5.5课程与场地→工作人员)中重复冗长的段落
# P300, P301, P302 内容重复描述工作人员功能
for i, p in enumerate(doc.paragraphs):
    text = p.text
    # 找到明显冗长重复的段落(工作人员相关)
    if "工作人员在系统中的操作权限主要包括以下几个方面" in text:
        # 这段和前面的P300、P274内容高度重复, 精简
        new_text = "工作人员的主要操作权限涵盖租赁订单管理、商品信息查看、场地状态管理等方面。"
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        print(f"  P{i}: condensed staff permissions")
    
    if "在滑雪场的日常运营管理中，工作人员承担着重要的前台业务职责" in text:
        # 这段和前面重复, 精简为一句
        new_text = "工作人员在系统中主要负责前台业务操作，登录后台后可使用商品信息管理、租赁订单处理及场地状态查看等功能。"
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        print(f"  P{i}: condensed staff description")
    
    if "在页面布局方面，工作人员后台界面采用" in text:
        # 精简布局描述
        new_text = "工作人员后台采用左侧导航加右侧内容区域的布局方式，菜单项根据角色权限进行了精简。商品管理页面提供按名称和分类检索的功能，租赁订单页面支持按状态筛选并可直接执行归还操作。"
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        print(f"  P{i}: condensed layout description")


# ============================================================
# 保存
# ============================================================
print("\n=== Saving ===")
doc.save(DST)
print(f"Saved to {DST}")
