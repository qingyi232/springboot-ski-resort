# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document('5.26.2.docx')

print("=" * 60)
print("最终验证")
print("=" * 60)

# 1. Chapter 4 structure
print("\n【第4章结构】")
in_ch4 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第4章' in text:
        in_ch4 = True
    elif in_ch4 and style == 'Heading 1':
        break
    if in_ch4 and style.startswith('Heading'):
        print(f"  {text}")

# Verify 4.4 content is database-related
print("\n  4.4节第一段内容验证:")
in_44 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '4.4' in text:
        in_44 = True
        continue
    if in_44 and text:
        print(f"  -> {text[:60]}")
        break

# Verify 4.5 content is business-flow
print("\n  4.5节第一段内容验证:")
in_45 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '4.5' in text:
        in_45 = True
        continue
    if in_45 and text:
        print(f"  -> {text[:60]}")
        break

# 2. Chapter 5 structure + content verification
print("\n【第5章结构】")
in_ch5 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style.startswith('Heading'):
        print(f"  {text}")

# Role order check
print("\n  角色顺序:")
in_ch5 = False
roles = []
for p in doc.paragraphs:
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if in_ch5 and style == 'Heading 2':
        for role in ['管理员', '用户', '工作人员', '教练']:
            if role in text and '工作' not in text[:5] if role != '工作人员' else role in text:
                roles.append(role)

expected = ['管理员', '用户', '工作人员', '教练']
ok = roles[:4] == expected
print(f"  实际: {roles}")
print(f"  期望: {expected}")
print(f"  {'[OK]' if ok else '[X]'}")

# Verify 5.3 content is 工作人员-related
print("\n  5.3节内容验证（应为工作人员内容）:")
in_53 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '5.3' in text:
        in_53 = True
        continue
    if in_53 and style.startswith('Heading 2'):
        break
    if in_53 and text and not style.startswith('Heading'):
        print(f"  -> {text[:80]}")
        has_worker = '工作人员' in text or '课程' in text or '场地' in text or '归还' in text or '订单' in text
        if has_worker:
            print(f"  [OK] 包含工作人员相关关键词")
        break

print("\n  5.4节内容验证（应为教练内容）:")
in_54 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 2' and '5.4' in text:
        in_54 = True
        continue
    if in_54 and style.startswith('Heading 2'):
        break
    if in_54 and text and not style.startswith('Heading'):
        print(f"  -> {text[:80]}")
        has_coach = '教练' in text or '预约' in text
        if has_coach:
            print(f"  [OK] 包含教练相关关键词")
        break

# 3. Image verification
print("\n【第5章图片映射】")
rels = doc.part.rels
img_map = {
    'media/image12.png': '商品管理',
    'media/image16.png': '数据统计',
    'media/image11.png': '用户登录',
    'media/image14.png': '工作人员后台',
    'media/image15.png': '工作人员模块',
    'media/image13.png': '教练预约',
}
in_ch5 = False
all_match = True
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if not in_ch5:
        continue
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            target = rels[embed].target_ref if embed in rels else '?'
            next_text = doc.paragraphs[i+1].text.strip()[:40] if i+1 < len(doc.paragraphs) else ''
            img_content = img_map.get(target, '?')
            match = True
            if '商品管理' in next_text and img_content != '商品管理':
                match = False
            elif '数据统计' in next_text and img_content != '数据统计':
                match = False
            elif '用户登录' in next_text and img_content != '用户登录':
                match = False
            elif '教练' in next_text and img_content != '教练预约':
                match = False
            elif '工作人员后台' in next_text and img_content != '工作人员后台':
                match = False
            elif '工作人员模块' in next_text and img_content != '工作人员模块':
                match = False
            
            status = '[OK]' if match else '[X]'
            if not match:
                all_match = False
            print(f"  {status} {target}({img_content}) -> {next_text}")

print(f"\n  图片总体: {'全部正确' if all_match else '存在错误'}")

# 4. 6.4 check
print("\n【6.4口语化检查】")
in_64 = False
issues = []
for p in doc.paragraphs:
    text = p.text.strip()
    style = p.style.name
    if '6.4' in text and style.startswith('Heading'):
        in_64 = True
    elif in_64 and style.startswith('Heading'):
        break
    elif in_64 and text:
        oral = ['还有一个问题', '理论上', '都能', '总的来说', '一个是', '另一个是', '有时候', '比如']
        found = [w for w in oral if w in text]
        if found:
            issues.append(found)
if issues:
    print(f"  [X] 仍有口语化: {issues}")
else:
    print(f"  [OK] 口语化已修复")

print("\n" + "=" * 60)
print("验证完成")
print("=" * 60)
