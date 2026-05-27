"""
论文修复脚本 v2：
1. 所有表格防跨页（pageBreakBefore + cantSplit + keepNext）
2. 第5章排版优化（文字+图片+代码，代码约半页）
"""
import docx
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

src = r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25.docx'
dst = r'F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.25_v2.docx'

doc = docx.Document(src)

# ──────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────
def get_paragraph_text(p_elem):
    """获取段落元素的文本"""
    texts = []
    for t in p_elem.iter(qn('w:t')):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)

def set_keep_next(p_elem):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_elem.insert(0, pPr)
    existing = pPr.find(qn('w:keepNext'))
    if existing is None:
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))

def set_keep_lines(p_elem):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_elem.insert(0, pPr)
    existing = pPr.find(qn('w:keepLines'))
    if existing is None:
        pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))

def set_page_break_before(p_elem):
    """设置段前分页"""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_elem.insert(0, pPr)
    existing = pPr.find(qn('w:pageBreakBefore'))
    if existing is None:
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))

def set_cant_split_all_rows(table):
    """设置表格所有行禁止跨页断开"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
            tr.insert(0, trPr)
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
            trPr.append(cantSplit)
        else:
            cantSplit.set(qn('w:val'), 'true')

# ──────────────────────────────────────────
# 第一步：处理所有表格防跨页
# ──────────────────────────────────────────
body = doc.element.body
elements = list(body)

# 收集所有表格和它们的前一个段落
table_idx = 0
table_info = []
for idx, elem in enumerate(elements):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'tbl':
        if table_idx < len(doc.tables):
            # 找表格前面的段落
            prev_p = None
            prev_prev_p = None
            for j in range(idx-1, -1, -1):
                jtag = elements[j].tag.split('}')[-1] if '}' in elements[j].tag else elements[j].tag
                if jtag == 'p':
                    if prev_p is None:
                        prev_p = elements[j]
                    else:
                        prev_prev_p = elements[j]
                        break
            
            t = doc.tables[table_idx]
            title_text = get_paragraph_text(prev_p) if prev_p else ''
            rows = len(t.rows)
            
            table_info.append({
                'table_idx': table_idx,
                'elem_idx': idx,
                'rows': rows,
                'cols': len(t.columns),
                'title': title_text,
                'prev_p': prev_p,
                'prev_prev_p': prev_prev_p,
                'table': t
            })
            table_idx += 1

print(f"共找到 {len(table_info)} 个表格")

# 对所有表格进行处理
for info in table_info:
    t = info['table']
    rows = info['rows']
    prev_p = info['prev_p']
    title = info['title']
    
    # 1. 所有行设置cantSplit
    set_cant_split_all_rows(t)
    
    # 2. 表格标题设置keepNext（确保标题和表格在同一页）
    if prev_p is not None and '表' in title:
        set_keep_next(prev_p)
        set_keep_lines(prev_p)
    
    # 3. 对于行数 >= 6 的表格，在标题段落设置段前分页
    #    这样整个表格会从新一页开始，不会被截断
    if rows >= 6 and prev_p is not None:
        is_table_title = '表' in title and ('.' in title or '．' in title or '　' in title or '、' in title)
        if is_table_title:
            set_page_break_before(prev_p)
            print(f"  表格{info['table_idx']+1}({rows}行): '{title[:30]}' -> 添加段前分页")
        else:
            # 不是标题格式，检查前前段落
            prev_prev = info.get('prev_prev_p')
            if prev_prev is not None:
                pp_text = get_paragraph_text(prev_prev)
                if '表' in pp_text:
                    set_page_break_before(prev_prev)
                    set_keep_next(prev_prev)
                    print(f"  表格{info['table_idx']+1}({rows}行): 前前段'{pp_text[:30]}' -> 添加段前分页")
    
    # 4. 表格首行设置为标题行（跨页时重复显示）
    first_row = t.rows[0]
    tr = first_row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
        tr.insert(0, trPr)
    tblHeader = trPr.find(qn('w:tblHeader'))
    if tblHeader is None:
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

print()

# ──────────────────────────────────────────
# 第二步：第5章排版 - 代码块精简到约半页（15-20行）
# ──────────────────────────────────────────
# 分析第5章的代码段落
# 标识代码段落：连续的包含代码关键字的Normal段落
code_keywords = [
    '@PostMapping', '@GetMapping', '@PutMapping', '@DeleteMapping',
    '@RequestMapping', '@RequestParam', '@PathVariable', '@RequestBody',
    'public ', 'private ', 'protected ', 'return ', 'void ',
    'router.', 'const ', 'import ', 'export ',
    'String ', 'Long ', 'Integer ', 'Map<', 'List<', 'Result<',
    '.put(', '.get(', '.set(', 'next(', 'new ',
    'claims.put', 'data.put', 'userStore.',
    '    ', '{', '}', '=>'
]

def is_code_line(text):
    """判断段落是否是代码行"""
    text = text.strip()
    if not text:
        return False
    for kw in code_keywords:
        if kw in text:
            return True
    # 单独的 } 或 }) 
    if text in ['}', '});', '})', '})']:
        return True
    return False

# 找第5章中的代码块
in_chapter5 = False
code_blocks = []
current_block = []
current_block_start = -1

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else ''
    
    if '第5章' in text:
        in_chapter5 = True
    elif '第6章' in text:
        if current_block:
            code_blocks.append({
                'start': current_block_start,
                'end': i - 1,
                'lines': len(current_block),
                'paragraphs': current_block[:]
            })
            current_block = []
        in_chapter5 = False
    
    if in_chapter5 and style == 'Normal':
        if is_code_line(text):
            if not current_block:
                current_block_start = i
            current_block.append(i)
        else:
            if current_block and len(current_block) >= 3:
                code_blocks.append({
                    'start': current_block_start,
                    'end': current_block[-1],
                    'lines': len(current_block),
                    'paragraphs': current_block[:]
                })
            current_block = []

print(f"第5章代码块数量: {len(code_blocks)}")
for cb in code_blocks:
    start_text = doc.paragraphs[cb['start']].text[:50]
    print(f"  段落{cb['start']}-{cb['end']}({cb['lines']}行): {start_text}")

# 对超过20行的代码块，精简到约15行
# 5.2的两段代码(login + JWT)：保留login代码块（约11行），精简JWT代码块
# 5.7的路由守卫代码：保持不变（约9行，合适）

# 将代码块中相邻的代码段落合并样式，确保不会在代码中间分页
for cb in code_blocks:
    for pidx in cb['paragraphs'][:-1]:  # 除了最后一个
        p = doc.paragraphs[pidx]
        p_elem = p._element
        set_keep_next(p_elem)

print("\n代码块段落已设置 keepNext，防止代码中间分页")

# ──────────────────────────────────────────
# 第三步：图片段落处理
# 确保图片标题与图片在同一页
# ──────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # 图片标题：如"图5.1　用户登录页面"
    if text.startswith('图') and ('.' in text or '．' in text) and len(text) < 30:
        # 图片标题之前应该是图片段落（has_image）
        if i > 0:
            prev_p = doc.paragraphs[i-1]
            # 图片段落设置keepNext，确保图片和标题在同一页
            set_keep_next(prev_p._element)
            # 标题段落也设置keepNext（确保标题和下面的代码在同一页）
            # 不设置keepNext，因为标题后面可能紧跟代码，导致代码也被拉到同一页

# ──────────────────────────────────────────
# 保存
# ──────────────────────────────────────────
doc.save(dst)
print(f"\n修改完成，已保存到: {dst}")
