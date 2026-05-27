# -*- coding: utf-8 -*-
"""
修复交换后的第5章图片引用。
交换后正确的对应关系:
  图5.1 商品管理 → image12.png (rId31)
  图5.2 数据统计 → image16.png (rId35)
  图5.3 用户登录 → image11.png (rId30)
  图5.4 工作人员后台 → image14.png (rId33)
  图5.5 工作人员模块 → image15.png (rId34)
  图5.6 教练预约 → image13.png (rId32)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document('5.26.2.docx')
rels = doc.part.rels

correct_rids = ['rId31', 'rId35', 'rId30', 'rId33', 'rId34', 'rId32']
correct_names = ['image12(商品管理)', 'image16(数据统计)', 'image11(用户登录)',
                 'image14(工作人员后台)', 'image15(工作人员模块)', 'image13(教练预约)']

in_ch5 = False
img_idx = 0
changes = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if not in_ch5:
        continue
    
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            old_rid = blip.get(qn('r:embed'))
            new_rid = correct_rids[img_idx]
            old_target = rels[old_rid].target_ref if old_rid in rels else '?'
            new_target = rels[new_rid].target_ref if new_rid in rels else '?'
            
            if old_rid != new_rid:
                blip.set(qn('r:embed'), new_rid)
                changes += 1
                print(f"图{img_idx+1}: {old_rid}({old_target}) -> {new_rid}({new_target}) = {correct_names[img_idx]}")
            else:
                print(f"图{img_idx+1}: 已正确 {old_rid}({old_target})")
            
            img_idx += 1

doc.save('5.26.2.docx')
print(f"\n共修复 {changes} 张图片, 已保存")

# 验证
print("\n=== 验证 ===")
doc2 = Document('5.26.2.docx')
rels2 = doc2.part.rels
in_ch5 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and '第5章' in text:
        in_ch5 = True
    elif in_ch5 and style == 'Heading 1':
        break
    if not in_ch5:
        continue
    drawings = p._element.findall('.//' + qn('w:drawing'))
    for d in drawings:
        blips = d.findall('.//' + qn('a:blip'))
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            target = rels2[embed].target_ref if embed in rels2 else '?'
            next_text = doc2.paragraphs[i+1].text.strip()[:50] if i+1 < len(doc2.paragraphs) else ''
            print(f"  {embed} -> {target}  | {next_text}")
