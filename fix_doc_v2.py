# -*- coding: utf-8 -*-
"""
5.26.1.docx modification script - v2 (fixed chapter boundaries)
"""
import copy
import re
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips
from lxml import etree

SRC = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.1.docx"
DST = r"F:\26毕设单\基于springboot的飞跃滑雪场管理系统的设计与实现\5.26.2.docx"

doc = Document(SRC)
body = doc.element.body


def para_replace(para, old, new):
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not para.runs:
        return False
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ""
    return True


def set_para_spacing(para, before_pt=0, after_pt=0):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))


def find_heading_para(text_pattern, style_prefix="Heading"):
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name.startswith(style_prefix):
            if text_pattern in p.text:
                return i
    return -1


def collect_elements_between_paras(start_para_idx, end_para_idx):
    start_elem = doc.paragraphs[start_para_idx]._element
    end_elem = doc.paragraphs[end_para_idx]._element
    collecting = False
    elements = []
    for child in list(body):
        if child is start_elem:
            collecting = True
        if collecting:
            elements.append(child)
        if child is end_elem:
            break
    return elements


def get_section_elements(heading_idx, next_heading_idx):
    start_elem = doc.paragraphs[heading_idx]._element
    end_limit = doc.paragraphs[next_heading_idx]._element
    elements = []
    collecting = False
    for child in list(body):
        if child is start_elem:
            collecting = True
        if collecting:
            if child is end_limit:
                break
            elements.append(child)
    return elements


# ============================================================
# STEP 1: 6.4 口语化修改
# ============================================================
print("=== Step 1: 6.4 ===")
p408 = doc.paragraphs[408]
para_replace(p408, "还有一个问题是教练预约", "此外，教练预约")
para_replace(p408, "理论上不应该允许", "不应允许")
print("  P408 done")

p409 = doc.paragraphs[409]
para_replace(p409, "还有一个在测试中发现的问题是关于并发的。", "此外，在并发控制方面也存在待优化之处。")
para_replace(p409, "理论上应该在数据库层面用乐观锁来处理", "应在数据库层面采用乐观锁机制加以处理")
para_replace(p409, "不过因为", "由于")
para_replace(p409, "所以这个问题在当前阶段影响不大", "该问题在当前阶段影响有限")
para_replace(p409, "但如果以后用户量增加的话，这个地方是需要优化的", "但随着用户规模扩大，该部分仍有优化的必要")
print("  P409 done")

p410 = doc.paragraphs[410]
para_replace(p410, "都能正常走通", "均可正常运行")
print("  P410 done")

# ============================================================
# STEP 2: 7.1 修改
# ============================================================
print("=== Step 2: 7.1 ===")
p415 = doc.paragraphs[415]
para_replace(p415, "笔者收获颇多", "收获颇多")
para_replace(p415, "这些在小型练习中难以遇到", "这些在课程中难以遇到")
print("  P415 done")

# ============================================================
# STEP 3: 参考文献重排
# ============================================================
print("=== Step 3: refs ===")

ref_map = {
    1: 1, 2: 2, 3: 3, 4: 4, 5: 5,
    6: 15, 7: 6, 8: 7, 9: 8, 10: 9,
    11: 10, 12: 11, 13: 12, 14: 13, 15: 14,
}

for i in range(56, 422):
    p = doc.paragraphs[i]
    text = p.text
    if not re.search(r'\[\d+\]', text):
        continue
    new_text = text
    for old_num in sorted(ref_map.keys(), reverse=True):
        new_text = new_text.replace(f'[{old_num}]', f'[###{ref_map[old_num]}###]')
    new_text = re.sub(r'\[###(\d+)###\]', r'[\1]', new_text)
    if new_text != text and p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""

# Remove duplicate citations
p75 = doc.paragraphs[75]
para_replace(p75, "[6][4]", "[6]")
p88 = doc.paragraphs[88]
para_replace(p88, "[5][12]", "[12]")

# Reorder reference list
ref_entries = {}
for i in range(423, 438):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    text = p.text
    match = re.match(r'\[(\d+)\]', text)
    if match:
        old_num = int(match.group(1))
        new_num = ref_map.get(old_num, old_num)
        new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', text)
        ref_entries[new_num] = new_text

sorted_nums = sorted(ref_entries.keys())
for idx, new_num in enumerate(sorted_nums):
    target_idx = 423 + idx
    if target_idx < len(doc.paragraphs):
        p = doc.paragraphs[target_idx]
        if p.runs:
            p.runs[0].text = ref_entries[new_num]
            for r in p.runs[1:]:
                r.text = ""

print("  refs reordered")

# ============================================================
# STEP 4: 三线表间距
# ============================================================
print("=== Step 4: table spacing ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if re.match(r'^(表[46]\.\d|续表[46]\.\d)', text):
        set_para_spacing(p, before_pt=3, after_pt=3)
print("  done")

# ============================================================
# STEP 5: 3.4 合并到 3.2
# ============================================================
print("=== Step 5: merge 3.4 into 3.2 ===")

move_map = [
    (103, 118, 122),
    (105, 123, 126),
    (107, 127, 130),
    (109, 131, 134),
]

moves_data = []
for target_para_idx, src_start, src_end in move_map:
    elems = collect_elements_between_paras(src_start, src_end)
    elems_to_move = elems[1:]
    moves_data.append((target_para_idx, elems_to_move))

for target_para_idx, elems_to_move in reversed(moves_data):
    target_elem = doc.paragraphs[target_para_idx]._element
    for elem in reversed(elems_to_move):
        target_elem.addnext(copy.deepcopy(elem))

elems_to_remove = collect_elements_between_paras(116, 134)
for elem in elems_to_remove:
    body.remove(elem)

print("  done")

# ============================================================
# STEP 6: 4.5 移到 4.4 前面
# ============================================================
print("=== Step 6: swap 4.4/4.5 ===")

h44_idx = find_heading_para("4.4")
h45_idx = find_heading_para("4.5")
h46_idx = find_heading_para("4.6")

if h44_idx >= 0 and h45_idx >= 0 and h46_idx >= 0:
    # Move flow paragraphs from 4.4 end to before 4.5's diagrams
    flow_para_idxs = []
    for i in range(h44_idx, h45_idx):
        p = doc.paragraphs[i]
        if any(kw in p.text for kw in ["商品租赁的整体业务流程", "用户登录的处理流程", "在完成功能模块划分"]):
            flow_para_idxs.append(i)

    # Collect 4.5 section elements
    sec45_elems = collect_elements_between_paras(h45_idx, h46_idx - 1)

    # Insert 4.5 before 4.4
    h44_elem = doc.paragraphs[h44_idx]._element
    for elem in reversed(sec45_elems):
        h44_elem.addprevious(copy.deepcopy(elem))
    for elem in sec45_elems:
        body.remove(elem)

    # Now move flow paragraphs to after new 4.5 heading (which is now before 4.4)
    h45_new = find_heading_para("4.5")
    if h45_new >= 0 and flow_para_idxs:
        # Re-find flow paragraphs (indices changed)
        new_flow_idxs = []
        for i, p in enumerate(doc.paragraphs):
            if any(kw in p.text for kw in ["商品租赁的整体业务流程", "用户登录的处理流程", "在完成功能模块划分"]):
                new_flow_idxs.append(i)

        h45_elem_new = doc.paragraphs[h45_new]._element
        for idx in reversed(new_flow_idxs):
            elem = doc.paragraphs[idx]._element
            h45_elem_new.addnext(copy.deepcopy(elem))
            body.remove(elem)

    # Swap section numbers
    h45_final = find_heading_para("4.5")
    h44_final = find_heading_para("4.4")
    if h45_final >= 0 and h44_final >= 0 and h45_final < h44_final:
        p45 = doc.paragraphs[h45_final]
        p44 = doc.paragraphs[h44_final]
        para_replace(p45, "4.5", "4.4")
        para_replace(p44, "4.4", "4.5")

    print("  done")

# ============================================================
# STEP 7: 第5章按角色重组
# ============================================================
print("=== Step 7: Ch5 reorg ===")

# Find Ch5 and Ch6 boundaries
ch5_start = find_heading_para("第5章")
ch6_start = find_heading_para("第6章")

if ch5_start < 0 or ch6_start < 0:
    print("  WARNING: cannot find Ch5/Ch6 boundaries")
else:
    # Find all Heading 2 within Ch5 (between ch5_start and ch6_start)
    ch5_h2 = {}
    for i in range(ch5_start, ch6_start):
        p = doc.paragraphs[i]
        if p.style and p.style.name == "Heading 2":
            text = p.text
            for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
                if key in text:
                    ch5_h2[key] = i
                    break

    print(f"  Ch5 sections: {ch5_h2}")

    if all(k in ch5_h2 for k in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]):
        # Move 5.3 (商品管理) after 5.1, before 5.2
        sec53_elems = get_section_elements(ch5_h2["5.3"], ch5_h2["5.4"])
        sec52_elem = doc.paragraphs[ch5_h2["5.2"]]._element
        for elem in reversed(sec53_elems):
            sec52_elem.addprevious(copy.deepcopy(elem))
        for elem in sec53_elems:
            body.remove(elem)
        print("  moved 5.3 after 5.1")

        # Re-find headings
        ch5_h2_2 = {}
        for i in range(ch5_start, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            if p.style and p.style.name == "Heading 1" and "第6章" in p.text:
                break
            if p.style and p.style.name == "Heading 2":
                for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
                    if key in p.text and key not in ch5_h2_2:
                        ch5_h2_2[key] = i
                        break

        print(f"  After 5.3 move: {ch5_h2_2}")

        # Move 5.6 (数据统计) after the moved 5.3, before 5.2
        if "5.6" in ch5_h2_2 and "5.2" in ch5_h2_2 and "5.7" in ch5_h2_2:
            sec56_elems = get_section_elements(ch5_h2_2["5.6"], ch5_h2_2["5.7"])
            user_elem = doc.paragraphs[ch5_h2_2["5.2"]]._element
            for elem in reversed(sec56_elems):
                user_elem.addprevious(copy.deepcopy(elem))
            for elem in sec56_elems:
                body.remove(elem)
            print("  moved 5.6 after 5.3")

        # Re-find headings again
        ch5_final = {}
        for i in range(ch5_start, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            if p.style and p.style.name == "Heading 1" and "第6章" in p.text:
                break
            if p.style and p.style.name == "Heading 2":
                for key in ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7"]:
                    if key in p.text and key not in ch5_final:
                        ch5_final[key] = i
                        break

        print(f"  Final order: {ch5_final}")

        # Rename headings
        # Current order should be: 5.1(管理员) 5.3(商品) 5.6(数据) 5.2(用户) 5.4(教练) 5.5(课程/工作人员) 5.7(安全)
        rename_map = {}
        if "5.1" in ch5_final:
            rename_map[ch5_final["5.1"]] = ("5.1　管理员模块的实现", "Heading 2")
        if "5.3" in ch5_final:
            rename_map[ch5_final["5.3"]] = ("5.1.1　商品管理功能", "Heading 3")
        if "5.6" in ch5_final:
            rename_map[ch5_final["5.6"]] = ("5.1.2　数据统计功能", "Heading 3")
        if "5.2" in ch5_final:
            rename_map[ch5_final["5.2"]] = ("5.2　用户模块的实现", "Heading 2")
        if "5.4" in ch5_final:
            rename_map[ch5_final["5.4"]] = ("5.3　教练模块的实现", "Heading 2")
        if "5.5" in ch5_final:
            rename_map[ch5_final["5.5"]] = ("5.4　工作人员模块的实现", "Heading 2")
        if "5.7" in ch5_final:
            rename_map[ch5_final["5.7"]] = ("5.5　系统安全实现", "Heading 2")

        for pi, (new_title, style_name) in rename_map.items():
            p = doc.paragraphs[pi]
            if p.runs:
                p.runs[0].text = new_title
                for r in p.runs[1:]:
                    r.text = ""
            if style_name == "Heading 3":
                p.style = doc.styles["Heading 3"]
            print(f"  P{pi}: -> {new_title}")

        # Add admin intro to 5.1
        if "5.1" in ch5_final:
            intro_para = doc.paragraphs[ch5_final["5.1"] + 1]
            old_text = intro_para.text
            new_intro = "管理员是系统的最高权限角色，负责系统整体的数据管理和运营配置。在介绍管理员的具体功能之前，先简要说明系统的开发环境。" + old_text
            if intro_para.runs:
                intro_para.runs[0].text = new_intro
                for r in intro_para.runs[1:]:
                    r.text = ""

        # Add user intro to 5.2
        if "5.2" in ch5_final:
            intro_para = doc.paragraphs[ch5_final["5.2"] + 1]
            old_text = intro_para.text
            new_intro = "普通用户是系统的最终使用者，通过前台页面访问各项服务功能。" + old_text
            if intro_para.runs:
                intro_para.runs[0].text = new_intro
                for r in intro_para.runs[1:]:
                    r.text = ""

        # Fix 5.4 (原课程与场地→工作人员) intro
        if "5.5" in ch5_final:
            intro_para = doc.paragraphs[ch5_final["5.5"] + 1]
            # Keep original content but might need adjustment
            pass

    print("  Ch5 reorg done")

# ============================================================
# STEP 8: 第5章精简文字
# ============================================================
print("=== Step 8: Ch5 text reduction ===")

for i, p in enumerate(doc.paragraphs):
    text = p.text
    if "工作人员在系统中的操作权限主要包括以下几个方面" in text:
        if p.runs:
            p.runs[0].text = "工作人员的主要操作权限涵盖租赁订单管理、商品信息查看、场地状态管理等方面。"
            for r in p.runs[1:]:
                r.text = ""

    if "在滑雪场的日常运营管理中，工作人员承担着重要的前台业务职责" in text:
        if p.runs:
            p.runs[0].text = "工作人员在系统中主要负责前台业务操作，登录后台后可使用商品信息管理、租赁订单处理及场地状态查看等功能。"
            for r in p.runs[1:]:
                r.text = ""

    if "在页面布局方面，工作人员后台界面采用" in text:
        if p.runs:
            p.runs[0].text = "工作人员后台采用左侧导航加右侧内容区域的布局方式，菜单项根据角色权限进行了精简。商品管理页面提供按名称和分类检索的功能，租赁订单页面支持按状态筛选并可直接执行归还操作。"
            for r in p.runs[1:]:
                r.text = ""

print("  done")

# ============================================================
# STEP 9: Update figure/table numbers affected by 4.4/4.5 swap
# ============================================================
print("=== Step 9: fix numbering ===")

# After swapping 4.4 and 4.5:
# Old 4.5 graphs (图4.3 登录流程, 图4.4 商品租赁流程) are now in section 4.4
# Old 4.4 graph (图4.2 E-R图) is now in section 4.5
# But the figure numbers should stay the same based on appearance order

# Fix table reference numbers in the text
# 原4.4中的表(表4.2-4.7) 现在在4.5中
# 原4.5中没有表
# 所以表编号不需要改

# Fix 4.4 subsection headings (数据库设计的子标题 4.4.1, 4.4.2 → 4.5.1, 4.5.2)
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 3":
        if "4.4.1" in p.text:
            para_replace(p, "4.4.1", "4.5.1")
        elif "4.4.2" in p.text:
            para_replace(p, "4.4.2", "4.5.2")

print("  done")

# ============================================================
# Save
# ============================================================
print("=== Saving ===")
doc.save(DST)
print(f"Saved to {DST}")
print("DONE")
